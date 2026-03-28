import customtkinter as ctk
from PIL import Image, ImageTk
from tkinter import Tk, ttk, StringVar, PhotoImage
from tkinter import font
import customtkinter as ctk
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from tkinter import filedialog
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from scipy.interpolate import make_interp_spline
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import font
from PIL import ImageFont, ImageTk
from docx2pdf import convert
import os
from matplotlib.backends.backend_pdf import PdfPages 
from tkinter.filedialog import asksaveasfilename
from mpl_toolkits.mplot3d import Axes3D 
import tempfile
import os


# Инициализация приложения
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("green")

root = ctk.CTk()
root.title('Расчет газовой турбины')
root.geometry('1280x800')
root.configure(bg="#1a1a1a", fg_color="#1a1a1a")

# Загрузка шрифта через Pillow (без установки)
try:
    # Создаем временный объект шрифта
    custom_font = ImageFont.truetype("coolveticaRg.ttf", size=16)
    # Регистрируем его в tkinter
    tk_font = font.Font(
        family=custom_font.getname()[0],  # Получаем имя шрифта
        size=16
    )
except Exception as e:
    print(f"Ошибка загрузки шрифта: {e}")
    tk_font = ("Arial", 16)  # Fallback на стандартный шрифт

tk_font1 = ctk.CTkFont(size=25, weight='bold')

# Установка иконки программы
try:
    logo_image = Image.open("logo1.png")
    photo = ImageTk.PhotoImage(logo_image)
    root.after(201, lambda: root.iconphoto(False, photo))
except FileNotFoundError:
    print("Логотип не найден.")
except Exception as e:
    print(f"Ошибка при загрузке логотипа: {e}")

# Загрузка изображений
try:
    menu_icon = ctk.CTkImage(light_image=Image.open("menu_d.png"), dark_image=Image.open("menu.png"), size=(40, 40))
    menu_white_icon = ctk.CTkImage(light_image=Image.open("menu_white.png"), dark_image=Image.open("menu_white.png"), size=(40, 40))

    save_menu_icon = ctk.CTkImage(light_image=Image.open("save_menu_d.png"), dark_image=Image.open("save_menu.png"), size=(40, 40))
    save_menu_white_icon = ctk.CTkImage(light_image=Image.open("save_menu_active.png"), dark_image=Image.open("save_menu_active.png"), size=(40, 40))

    calc_icon = ctk.CTkImage(light_image=Image.open("calc_d.png"), dark_image=Image.open("calc.png"), size=(40, 40))
    grafik_icon = ctk.CTkImage(light_image=Image.open("grafik_d.png"), dark_image=Image.open("grafik.png"), size=(40, 40))
    settings_icon = ctk.CTkImage(light_image=Image.open("settings_d.png"), dark_image=Image.open("settings.png"), size=(40, 40))
    calc_white_icon = ctk.CTkImage(light_image=Image.open("calc_white.png"), dark_image=Image.open("calc_white.png"), size=(40, 40))
    grafik_white_icon = ctk.CTkImage(light_image=Image.open("grafik_white.png"), dark_image=Image.open("grafik_white.png"), size=(40, 40))
    settings_white_icon = ctk.CTkImage(light_image=Image.open("settings_white.png"), dark_image=Image.open("settings_white.png"), size=(40, 40))
    word_icon = ctk.CTkImage(light_image=Image.open("word.png"), dark_image=Image.open("word.png"), size=(20, 20))
    pdf_icon = ctk.CTkImage(light_image=Image.open("pdf.png"), dark_image=Image.open("pdf.png"), size=(20, 20))

    func_icon = ctk.CTkImage(light_image=Image.open("func.png"), dark_image=Image.open("func.png"), size=(20, 20))
    d_curve_icon = ctk.CTkImage(light_image=Image.open("3d_curve.png"), dark_image=Image.open("3d_curve.png"), size=(20, 20))

    trash_icon = ctk.CTkImage(light_image=Image.open("mush.png"), dark_image=Image.open("mush.png"), size=(20, 20))
    cs_icon = ctk.CTkImage(light_image=Image.open("cs.black.png"), dark_image=Image.open("cs.black.png"), size=(20, 20))
    inp_icon = ctk.CTkImage(light_image=Image.open("strelka.white.png"), dark_image=Image.open("strelka.white.png"), size=(20, 20))
    get_icon = ctk.CTkImage(light_image=Image.open("start.black.png"), dark_image=Image.open("start.black.png"), size=(45, 25))
    galka_icon = ctk.CTkImage(light_image=Image.open("galka.png"), dark_image=Image.open("galka.png"), size=(20, 20))
    galka_active_icon = ctk.CTkImage(light_image=Image.open("galka_active.png"), dark_image=Image.open("galka_active.png"), size=(20, 20))

     # Добавляем новые иконки для графиков
    graph_icon = ctk.CTkImage(light_image=Image.open("fi_pk_d.png"), dark_image=Image.open("fi_pk.png"), size=(100, 100))
    graph2_icon = ctk.CTkImage(light_image=Image.open("kpd_e_d.png"), dark_image=Image.open("kpd_e.png"), size=(100, 100))
    graph3_icon = ctk.CTkImage(light_image=Image.open("he_pk_d.png"), dark_image=Image.open("he_pk.png"), size=(100, 100))
    graph4_icon = ctk.CTkImage(light_image=Image.open("kpd_he_d.png"), dark_image=Image.open("kpd_he.png"), size=(100, 100))

    widget_icon = ctk.CTkImage(light_image=Image.open("widget_d.png"), dark_image=Image.open("widget.png"), size=(160, 160))
    widget_active_icon = ctk.CTkImage(light_image=Image.open("widget_d_active.png"), dark_image=Image.open("widget_active.png"), size=(160, 160))

    pdf_save_icon = ctk.CTkImage(light_image=Image.open("save_pdf_d.png"), dark_image=Image.open("save_pdf.png"), size=(200, 200))
    word_save_icon = ctk.CTkImage(light_image=Image.open("save_d.png"), dark_image=Image.open("save.png"), size=(200, 200))

    graph1_1_icon = ctk.CTkImage(light_image=Image.open("fi_pk1.png"), dark_image=Image.open("fi_pk1.png"), size=(100, 100))
    graph2_2_icon = ctk.CTkImage(light_image=Image.open("kpd_e1.png"), dark_image=Image.open("kpd_e1.png"), size=(100, 100))
    graph3_3_icon = ctk.CTkImage(light_image=Image.open("he_pk1.png"), dark_image=Image.open("he_pk1.png"), size=(100, 100))
    graph4_4_icon = ctk.CTkImage(light_image=Image.open("kpd_he1.png"), dark_image=Image.open("kpd_he1.png"), size=(100, 100))

    tema_icon = ctk.CTkImage(light_image=Image.open("tema_night.png"), dark_image=Image.open("tema_night_active.png"), size=(280, 180))
    tema1_icon = ctk.CTkImage(light_image=Image.open("tema_d_active.png"), dark_image=Image.open("tema_d.png"), size=(280, 180))
except FileNotFoundError as e:
    print(f"Изображение не найдено: {e}")



# Левый фрейм меню
left_frame = ctk.CTkFrame(root, width=70, fg_color="#151515")
left_frame.grid(row=0, column=0, sticky="nsew", rowspan=10)
left_frame.grid_propagate(False)

# Настройка grid для left_frame
left_frame.grid_rowconfigure(0, weight=0)  # Кнопка меню
left_frame.grid_rowconfigure(1, weight=0)  # Кнопка расчета
left_frame.grid_rowconfigure(2, weight=0)  # Кнопка графиков
left_frame.grid_rowconfigure(3, weight=0)  # Кнопка сохранения
left_frame.grid_rowconfigure(4, weight=1)  # Пустое пространство
left_frame.grid_rowconfigure(5, weight=0)  # Кнопка настроек
left_frame.grid_columnconfigure(0, weight=1)

root.grid_rowconfigure(0, weight=1)
root.grid_columnconfigure(0, weight=0)
root.grid_columnconfigure(1, weight=1)

# Основная область
main_frame = ctk.CTkFrame(root, fg_color="#1f1f1f")
main_frame.grid(row=0, column=1, sticky="nsew")

main_frame.grid_rowconfigure(0, weight=1)
main_frame.grid_columnconfigure(0, weight=1)

# Фреймы для контента
frame_calc = ctk.CTkFrame(main_frame, fg_color="#1f1f1f")
# frame_calc.grid(row=0, column=0, sticky="nsew", padx=(250, 0), pady=(10,0))
frame_grafik = ctk.CTkFrame(main_frame, fg_color="#1f1f1f")
# frame_grafik.grid(row=0, column=0, sticky="nsew", padx=(250, 0), pady=(10,0))
frame_settings = ctk.CTkFrame(main_frame, fg_color="#1f1f1f")

frame_save = ctk.CTkFrame(main_frame, fg_color="#1f1f1f")

# Функция переключения меню
def toggle_menu():
    current_width = left_frame.cget("width")
    new_width = 210 if current_width < 100 else 70
    left_frame.configure(width=new_width)
    
    # Меняем иконку меню в зависимости от состояния
    if new_width == 210:
        menu_btn.configure(image=menu_white_icon)
    else:
        menu_btn.configure(image=menu_icon)
    
    for btn in [calc_btn, grafik_btn, settings_btn, save_btn]:
        if new_width == 70:
            btn.configure(text="", anchor="w", width=400, height=40)
            if btn == calc_btn and calc_icon:
                btn.configure(image=calc_icon if not getattr(btn, 'active', False) else calc_white_icon, anchor="c")
            elif btn == grafik_btn and grafik_icon:
                btn.configure(image=grafik_icon if not getattr(btn, 'active', False) else grafik_white_icon, anchor="c")
            elif btn == settings_btn and settings_icon:
                btn.configure(image=settings_icon if not getattr(btn, 'active', False) else settings_white_icon, anchor="c")
            elif btn == save_btn and save_menu_icon:
                btn.configure(image=save_menu_icon if not getattr(btn, 'active', False) else save_menu_white_icon, anchor="c")
        else:
            btn.configure(anchor="w", text=btn.full_text)
            if btn == calc_btn and calc_icon:
                btn.configure(image=calc_icon if not getattr(btn, 'active', False) else calc_white_icon, anchor="w")
            elif btn == grafik_btn and grafik_icon:
                btn.configure(image=grafik_icon if not getattr(btn, 'active', False) else grafik_white_icon, anchor="w")
            elif btn == settings_btn and settings_icon:
                btn.configure(image=settings_icon if not getattr(btn, 'active', False) else settings_white_icon, anchor="w")
            elif btn == save_btn and save_menu_icon:
                btn.configure(image=save_menu_icon if not getattr(btn, 'active', False) else save_menu_white_icon, anchor="w")

# Функции переключения с изменением цвета иконки
def show_calc():
    frame_grafik.grid_forget()
    frame_settings.grid_forget()
    frame_save.grid_forget()
    frame_calc.grid(row=0, column=0, sticky="nsew")
    
    # Get current theme
    current_theme = ctk.get_appearance_mode()
    
    # Set active state for calc button
    calc_btn.active = True
    calc_btn.configure(text_color="white")
    if calc_white_icon:
        current_width = left_frame.cget("width")
        if current_width == 70:
            calc_btn.configure(image=calc_white_icon, anchor="c")
        else:
            calc_btn.configure(image=calc_white_icon, anchor="w")
    
    # Reset state for other buttons
    grafik_btn.active = False
    settings_btn.active = False
    save_btn.active = False
    
    # Set colors based on theme
    if current_theme == "Dark":
        grafik_btn.configure(text_color="#6a6a6a")
        settings_btn.configure(text_color="#6a6a6a")
        save_btn.configure(text_color="#6a6a6a")
        if grafik_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                grafik_btn.configure(image=grafik_icon, anchor="c")
            else:
                grafik_btn.configure(image=grafik_icon, anchor="w")
        if settings_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                settings_btn.configure(image=settings_icon, anchor="c")
            else:
                settings_btn.configure(image=settings_icon, anchor="w")
        if save_menu_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                save_btn.configure(image=save_menu_icon, anchor="c")
            else:
                save_btn.configure(image=save_menu_icon, anchor="w")
    else:
        grafik_btn.configure(text_color="#a0a4ac")
        settings_btn.configure(text_color="#a0a4ac")
        save_btn.configure(text_color="#a0a4ac")
        if grafik_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                grafik_btn.configure(image=grafik_icon, anchor="c")
            else:
                grafik_btn.configure(image=grafik_icon, anchor="w")
        if settings_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                settings_btn.configure(image=settings_icon, anchor="c")
            else:
                settings_btn.configure(image=settings_icon, anchor="w")
        if save_menu_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                save_btn.configure(image=save_menu_icon, anchor="c")
            else:
                save_btn.configure(image=save_menu_icon, anchor="w")

def show_grafik():
    frame_calc.grid_forget()
    frame_settings.grid_forget()
    frame_save.grid_forget()
    frame_grafik.grid(row=0, column=0, sticky="nsew")
    
    # Get current theme
    current_theme = ctk.get_appearance_mode()
    
    # Set active state for grafik button
    grafik_btn.active = True
    grafik_btn.configure(text_color="white")
    if grafik_white_icon:
        current_width = left_frame.cget("width")
        if current_width == 70:
            grafik_btn.configure(image=grafik_white_icon, anchor="c")
        else:
            grafik_btn.configure(image=grafik_white_icon, anchor="w")
    
    # Reset state for other buttons
    calc_btn.active = False
    settings_btn.active = False
    save_btn.active = False
    
    # Set colors based on theme
    if current_theme == "Dark":
        calc_btn.configure(text_color="#6a6a6a")
        settings_btn.configure(text_color="#6a6a6a")
        save_btn.configure(text_color="#6a6a6a")
        if calc_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                calc_btn.configure(image=calc_icon, anchor="c")
            else:
                calc_btn.configure(image=calc_icon, anchor="w")
        if settings_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                settings_btn.configure(image=settings_icon, anchor="c")
            else:
                settings_btn.configure(image=settings_icon, anchor="w")
        if save_menu_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                save_btn.configure(image=save_menu_icon, anchor="c")
            else:
                save_btn.configure(image=save_menu_icon, anchor="w")
    else:
        calc_btn.configure(text_color="#a0a4ac")
        settings_btn.configure(text_color="#a0a4ac")
        save_btn.configure(text_color="#a0a4ac")
        if calc_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                calc_btn.configure(image=calc_icon, anchor="c")
            else:
                calc_btn.configure(image=calc_icon, anchor="w")
        if settings_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                settings_btn.configure(image=settings_icon, anchor="c")
            else:
                settings_btn.configure(image=settings_icon, anchor="w")
        if save_menu_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                save_btn.configure(image=save_menu_icon, anchor="c")
            else:
                save_btn.configure(image=save_menu_icon, anchor="w")



def show_settings():
    frame_calc.grid_forget()
    frame_grafik.grid_forget()
    frame_save.grid_forget()
    frame_settings.grid(row=0, column=0, sticky="nsew")
    
    # Get current theme
    current_theme = ctk.get_appearance_mode()
    
    # Set active state for settings button
    settings_btn.active = True
    settings_btn.configure(text_color="white")
    if settings_white_icon:
        current_width = left_frame.cget("width")
        if current_width == 70:
            settings_btn.configure(image=settings_white_icon, anchor="c")
        else:
            settings_btn.configure(image=settings_white_icon, anchor="w")
    
    # Reset state for other buttons
    calc_btn.active = False
    grafik_btn.active = False
    save_btn.active = False
    
    # Set colors based on theme
    if current_theme == "Dark":
        calc_btn.configure(text_color="#6a6a6a")
        grafik_btn.configure(text_color="#6a6a6a")
        save_btn.configure(text_color="#6a6a6a")
        if calc_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                calc_btn.configure(image=calc_icon, anchor="c")
            else:
                calc_btn.configure(image=calc_icon, anchor="w")
        if grafik_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                grafik_btn.configure(image=grafik_icon, anchor="c")
            else:
                grafik_btn.configure(image=grafik_icon, anchor="w")
        if save_menu_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                save_btn.configure(image=save_menu_icon, anchor="c")
            else:
                save_btn.configure(image=save_menu_icon, anchor="w")
    else:
        calc_btn.configure(text_color="#a0a4ac")
        grafik_btn.configure(text_color="#a0a4ac")
        save_btn.configure(text_color="#a0a4ac")
        if calc_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                calc_btn.configure(image=calc_icon, anchor="c")
            else:
                calc_btn.configure(image=calc_icon, anchor="w")
        if grafik_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                grafik_btn.configure(image=grafik_icon, anchor="c")
            else:
                grafik_btn.configure(image=grafik_icon, anchor="w")
        if save_menu_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                save_btn.configure(image=save_menu_icon, anchor="c")
            else:
                save_btn.configure(image=save_menu_icon, anchor="w")


def show_save():
    frame_calc.grid_forget()
    frame_settings.grid_forget()
    frame_grafik.grid_forget()
    frame_save.grid(row=0, column=0, sticky="nsew")
    
    # Get current theme
    current_theme = ctk.get_appearance_mode()
    
    # Set active state for grafik button
    save_btn.active = True
    save_btn.configure(text_color="white")
    if save_menu_white_icon:
        current_width = left_frame.cget("width")
        if current_width == 70:
            save_btn.configure(image=save_menu_white_icon, anchor="c")
        else:
            save_btn.configure(image=save_menu_white_icon, anchor="w")
    
    # Reset state for other buttons
    calc_btn.active = False
    grafik_btn.active = False
    settings_btn.active = False
    
    # Set colors based on theme
    if current_theme == "Dark":
        calc_btn.configure(text_color="#6a6a6a")
        settings_btn.configure(text_color="#6a6a6a")
        grafik_btn.configure(text_color="#6a6a6a")
        if calc_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                calc_btn.configure(image=calc_icon, anchor="c")
            else:
                calc_btn.configure(image=calc_icon, anchor="w")
        if settings_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                settings_btn.configure(image=settings_icon, anchor="c")
            else:
                settings_btn.configure(image=settings_icon, anchor="w")
        if grafik_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                grafik_btn.configure(image=grafik_icon, anchor="c")
            else:
                grafik_btn.configure(image=grafik_icon, anchor="w")
    else:
        calc_btn.configure(text_color="#a0a4ac")
        settings_btn.configure(text_color="#a0a4ac")
        grafik_btn.configure(text_color="#a0a4ac")
        if calc_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                calc_btn.configure(image=calc_icon, anchor="c")
            else:
                calc_btn.configure(image=calc_icon, anchor="w")
        if settings_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                settings_btn.configure(image=settings_icon, anchor="c")
            else:
                settings_btn.configure(image=settings_icon, anchor="w")
        if grafik_icon:
            current_width = left_frame.cget("width")
            if current_width == 70:
                grafik_btn.configure(image=grafik_icon, anchor="c")
            else:
                grafik_btn.configure(image=grafik_icon, anchor="w")

# Инициализация кнопок с атрибутами
def create_nav_button(row, icon, white_icon, text, command):
    btn = ctk.CTkButton(
        left_frame, 
        image=icon, 
        text=None,
        command=command, 
        width=40,
        height=40,
        anchor="w",
        fg_color="transparent", 
        hover_color="#303030",
        text_color="#6a6a6a",
        corner_radius=5
    )
    btn.full_text = text
    btn.icon = icon
    btn.white_icon = white_icon
    btn.active = False
    btn.grid(row=row, column=0, sticky="nsew", pady=(0, 10), padx=0)
    return btn

def update_label_positions():
    # Проверяем, пуст ли текст в label36
    if not label36.cget("text"):
        # Если текст пуст, сдвигаем следующие элементы вверх
        label50.grid(row=8, pady=5)  # Сдвигаем label50 на позицию label36
        label25.grid(row=9, pady=5)
        label26.grid(row=10, pady=5)
        label27.grid(row=11, pady=5)
        label28.grid(row=12, pady=5)
        label29.grid(row=13, pady=5)
        label30.grid(row=14, pady=5)
        label31.grid(row=15, pady=5)
        label32.grid(row=16, pady=5)
        label33.grid(row=17, pady=5)
        label200.grid(row=20, pady=5)
        lb8.grid(row=18, pady=(15,5))
        separator2.grid(row=19, pady=2)
        label34.grid(row=21, pady=5)
        label35.grid(row=22, pady=5)
    else:
        # Если текст есть, возвращаем стандартное расположение
        label36.grid(row=8, pady=5)
        label50.grid(row=9, pady=5)
        label25.grid(row=10, pady=5)
        label26.grid(row=11, pady=5)
        label27.grid(row=12, pady=5)
        label28.grid(row=13, pady=5)
        label29.grid(row=14, pady=5)
        label30.grid(row=15, pady=5)
        label31.grid(row=16, pady=5)
        label32.grid(row=17, pady=5)
        label33.grid(row=18, pady=5)
        label200.grid(row=21, pady=5)
        lb8.grid(row=19, pady=(15,5))
        separator2.grid(row=20, pady=2)
        label34.grid(row=22, pady=5)
        label35.grid(row=23, pady=5)

# Кнопка меню
menu_btn = ctk.CTkButton(left_frame, 
                        image=menu_icon, 
                        text=None, 
                        width=40, 
                        height=40, 
                        fg_color="transparent", 
                        hover_color="#303030", 
                        command=toggle_menu,
                        anchor="center")
menu_btn.grid(row=0, column=0, pady=(10, 20), padx=0, sticky="nsew")  

# Функция для создания кнопок навигации
def create_nav_button(row, icon, text, command):
    btn = ctk.CTkButton(left_frame, 
                       image=icon, 
                       text=None,
                       command=command, 
                       width=40,
                       height=40,
                       anchor="w",
                       fg_color="transparent", 
                       hover_color="#303030",
                       text_color="#6a6a6a",
                       corner_radius=5)
    
    btn.grid(row=row, column=0, sticky="nsew", pady=(0, 10), padx=0)
    btn.full_text = text
    btn.active = False  # Добавляем атрибут для отслеживания активного состояния
    return btn

# Функции
def replace_comma_with_dot(value):
    return value.replace(',', '.')


def toggle_color_theme():
    current_theme = ctk.get_appearance_mode()
    if current_theme == "Light":
        ctk.set_appearance_mode("Dark")
        # Обновляем только существующие и отображенные фреймы
        # Обновляем цвет всех фреймов с графиками
        for frame in [graph_frame, graph_frame2, graph_frame3, graph_frame4]:
            frame.configure(fg_color="#252525")


        root.configure(bg="#1a1a1a", fg_color="#1a1a1a")
        btn2.configure(fg_color="#505052", hover_color="#9b2d30")
        btn4.configure(fg_color="#505052", hover_color="#9b2d30")
        btn6.configure(fg_color="#505052", hover_color="#9b2d30")
        btn8.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn10.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn13.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn15.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn17.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn19.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn21.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn23.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn25.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn27.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn29.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn101.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn31.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn33.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn35.configure(fg_color="#505052", hover_color="#9b2d30") 
        btn37.configure(fg_color="#505052", hover_color="#9b2d30") 

        btn1_1.configure(fg_color="#505052", hover_color="#73777d")

        save_button.configure(fg_color="#009dda", hover_color="#2dcbff")
        clear1_1.configure(fg_color="#20252e", hover_color="#9b2d30")
        clear11_1.configure(fg_color="#20252e", hover_color="#9b2d30")
        # savegraf2_button.configure(fg_color="#009dda", hover_color="#2dcbff")
        clear11_1.configure(fg_color="#20252e", hover_color="#9b2d30")

        clear12_12.configure(fg_color="#20252e", hover_color="#9b2d30")

        left_frame.configure(fg_color="#151515")
        main_frame.configure(fg_color="#1f1f1f")
        frame_calc.configure(fg_color="#1f1f1f")
        frame_grafik.configure(fg_color="#1f1f1f")
        frame_settings.configure(fg_color="#1f1f1f")
        frame_save.configure(fg_color="#1f1f1f")
        calc_content.configure(fg_color="#1f1f1f")
        grafik_content.configure(fg_color="#1f1f1f")
        settings_content.configure(fg_color="#1f1f1f")
        save_content.configure(fg_color="#1f1f1f")
        calc_right_frame.configure(fg_color="#1f1f1f")
        grafik_right_frame.configure(fg_color="#1f1f1f")
        settings_right_frame.configure( fg_color="#151515")
        save_right_frame.configure( fg_color="#151515")
        left_method_frame.configure(fg_color="#1f1f1f")
        left_method_frame1.configure(fg_color="#1f1f1f")

        right_content_frame.configure(fg_color="#252525", border_color="#454545") #scrollbar_button_color="#606060", #  scrollbar_button_hover_color="#808080"
        right_content_frame1.configure(fg_color="#1f1f1f", border_color="#1f1f1f")
        
        display_frame.configure(fg_color="#151515", border_color="#454545")
        settings_frame.configure(fg_color="#151515", border_color="#454545")
        toggle_graph_frame_btn.configure(fg_color="#151515", hover_color="#252525",border_color="#454545")
        # graph_frame.configure(fg_color="#252525")
        toggle_graph_frame_btn2.configure(fg_color="#151515",hover_color="#252525", border_color="#454545")
        # graph_frame2.configure(fg_color="#252525")
        toggle_graph_frame_btn3.configure(fg_color="#151515", hover_color="#252525", border_color="#454545")
        # graph_frame3.configure(fg_color="#252525")
        toggle_graph_frame_btn4.configure(fg_color="#151515",hover_color="#252525", border_color="#454545")
        # graph_frame4.configure(fg_color="#252525")
       
        settings_label3.configure(text_color="#6a6a6a")
        settings_label5.configure(text_color="#6a6a6a")
        save_label3.configure(text_color="#6a6a6a")
        # settings_label7.configure(text_color="#6a6a6a")
        # settings_label9.configure(text_color="#6a6a6a")

        settings_btn.configure(text_color="#6a6a6a") #6a6a6a
        grafik_btn.configure(text_color="#6a6a6a")
        menu_btn.configure(text_color="#6a6a6a")
        separ.configure(fg_color="#252525", hover_color="#505052")
        widget.configure(fg_color="#151515", hover_color="#252525")
        word_1.configure(fg_color="#151515", hover_color="#252525")
        pdf_1.configure(fg_color="#151515", hover_color="#252525")

        tema_button.configure(fg_color="#151515", hover_color="#252525")
        tema_button1.configure(fg_color="#151515", hover_color="#252525")

        

        # theme_switch.configure(text="Светлая тема")
    else:
        ctk.set_appearance_mode("Light")
        for frame in [graph_frame, graph_frame2, graph_frame3, graph_frame4]:
            frame.configure(fg_color="#868a91")
        
        

        root.configure(bg="#868a91", fg_color="#868a91")
        btn2.configure(fg_color="#868a91", hover_color="#d05e61")
        btn4.configure(fg_color="#868a91", hover_color="#d05e61")
        btn6.configure(fg_color="#868a91", hover_color="#d05e61")
        btn8.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn10.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn13.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn15.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn17.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn19.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn21.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn23.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn25.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn27.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn29.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn101.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn31.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn33.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn35.configure(fg_color="#868a91", hover_color="#d05e61") 
        btn37.configure(fg_color="#868a91", hover_color="#d05e61") 

        separ.configure(fg_color="#414142", hover_color="#68696b")

        btn1_1.configure(fg_color="#505052", hover_color="#73777d")

        save_button.configure(fg_color="#44d7ff", hover_color="#00b3f3")
        clear1_1.configure(fg_color="#68696b", hover_color="#d05e61")
        clear11_1.configure(fg_color="#68696b", hover_color="#d05e61")
        clear12_12.configure(fg_color="#68696b", hover_color="#d05e61")
        # savegraf2_button.configure(fg_color="#44d7ff", hover_color="#00b3f3")
        clear11_11.configure(fg_color="#68696b", hover_color="#d05e61")

        left_frame.configure(fg_color="#414142")
        main_frame.configure(fg_color="#868a91")
        frame_calc.configure(fg_color="#868a91")
        frame_grafik.configure(fg_color="#868a91")
        frame_settings.configure(fg_color="#868a91")
        frame_save.configure(fg_color="#868a91")
        calc_content.configure(fg_color="#868a91")
        grafik_content.configure(fg_color="#868a91")
        settings_content.configure(fg_color="#868a91")
        save_content.configure(fg_color="#868a91")
        calc_right_frame.configure(fg_color="#868a91")
        grafik_right_frame.configure(fg_color="#868a91")
        settings_right_frame.configure( fg_color="#414142")
        save_right_frame.configure( fg_color="#414142")
        left_method_frame.configure(fg_color="#868a91")
        left_method_frame1.configure(fg_color="#868a91")

        settings_btn.configure(text_color="#a0a4ac") #6a6a6a"
        grafik_btn.configure(text_color="#a0a4ac")
        menu_btn.configure(text_color="#a0a4ac")

        right_content_frame.configure(fg_color="#414142", border_color="#b0b0b0") #scrollbar_button_color="#606060", #  scrollbar_button_hover_color="#808080"
        right_content_frame1.configure(fg_color="#868a91", border_color="#868a91")

        display_frame.configure(fg_color="white", border_color="black")
        settings_frame.configure(fg_color="#414142", border_color="#b0b0b0")
        toggle_graph_frame_btn.configure(fg_color="#414142", hover_color="#73777d",border_color="#b0b0b0")
        
        toggle_graph_frame_btn2.configure(fg_color="#414142",hover_color="#73777d", border_color="#b0b0b0")
        toggle_graph_frame_btn3.configure(fg_color="#414142", hover_color="#73777d", border_color="#b0b0b0")
        toggle_graph_frame_btn4.configure(fg_color="#414142",hover_color="#73777d", border_color="#b0b0b0")
        # graph_frame.configure(fg_color="#868a91")
        # graph_frame2.configure(fg_color="#868a91")
        # graph_frame3.configure(fg_color="#868a91")
        # graph_frame4.configure(fg_color="#868a91")
        
        settings_label3.configure(text_color="#868a91")
        settings_label5.configure(text_color="#868a91")
        save_label3.configure(text_color="#868a91")
        # settings_label7.configure(text_color="#868a91")
        # settings_label9.configure(text_color="#868a91")

        widget.configure(fg_color="#414142", hover_color="#73777d")

        word_1.configure(fg_color="#414142", hover_color="#68696b")
        pdf_1.configure(fg_color="#414142", hover_color="#68696b")

        tema_button.configure(fg_color="#414142", hover_color="#73777d")
        tema_button1.configure(fg_color="#414142", hover_color="#73777d")

        # Всегда обновляем тему фрейма, даже если он скрыт
    update_separ_frame_theme()
    
    # Обновляем иконку кнопки
    if separ_frame.winfo_ismapped():
        separ.configure(image=galka_active_icon)
    else:
        separ.configure(image=galka_icon)

    # Добавляем обновление виджета и его фреймов
    update_widget_frames_theme()

   # Перерисовываем графики, если они существуют
    redraw_existing_plots()

    # Если активен виджет, обновляем подфреймы и графики
    if widget.active:
        toggle_widget()  # Деактивируем
        toggle_widget()  # Активируем снова с новыми настройками темы

def redraw_existing_plots():
    """Перерисовывает все существующие графики с учетом текущей темы"""
    # Если активен виджет с 4 фреймами
    if widget.active and subframes:
        # Перерисовываем графики во всех подфреймах
        for frame in subframes:
            create_empty_plot_in_frame(frame)
        return
    
    # Получаем текущий активный график (если виджет не активен)
    active_graph = None
    if graph_frame.winfo_ismapped():
        active_graph = "phi"
    elif graph_frame2.winfo_ismapped():
        active_graph = "eta"
    elif graph_frame3.winfo_ismapped():
        active_graph = "he"
    elif graph_frame4.winfo_ismapped():
        active_graph = "eta_he"
    
    # Если есть активный график, перерисовываем его
    if active_graph:
        try:
            # Для каждого типа графика проверяем заполненность полей и перерисовываем
            if active_graph == "phi":
                if (entry_min_phi.get() and entry_max_phi.get() and 
                    entry_step_phi.get()):
                    plot_phi_vs_pk()
                else:
                    create_empty_plot()
            
            elif active_graph == "eta":
                if (entry_min_eta.get() and entry_max_eta.get() and 
                    entry_step_eta.get()):
                    plot_eta_vs_pk()
                else:
                    create_empty_plot()
            
            elif active_graph == "he":
                if (entry_min_he.get() and entry_max_he.get() and 
                    entry_step_he.get()):
                    plot_he_vs_pk()
                else:
                    create_empty_plot()
            
            elif active_graph == "eta_he":
                if (entry_min_eta_he.get() and entry_max_eta_he.get() and 
                    entry_step_eta_he.get()):
                    plot_eta_he_vs_pk()
                else:
                    create_empty_plot()
        
        except Exception as e:
            print(f"Ошибка при перерисовке графика: {e}")
            create_empty_plot()
    else:
        # Если нет активного графика, создаем пустой
        create_empty_plot()
    
    # Обновляем цвет фреймов с графиками
    current_theme = ctk.get_appearance_mode()
    graph_color = "#252525" if current_theme == "Dark" else "#868a91"
    
    for frame in [graph_frame, graph_frame2, graph_frame3, graph_frame4]:
        frame.configure(fg_color=graph_color)




# Добавляем функцию для обновления подфреймов виджета
def update_widget_frames_theme():
    global subframes
    if not widget.active:
        return
    
    current_theme = ctk.get_appearance_mode()
    new_fg = "#151515" if current_theme == "Dark" else "white"
    new_border = "#454545" if current_theme == "Dark" else "black"
    
    for frame in subframes:
        frame.configure(fg_color=new_fg, border_color=new_border)
        # Обновляем графики внутри фреймов
        for child in frame.winfo_children():
            if isinstance(child, FigureCanvasTkAgg):
                create_empty_plot_in_frame(frame)


def get1():
    try:
        value = replace_comma_with_dot(inp1.get())
        res1 = 'Tн=' + str(float(value)) + ' K'
        label1.configure(text=res1, text_color="white")
    except ValueError:
        label1.configure(text='Введите число!', text_color="#b51b1b")


def delete1():
    inp1.delete(0, ctk.END)
    label1.configure(text='Температура окружающей среды Tн, K', text_color="#bfbfbf")

def get2():
    try:
        value = replace_comma_with_dot(inp2.get())
        res2 = 'Pн=' + str(float(value)) + ' Па'
        label2.configure(text=res2, text_color="white")
    except ValueError:
        label2.configure(text='Введите число!', text_color="#b51b1b")


def delete2():
    inp2.delete(0, ctk.END)
    label2.configure(text='Давление окружающей среды Pн, Па', text_color="#bfbfbf")

def get3():
    try:
        value = replace_comma_with_dot(inp3.get())
        res3 = 'T3=' + str(float(value)) + ' K'
        label3.configure(text=res3, text_color="white")
    except ValueError:
        label3.configure(text='Введите число!', text_color="#b51b1b")

def delete3():
    inp3.delete(0, ctk.END)
    label3.configure(text='Начальная температура газа перед турбиной T3, K', text_color="#bfbfbf")

def get4():
    try:
        value = replace_comma_with_dot(inp4.get())
        res4 = 'R=' + str(float(value)) + ' Дж/(кг*К)'
        label4.configure(text=res4, text_color="white")
    except ValueError:
        label4.configure(text='Введите число!', text_color="#b51b1b")

def delete4():
    inp4.delete(0, ctk.END)   # удаление введенного текста
    label4.configure(text='Газовая постоянная для воздуха R, Дж/(кг*К)', text_color="#bfbfbf")

def get5():
    try:
        pcs = replace_comma_with_dot(inp5.get())
        res5 = 'σвх=' + str(float(pcs)) 
        label5.configure(text=res5, text_color="white")
        pcs1 = float(pcs)

        if pcs1 > 1:
            label5.configure(text='Невозможное значение σвх', text_color="#b51b1b")
        elif pcs1 < 0:
            label5.configure(text='Невозможное значение σвх', text_color="#b51b1b")
        else:
            print ('Возможное значение')
    except ValueError:
        label5.configure(text='Введите число!', text_color="#b51b1b")

def delete5():
    inp5.delete(0, ctk.END)   # удаление введенного текста
    label5.configure(text='Коэффициент потерь на входе в компрессор σвх', text_color="#bfbfbf")
  
    
#==================================================================================================================================================
def get7():
    try:
        pcs = replace_comma_with_dot(inp7.get())
        res7 = 'σвых*=' + str(float(pcs)) 
        label9.configure(text=res7, text_color="white")
        pcss1 = float(pcs)

        if pcss1 > 1:
            label9.configure(text='Невозможное значение σвых*', text_color="#b51b1b")
        elif pcss1 < 0:
            label9.configure(text='Невозможное значение σвых*', text_color="#b51b1b")
        else:
            print ('Возможное значение', text_color="#b51b1b")
    except ValueError:
        label9.configure(text='Введите число!', text_color="#b51b1b")

def delete7():
    inp7.delete(0, ctk.END)   # удаление введенного текста
    label9.configure(text='Коэффициент потерь давления воздуха в выходном устройстве σвых*', text_color="#bfbfbf")
#==================================================================================================================================================
def get8():
    try:
        pcs = replace_comma_with_dot(inp8.get()) 
        res8 = 'σvtepl*=' + str(float(pcs)) 
        label10.configure(text=res8, text_color="white")
        pcsss1 = float(pcs)

        if pcsss1 > 1:
            label10.configure(text='Невозможное значение σvtepl*', text_color="#b51b1b")
        elif pcsss1 < 0:
            label10.configure(text='Невозможное значение σvtepl*', text_color="#b51b1b")
        else:
            print ('Возможное значение')
    except ValueError:
        label10.configure(text='Введите число!', text_color="#b51b1b")

def delete8():
    inp8.delete(0, ctk.END)   # удаление введенного текста
    label10.configure(text='Коэффициент потерь давления воздуха перед кс σvtepl*', text_color="#bfbfbf")
#==================================================================================================================================================
def get9():
    try:
        pcs = replace_comma_with_dot(inp9.get()) 
        res9 = 'σкс=' + str(float(pcs)) 
        label11.configure(text=res9, text_color="white")
        pcssss1 = float(pcs)

        if pcssss1 > 1: 
            label11.configure(text='Невозможное значение σкс', text_color="#b51b1b")
        elif pcssss1 < 0:
            label11.configure(text='Невозможное значение σкс', text_color="#b51b1b")
        else:
            print ('Возможное значение')
    except ValueError:
        label11.configure(text='Введите число!', text_color="#b51b1b")

def delete9():
    inp9.delete(0, ctk.END)   # удаление введенного текста
    label11.configure(text='Коэффициент потерь давления воздуха в кс σкс', text_color="#bfbfbf")
#==================================================================================================================================================
def get10():
    try:
        pcs = replace_comma_with_dot(inp10.get()) 
        res10 = 'ηпол=' + str(float(pcs)) + '%'
        label12.configure(text=res10, text_color="white")
        pcsssss1 = float(pcs)

        if pcsssss1 > 100:
            label12.configure(text='Невозможное значение ηпол %', text_color="#b51b1b")
        elif pcsssss1 < 0:
            label12.configure(text='Невозможное значение ηпол %', text_color="#b51b1b")
        else:
            print ('Возможное значение')
    except ValueError:
        label12.configure(text='Введите число!', text_color="#b51b1b")

def delete10():
    inp10.delete(0, ctk.END)   # удаление введенного текста
    label12.configure(text='Политропный КПД турбины ηпол %', text_color="#bfbfbf")
#==================================================================================================================================================
def get11():
    try:
        pcs = replace_comma_with_dot(inp11.get()) 
        res11 = 'ηмт=' + str(float(pcs)) + '%'
        label13.configure(text=res11, text_color="white")
        p11 = float(pcs)

        if p11 > 100:
            label13.configure(text='Невозможное значение ηмт %', text_color="#b51b1b")
        elif p11 < 0:
            label13.configure(text='Невозможное значение ηмт %', text_color="#b51b1b")
        else:
            print ('Возможное значение')
    except ValueError:
        label13.configure(text='Введите число!', text_color="#b51b1b")

def delete11():
    inp11.delete(0, ctk.END)   # удаление введенного текста
    label13.configure(text='Механический КПД турбины ηмт %', text_color="#bfbfbf")
#==================================================================================================================================================
def get12():
    try:
        pcs = replace_comma_with_dot(inp12.get())  
        res12 = 'ηмк=' + str(float(pcs)) + '%'
        label14.configure(text=res12, text_color="white")
        p111 = float(pcs)

        if p111 > 100:
            label14.configure(text='Невозможное значение ηмк %', text_color="#b51b1b")
        elif p111 < 0:
            label14.configure(text='Невозможное значение ηмк %', text_color="#b51b1b")
        else:
            print ('Возможное значение', text_color="#b51b1b")
    except ValueError:
        label14.configure(text='Введите число!', text_color="#b51b1b")

def delete12():
    inp12.delete(0, ctk.END)   # удаление введенного текста
    label14.configure(text='Механический КПД компрессора ηмк %', text_color="#bfbfbf")
#==================================================================================================================================================
def get13():
    try:
        pcs = replace_comma_with_dot(inp13.get()) 
        res13 = 'ηкад=' + str(float(pcs)) + '%'
        label15.configure(text=res13, text_color="white")
        pi111 = float(pcs)

        if pi111 > 100:
            label15.configure(text='Невозможное значение ηкад %', text_color="#b51b1b")
        elif pi111 < 0:
            label15.configure(text='Невозможное значение ηкад %', text_color="#b51b1b")
        else:
            print ('Возможное значение')
    except ValueError:
        label15.configure(text='Введите число!', text_color="#b51b1b")

def delete13():
    inp13.delete(0, ctk.END)   # удаление введенного текста
    label15.configure(text='Адиабатический КПД компрессора ηкад %', text_color="#bfbfbf")

#==================================================================================================================================================
def get14():
    try:
        pcs = replace_comma_with_dot(inp14.get()) 
        res14 = 'μ=' + str(float(pcs)) + '%'
        label16.configure(text=res14, text_color="white")
        piс111 = float(pcs)

        if piс111 > 100:
            label16.configure(text='Невозможное значение μ %', text_color="#b51b1b")
        elif piс111 < 0:
            label16.configure(text='Невозможное значение μ %', text_color="#b51b1b")
        else:
            print ('Возможное значение')
    except ValueError:
        label16.configure(text='Введите число!', text_color="#b51b1b")

def delete14():
    inp14.delete(0, ctk.END)   # удаление введенного текста
    label16.configure(text='Степень рекуперации μ %', text_color="#bfbfbf")
#==================================================================================================================================================
def get15():
    try:
        value = replace_comma_with_dot(inp15.get())
        res15 = 'Kв=' + str(float(value)) 
        label17.configure(text=res15, text_color="white")
    except ValueError:
        label17.configure(text='Введите число!', text_color="#b51b1b")

def delete15():
    inp15.delete(0, ctk.END)   # удаление введенного текста
    label17.configure(text='Показатель изоэнтропы для воздуха Kв', text_color="#bfbfbf")

#==================================================================================================================================================
def get16():
    try:
        value = replace_comma_with_dot(inp16.get())
        res16 = 'Kг=' + str(float(value)) 
        label18.configure(text=res16, text_color="white")
    except ValueError:
        label18.configure(text='Введите число!', text_color="#b51b1b")

def delete16():
    inp16.delete(0, ctk.END)   # удаление введенного текста
    label18.configure(text='Показатель изоэнтропы для газа Kг', text_color="#bfbfbf")
#==================================================================================================================================================
def get17():
    try:
        value = replace_comma_with_dot(inp17.get()) 
        res16 = 'Cв=' + str(float(value)) + ' Дж/(кг*К)'
        label19.configure(text=res16, text_color="white")
    except ValueError:
        label19.configure(text='Введите число!', text_color="#b51b1b")

def delete17():
    inp17.delete(0, ctk.END)   # удаление введенного текста
    label19.configure(text='Теплоёмкость воздуха Cв, Дж/(кг*К)', text_color="#bfbfbf")

#==================================================================================================================================================
def get18():
    try:
        value = replace_comma_with_dot(inp18.get()) 
        res16 = 'Cг=' + str(float(value)) + ' Дж/(кг*К)'
        label20.configure(text=res16, text_color="white")
    except ValueError:
        label20.configure(text='Введите число!', text_color="#b51b1b")

def delete18():
    inp18.delete(0, ctk.END)   # удаление введенного текста
    label20.configure(text='Теплоёмкость газа Cг, Дж/(кг*К)', text_color="#bfbfbf")
#==================================================================================================================================================
def get100():
    try:
        value = replace_comma_with_dot(inp100.get()) 
        res16 = 'Ne=' + str(float(value)) + ' кВт'
        label100.configure(text=res16, text_color="white")
    except ValueError:
        label100.configure(text='Введите число!', text_color="#b51b1b")

def delete100():
    inp100.delete(0, ctk.END)   # удаление введенного текста
    label100.configure(text='Эффективная мощность Ne, кВт', text_color="#bfbfbf")
#==================================================================================================================================================
def get6():
    try:
        value = replace_comma_with_dot(inp6.get()) 
        pk = 'πк=' + str(float(value)) 
        label8.configure(text=pk, text_color="white")
    except ValueError:
        label8.configure(text='Введите число!', text_color="#b51b1b")

def delete6():
    inp6.delete(0, ctk.END)   # удаление введенного текста
    label8.configure(text='πк - степень повышения давления', text_color="#bfbfbf") # возвращение в строку текста вводимой величины

def clear():
    try:
        # Вызов всех функций delete
        delete1()
        delete2()
        delete3()
        delete4()
        delete5()
        delete6()
        delete7()
        delete8()
        delete9()
        delete10()
        delete11()
        delete12()
        delete13()
        delete14()
        delete15()
        delete16()
        delete17()
        delete18()
        delete100()

        if lang.get() == 'без регенерации':
            label6.configure(text='Давление воздуха перед компрессором P1*, Па', text_color="#bfbfbf")
            label7.configure(text='Температура воздуха перед компрессором T1*, К', text_color="#bfbfbf")
            label21.configure(text='Давление воздуха за компрессором P2*, Па', text_color="#bfbfbf")
            label22.configure(text='Температура воздуха за компрессором T2*, К', text_color="#bfbfbf")
            label31.configure(text='Температура воздуха перед кс T2*, К', text_color="#bfbfbf")
            label23.configure(text='Работа изоэнтропийного перепада в компрессоре Hok*, кДж/кг', text_color="#bfbfbf")
            label24.configure(text='Полезная работа в копрессоре Hk, кДж/кг', text_color="#bfbfbf")
            label36.configure(text='', text_color="#bfbfbf")
            update_label_positions()  
            label50.configure(text='Давление газа перед турбиной P3*, Па', text_color="#bfbfbf")
            label25.configure(text='Давление газа за турбиной P4*, Па', text_color="#bfbfbf")
            label26.configure(text='Степень расширения газа в турбине πт*', text_color="#bfbfbf")
            label27.configure(text='Работа изоэнтропийного перепада в турбине Hот*, кДж/кг', text_color="#bfbfbf")
            label28.configure(text='Полезная работа в турбине Hт, кДж/кг', text_color="#bfbfbf")
            label29.configure(text='Температура газа за турбиной T4*, К', text_color="#bfbfbf")
            label30.configure(text='Расход воздуха через компрессор Gв, кг/с', text_color="#bfbfbf")
            label32.configure(text="Расход теплоты с учетом потерь тепла в кс Q1', кДж/кг", text_color="#bfbfbf")
            label33.configure(text='Расход теплоты Q1, кДж/кг', text_color="#bfbfbf")
            label34.configure(text='Эффективный КПД установки ηе', text_color="#bfbfbf")
            label35.configure(text='Коэффициент полезной работы φ', text_color="#bfbfbf")
            label200.configure(text='Эффективная удельная работа He, кДж/кг', text_color="#bfbfbf")

        elif lang.get() == 'c регенерацией':
            label6.configure(text='Давление воздуха перед компрессором P1*, Па', text_color="#bfbfbf")
            label7.configure(text='Температура воздуха перед компрессором T1*, К', text_color="#bfbfbf")
            label21.configure(text='Давление воздуха за компрессором P2*, Па', text_color="#bfbfbf")
            label22.configure(text='Температура воздуха за компрессором T2*, К', text_color="#bfbfbf")
            label23.configure(text='Работа изоэнтропийного перепада в компрессоре Hok*, кДж/кг', text_color="#bfbfbf")
            label24.configure(text='Полезная работа в копрессоре Hk, кДж/кг', text_color="#bfbfbf")
            label36.configure(text='Давление воздуха перед РВ P5*, Па', text_color="#bfbfbf")
            update_label_positions()  
            label50.configure(text='Давление газа перед турбиной P3*, Па', text_color="#bfbfbf")
            label25.configure(text='Давление газа за турбиной P4*, Па', text_color="#bfbfbf")
            label26.configure(text='Степень расширения газа в турбине πт*', text_color="#bfbfbf")
            label27.configure(text='Работа изоэнтропийного перепада в турбине Hот*, кДж/кг', text_color="#bfbfbf")
            label28.configure(text='Полезная работа в турбине Hт, кДж/кг', text_color="#bfbfbf")
            label29.configure(text='Температура газа за турбиной T4*, К', text_color="#bfbfbf")
            label30.configure(text='Расход воздуха через компрессор Gв, кг/с', text_color="#bfbfbf")
            label31.configure(text='Температура воздуха перед кс Ts*, К', text_color="#bfbfbf")
            label32.configure(text="Расход теплоты с учетом потерь тепла в кс Q1', кДж/кг", text_color="#bfbfbf")
            label33.configure(text='Расход теплоты Q1, кДж/кг', text_color="#bfbfbf")
            label34.configure(text='Эффективный КПД установки ηе', text_color="#bfbfbf")
            label35.configure(text='Коэффициент полезной работы φ', text_color="#bfbfbf")
            label200.configure(text='Эффективная удельная работа He, кДж/кг', text_color="#bfbfbf")

        else:
            print("Не выбран режим очистки")
    except Exception as e:
        print("Ошибка", e)

def raschet():
    try:
        if lang.get() == 'без регенерации':
            # Расчёты для режима без регенерации
            P1_ = float(replace_comma_with_dot(inp2.get()))
            Pcs_ = float(replace_comma_with_dot(inp5.get()))
            a = P1_ * Pcs_ 
            a_rounded = (round(a, 2)) #округление числа до 2 знаков
            result6 = 'P1*=' + str(a_rounded) + ' Па'
            label6.configure(text=result6, text_color="white")

            Tn_ = float(replace_comma_with_dot(inp1.get()))
            Tn_rounded = (round(Tn_, 1)) #округление числа до 1 знака
            result7 = 'T1*=' + str(Tn_rounded) + ' K'
            label7.configure(text=result7, text_color="white")

            Pk_ = float(replace_comma_with_dot(inp6.get()))
            P2_ = Pk_ * a
            P2_rounded = (round(P2_, 2)) #округление числа до 2 знаков
            result8 = 'P2*=' + str(P2_rounded) + ' Па'
            label21.configure(text=result8, text_color="white")

            Kv_ = float(replace_comma_with_dot(inp15.get()))
            T2_ = Tn_ * Pk_ ** ((Kv_- 1) / Kv_)
            T2_rounded = (round(T2_, 1)) #округление числа до 1 знака
            result9 = 'T2*=' + str(T2_rounded) + ' К'
            label22.configure(text=result9, text_color="white")
            label31.configure(text=result9, text_color="white")

            Cv_ = float(replace_comma_with_dot(inp17.get()))
            Hok_ = Cv_ * Tn_ * ((Pk_ ** ((Kv_ - 1) / Kv_)) - 1)
            Hok_r = Hok_ / 1000
            Hok_rounded = (round(Hok_r, 3)) 
            result10 = 'Hok*=' + str(Hok_rounded) + ' кДж/кг'
            label23.configure(text=result10, text_color="white")

            KPDad_ = float(replace_comma_with_dot(inp13.get()))
            KPDad = KPDad_ / 100
            Hk_ = Hok_ / KPDad 
            Hk_r = Hk_ / 1000
            Hk_rounded = (round(Hk_r, 3)) 
            re11 = 'Hk=' + str(Hk_rounded) + ' кДж/кг'
            label24.configure(text=re11, text_color="white")

            label36.configure(text='')
            update_label_positions()  

            Sigmacs_ = float(replace_comma_with_dot(inp9.get()))
            Sigmacs = Sigmacs_ 
            P3_ = Sigmacs * P2_ 
            P3_rounded = (round(P3_, 2))
            re12 = 'P3*=' + str(P3_rounded) + ' Па'
            label50.configure(text=re12, text_color="white")

            Sigmavix_ = float(replace_comma_with_dot(inp7.get()))
            Sigmavix = Sigmavix_ 
            P4_ = P1_ / Sigmavix
            P4_rounded = (round(P4_, 2))
            re13 = 'P4*=' + str(P4_rounded) + ' Па'
            label25.configure(text=re13, text_color="white")

            PiT_ = P3_ / P4_
            PiT_rounded = (round(PiT_, 1))
            re14 = 'πт*=' + str(PiT_rounded)
            label26.configure(text=re14, text_color="white")

            Cg_ = float(replace_comma_with_dot(inp18.get()))
            T3_ = float(replace_comma_with_dot(inp3.get()))
            Kg_ = float(replace_comma_with_dot(inp16.get()))
            Kg = (Kg_ - 1) / Kg_ 
            Kgg = -1 * Kg
            Hot_ = Cg_ * T3_ * (1 - PiT_ ** Kgg)
            Hot_r = Hot_ / 1000
            Hot_rounded = (round(Hot_r, 3)) 
            re15 = 'Hот*=' + str(Hot_rounded) + ' кДж/кг'
            label27.configure(text=re15, text_color="white")

            KPDpol_ = float(replace_comma_with_dot(inp10.get()))
            Htt = Hot_ * KPDpol_ / 100
            Htt_r = Htt / 1000
            Htt_rounded = (round(Htt_r, 3)) 
            re16 = 'Hт=' + str(Htt_rounded) + ' кДж/кг'
            label28.configure(text=re16, text_color="white")

            T4_ = T3_ * PiT_ ** Kgg
            T4_rounded = (round(T4_, 1))
            re17 = 'T4*=' + str(T4_rounded) + ' К'
            label29.configure(text=re17, text_color="white")

            Ne_ = float(replace_comma_with_dot(inp100.get()))
            KPDmt_ = float(replace_comma_with_dot(inp11.get()))
            KPDmk_ = float(replace_comma_with_dot(inp12.get()))
            KPDmt = KPDmt_ / 100
            KPDmk = KPDmk_ / 100
            Gv_ = (1000 * Ne_) / ((Htt * KPDmt) - (Hk_ / KPDmk))
            Gv_rounded = (round(Gv_, 2))
            re18 = 'Gв=' + str(Gv_rounded) + ' кг/с'
            label30.configure(text=re18, text_color="white")

            Q_1_ = Cg_ * (T3_ - T2_)
            Q_1_r = Q_1_ / 1000
            Q_1_rounded = (round(Q_1_r, 3)) 
            re19 = "Q1'=" + str(Q_1_rounded) + " кДж/кг"
            label32.configure(text=re19, text_color="white")

            Q1_ = Q_1_ / Sigmacs
            Q1_r = Q1_ / 1000
            Q1_rounded = (round(Q1_r, 3)) 
            re20 = "Q1=" + str(Q1_rounded) + " кДж/кг"
            label33.configure(text=re20, text_color="white")

            HNe_ = ((Htt * KPDmt) - (Hk_ / KPDmk)) / Q1_
            HNe_rounded = (round(HNe_, 3))
            re21 = 'ηе=' + str(HNe_rounded)
            # Проверяем условие и обновляем label34
            if HNe_ > 1:
                label34.configure(text='ηе=1', text_color="white")  # Если HNe_ больше 1, выводим 1
            else:
                label34.configure(text=re21, text_color="white")  # Иначе выводим значение 

            Fi_ = ((Htt * KPDmt) - (Hk_ / KPDmk)) / (Htt * KPDmt)
            Fi_rounded = (round(Fi_, 3))
            re22 = 'φ=' + str(Fi_rounded)
            # Проверяем условие и обновляем label35
            if Fi_ > 1:
                label35.configure(text='φ=1', text_color="white")  # Если Fi_ больше 1, выводим 1
            else:
                label35.configure(text=re22, text_color="white")  # Иначе выводим значение Fi_

            HHHe = ((Htt * KPDmt) - (Hk_ / KPDmk))
            HHHe_r = HHHe / 1000
            HHHe_rounded = (round(HHHe_r, 3)) 
            re200 = 'He=' + str(HHHe_rounded) + ' кДж/кг'
            label200.configure(text=re200, text_color="white")
#==================================================================================================================================================
        elif lang.get() == 'c регенерацией':
            # Расчёты для режима с регенерацией
            P1_ = float(replace_comma_with_dot(inp2.get()))
            Pcs_ = float(replace_comma_with_dot(inp5.get()))
            a = P1_ * Pcs_ 
            a_rounded = (round(a, 2)) #округление числа до 2 знаков
            result6 = 'P1*=' + str(a_rounded) + ' Па'
            label6.configure(text=result6, text_color="white")

            Tn_ = float(replace_comma_with_dot(inp1.get()))
            Tn_rounded = (round(Tn_, 1)) #округление числа до 1 знака
            result7 = 'T1*=' + str(Tn_rounded) + ' K'
            label7.configure(text=result7, text_color="white")

            Pk_ = float(replace_comma_with_dot(inp6.get()))
            P2_ = Pk_ * a
            P2_rounded = (round(P2_, 2)) #округление числа до 2 знаков
            result8 = 'P2*=' + str(P2_rounded) + ' Па'
            label21.configure(text=result8, text_color="white")

            Kv_ = float(replace_comma_with_dot(inp15.get()))
            T2_ = Tn_ * Pk_ ** ((Kv_- 1) / Kv_)
            T2_rounded = (round(T2_, 1)) #округление числа до 1 знака
            result9 = 'T2*=' + str(T2_rounded) + ' К'
            label22.configure(text=result9, text_color="white")

            Cv_ = float(replace_comma_with_dot(inp17.get()))
            Hok_ = Cv_ * Tn_ * ((Pk_ ** ((Kv_ - 1) / Kv_)) - 1)
            Hok_rr = Hok_ / 1000
            Hok_rounded = (round(Hok_rr, 3)) 
            result10 = 'Hok*=' + str(Hok_rounded) + ' кДж/кг'
            label23.configure(text=result10, text_color="white")

            KPDad_ = float(replace_comma_with_dot(inp13.get()))
            KPDad = KPDad_ / 100
            Hk_ = Hok_ / KPDad 
            Hk_rr = Hk_ / 1000
            Hk_rounded = (round(Hk_rr, 3)) 
            re11 = 'Hk=' + str(Hk_rounded) + ' кДж/кг'
            label24.configure(text=re11, text_color="white")

            Sigmavtepl_ = float(replace_comma_with_dot(inp8.get()))
            Sigmavtepl = Sigmavtepl_
            P5_ = Sigmavtepl * P2_
            P5_rounded = (round(P5_, 2))
            resu1 = 'P5*=' + str(P5_rounded) + ' Па'
            label36.configure(text=resu1, text_color="white")
            update_label_positions()  # Добавьте эту строку

            Sigmacs_ = float(replace_comma_with_dot(inp9.get()))
            Sigmacs = Sigmacs_ 
            P3_ = Sigmacs * P5_ 
            P3_rounded = (round(P3_, 2))
            re12 = 'P3*=' + str(P3_rounded) + ' Па'
            label50.configure(text=re12, text_color="white")

            Sigmavix_ = float(replace_comma_with_dot(inp7.get()))
            Sigmavix = Sigmavix_ 
            P4_ = P1_ / (Sigmavix * Sigmavtepl)
            P4_rounded = (round(P4_, 2))
            re13 = 'P4*=' + str(P4_rounded) + ' Па'
            label25.configure(text=re13, text_color="white")

            PiT_ = P3_ / P4_
            PiT_rounded = (round(PiT_, 1))
            re14 = 'πт*=' + str(PiT_rounded)
            label26.configure(text=re14, text_color="white")

            Cg_ = float(replace_comma_with_dot(inp18.get()))
            T3_ = float(replace_comma_with_dot(inp3.get()))
            Kg_ = float(replace_comma_with_dot(inp16.get()))
            Kg = (Kg_ - 1) / Kg_ 
            Kgg = -1 * Kg
            Hot_ = Cg_ * T3_ * (1 - PiT_ ** Kgg)
            Hot_rr = Hot_ / 1000
            Hot_rounded = (round(Hot_rr, 3)) 
            re15 = 'Hот*=' + str(Hot_rounded) + ' кДж/кг'
            label27.configure(text=re15, text_color="white")

            KPDpol_ = float(replace_comma_with_dot(inp10.get()))
            Htt = Hot_ * KPDpol_ / 100
            Htt_rr = Htt / 1000
            Htt_rounded = (round(Htt_rr, 3)) 
            re16 = 'Hт=' + str(Htt_rounded) + ' кДж/кг'
            label28.configure(text=re16, text_color="white")

            T4_ = T3_ * PiT_ ** Kgg
            T4_rounded = (round(T4_, 1))
            re17 = 'T4*=' + str(T4_rounded) + ' К'
            label29.configure(text=re17, text_color="white")

            Ne_ = float(replace_comma_with_dot(inp100.get()))
            KPDmt_ = float(replace_comma_with_dot(inp11.get()))
            KPDmk_ = float(replace_comma_with_dot(inp12.get()))
            KPDmt = KPDmt_ / 100
            KPDmk = KPDmk_ / 100
            Gv_ = (1000 * Ne_) / ((Htt * KPDmt) - (Hk_ / KPDmk))
            Gv_rounded = (round(Gv_, 2))
            re18 = 'Gв=' + str(Gv_rounded) + ' кг/с'
            label30.configure(text=re18, text_color="white")

            Mu_ = float(replace_comma_with_dot(inp14.get()))
            Mu = Mu_ / 100
            T5_ = T2_ + Mu * (T4_ - T2_)
            T5_rounded = (round(T5_, 1))
            re199 = 'T5*=' + str(T5_rounded) + ' К'
            label31.configure(text=re199, text_color="white")

            Q_1_ = Cg_ * (T3_ - T5_)
            Q_1_rr = Q_1_ / 1000
            Q_1_rounded = (round(Q_1_rr, 3)) 
            re19 = "Q1'=" + str(Q_1_rounded) + " кДж/кг"
            label32.configure(text=re19, text_color="white")

            Q1_ = Q_1_ / Sigmacs
            Q1_rr = Q1_ / 1000
            Q1_rounded = (round(Q1_rr, 3)) 
            re20 = "Q1=" + str(Q1_rounded) + " кДж/кг"
            label33.configure(text=re20, text_color="white")

            HNe_ = ((Htt * KPDmt) - (Hk_ / KPDmk)) / Q1_
            HNe_rounded = (round(HNe_, 3))
            re21 = 'ηе=' + str(HNe_rounded)
            
            # Проверяем условие и обновляем label34
            if HNe_ > 1:
                label34.configure(text='ηе=1', text_color="white")  # Если HNe_ больше 1, выводим 1
            else:
                label34.configure(text=re21, text_color="white")  # Иначе выводим значение 

            Fi_ = ((Htt * KPDmt) - (Hk_ / KPDmk)) / (Htt * KPDmt)
            Fi_rounded = (round(Fi_, 3))
            re22 = 'φ=' + str(Fi_rounded)
            # Проверяем условие и обновляем label35
            if Fi_ > 1:
                label35.configure(text='φ=1')  # Если Fi_ больше 1, выводим 1
            else:
                label35.configure(text=re22, text_color="white")  # Иначе выводим значение Fi_

            HHHe = ((Htt * KPDmt) - (Hk_ / KPDmk))
            HHHe_rr = HHHe / 1000
            HHHe_rounded = (round(HHHe_rr, 3))
            re200 = 'He=' + str(HHHe_rounded) + ' кДж/кг'
            label200.configure(text=re200, text_color="white")

            
        else:
            print("Не выбран режим расчёта!")
    except Exception as e:
        print("Ошибка в расчётах:", e)
#==================================================================================================================================================

#==================================================================================================================================================
def vvod(action=None):
    current_theme = ctk.get_appearance_mode()  # Получаем текущую тему
    graph_frame.grid_remove()
    graph_frame2.grid_remove()
    graph_frame3.grid_remove()
    graph_frame4.grid_remove()
    remove_all_plot_buttons()
    

    if action == 'без регенерации':
        inp14.configure(state="readonly")  # Блокируем поле ввода для режима без регенерации
        if current_theme == "Dark":
            label16.configure(text_color="#616161")  # Серый цвет для темной темы
        else:
            label16.configure(text_color="#616161")  # Серый цвет для светлой темы
    
        
        label6.configure(text='Давление воздуха перед компрессором P1*, Па', text_color="#bfbfbf")
        label7.configure(text='Температура воздуха перед компрессором T1*, К', text_color="#bfbfbf")
        label21.configure(text='Давление воздуха за компрессором P2*, Па', text_color="#bfbfbf")
        label22.configure(text='Температура воздуха за компрессором T2*, К', text_color="#bfbfbf")
        label31.configure(text='Температура воздуха перед кс T2*, К', text_color="#bfbfbf")
        label23.configure(text='Работа изоэнтропийного перепада в компрессоре Hok*, кДж/кг', text_color="#bfbfbf")
        label24.configure(text='Полезная работа в копрессоре Hk, кДж/кг', text_color="#bfbfbf")
        label36.configure(text='', text_color="#bfbfbf")
        update_label_positions() 
        label50.configure(text='Давление газа перед турбиной P3*, Па', text_color="#bfbfbf")
        label25.configure(text='Давление газа за турбиной P4*, Па', text_color="#bfbfbf")
        label26.configure(text='Степень расширения газа в турбине πт*', text_color="#bfbfbf")
        label27.configure(text='Работа изоэнтропийного перепада в турбине Hот*, кДж/кг', text_color="#bfbfbf")
        label28.configure(text='Полезная работа в турбине Hт, кДж/кг', text_color="#bfbfbf")
        label29.configure(text='Температура газа за турбиной T4*, К', text_color="#bfbfbf")
        label30.configure(text='Расход воздуха через компрессор Gв, кг/с', text_color="#bfbfbf")
        label32.configure(text="Расход теплоты с учетом потерь тепла в кс Q1', кДж/кг", text_color="#bfbfbf")
        label33.configure(text='Расход теплоты Q1, кДж/кг', text_color="#bfbfbf")
        label34.configure(text='Эффективный КПД установки ηе', text_color="#bfbfbf")
        label35.configure(text='Коэффициент полезной работы φ', text_color="#bfbfbf")
        label200.configure(text='Эффективная удельная работа He, кДж/кг', text_color="#bfbfbf")

    elif action == 'c регенерацией':
        graph_frame.grid_remove()
        graph_frame2.grid_remove()
        graph_frame3.grid_remove()
        graph_frame4.grid_remove()
        remove_all_plot_buttons()
        
        inp14.configure(state="normal")  # Разблокируем поле ввода для режима с регенерацией
        if current_theme == "Dark":
            label16.configure(text_color="#bfbfbf")  # Белый цвет для темной темы
        else:
            label16.configure(text_color="#bfbfbf")  
        
        
        label6.configure(text='Давление воздуха перед компрессором P1*, Па', text_color="#bfbfbf")
        label7.configure(text='Температура воздуха перед компрессором T1*, К', text_color="#bfbfbf")
        label21.configure(text='Давление воздуха за компрессором P2*, Па', text_color="#bfbfbf")
        label22.configure(text='Температура воздуха за компрессором T2*, К', text_color="#bfbfbf")
        label23.configure(text='Работа изоэнтропийного перепада в компрессоре Hok*, кДж/кг', text_color="#bfbfbf")
        label24.configure(text='Полезная работа в копрессоре Hk, кДж/кг', text_color="#bfbfbf")
        label36.configure(text='Давление воздуха перед РВ P5*, Па', text_color="#bfbfbf")
        update_label_positions() 
        label50.configure(text='Давление газа перед турбиной P3*, Па', text_color="#bfbfbf")
        label25.configure(text='Давление газа за турбиной P4*, Па', text_color="#bfbfbf")
        label26.configure(text='Степень расширения газа в турбине πт*', text_color="#bfbfbf")
        label27.configure(text='Работа изоэнтропийного перепада в турбине Hот*, кДж/кг', text_color="#bfbfbf")
        label28.configure(text='Полезная работа в турбине Hт, кДж/кг', text_color="#bfbfbf")
        label29.configure(text='Температура газа за турбиной T4*, К', text_color="#bfbfbf")
        label30.configure(text='Расход воздуха через компрессор Gв, кг/с', text_color="#bfbfbf")
        label31.configure(text='Температура воздуха перед кс Ts*, К', text_color="#bfbfbf")
        label32.configure(text="Расход теплоты с учетом потерь тепла в кс Q1', кДж/кг", text_color="#bfbfbf")
        label33.configure(text='Расход теплоты Q1, кДж/кг', text_color="#bfbfbf")
        label34.configure(text='Эффективный КПД установки ηе', text_color="#bfbfbf")
        label35.configure(text='Коэффициент полезной работы φ', text_color="#bfbfbf")
        label200.configure(text='Эффективная удельная работа He, кДж/кг', text_color="#bfbfbf")
#==================================================================================================================================================

#==================================================================================================================================================
#Блок работы с документами
from docx import Document
from lxml import etree

def replace_dot(value):
    """
    Заменяет запятую на точку в строке, если она присутствует.
    """
    return value.replace(',', '.')

def replace_dot_with_comma(value):
    """
    Заменяет точку на запятую в строке, если она присутствует.
    """
    return value.replace('.', ',')

def divide_by_100(value):
    """
    Делит значение на 100 и возвращает строку с замененной точкой на запятую.
    """
    try:
        # Заменяем запятую на точку для корректного преобразования в число
        value_with_dot = replace_dot(value)
        # Преобразуем значение в число, делим на 100 и обратно в строку
        divided_value = str(float(value_with_dot) / 100)
        # Заменяем точку на запятую для записи в Word
        return replace_dot_with_comma(divided_value)
    except ValueError:
        # Если значение не является числом, возвращаем его без изменений
        return value
def fill_template():
    try:
        # Определяем режим расчета
        if lang.get() == 'без регенерации':
            
                # Открываем шаблон
                doc = Document('шаблон_без_регенерации.docx')

                # Словарь для замены значений после "=" в формулах
                replacements = {
                    "Температура окружающей среды": replace_dot_with_comma(inp1.get()),  # Температура окружающей среды
                    "Давление окружающей среды": replace_dot_with_comma(inp2.get()),      # Давление окружающей среды
                    "Начальная температура газа перед турбиной": replace_dot_with_comma(inp3.get()),  # Начальная температура газа перед турбиной
                    "Газовая постоянная воздуха": replace_dot_with_comma(inp4.get()),      # Газовая постоянная воздуха
                    "Показатель изоэнтропы для воздуха": replace_dot_with_comma(inp15.get()),  # Показатель изоэнтропы для воздуха
                    "Показатель изоэнтропы для газа": replace_dot_with_comma(inp16.get()),     # Показатель изоэнтропы для газа
                    "Теплоемкость воздуха": replace_dot_with_comma(inp17.get()),              # Теплоемкость воздуха
                    "Теплоемкость газа": replace_dot_with_comma(inp18.get()),                # Теплоемкость газа
                    "Коэффициент потерь на входе в компрессор": replace_dot_with_comma(inp5.get()),  # Коэффициент потерь на входе в компрессор (делим на 100)
                    "Коэффициент потерь давления воздуха в вых. устройстве": replace_dot_with_comma(inp7.get()),  # Коэффициент потерь давления воздуха в вых. устройстве (делим на 100)
                    "Коэффициент потерь давления воздуха перед КС": replace_dot_with_comma(inp8.get()),  # Коэффициент потерь давления воздуха перед КС (делим на 100)
                    "Коэффициент потерь давления воздуха в КС": replace_dot_with_comma(inp9.get()),      # Коэффициент потерь давления воздуха в КС (делим на 100)
                    "Политропный КПД турбины": divide_by_100(inp10.get()),                      # Политропный КПД турбины (делим на 100)
                    "Механический КПД турбины": divide_by_100(inp11.get()),                     # Механический КПД турбины (делим на 100)
                    "Механический КПД компрессора": divide_by_100(inp12.get()),                 # Механический КПД компрессора (делим на 100)
                    "Адиабатический КПД компрессора": divide_by_100(inp13.get()),               # Адиабатический КПД компрессора (делим на 100)
                    "Степень повышения давления в компрессоре": replace_dot_with_comma(inp6.get()),      # Степень повышения давления в компрессоре
                    "Эффективная мощность": replace_dot_with_comma(inp100.get()),                        # Эффективная мощность
                }

                # Проходим по всем параграфам в документе
                for paragraph in doc.paragraphs:
                    # Проверяем, содержит ли параграф текст, который есть в словаре replacements
                    for key in replacements:
                        if key in paragraph.text:
                            

                            # Получаем XML-структуру параграфа
                            p_xml = paragraph._element.xml

                            # Парсим XML с помощью lxml
                            p_tree = etree.fromstring(p_xml)

                            # Ищем все формулы (элементы <m:t>)
                            formulas = p_tree.xpath('.//m:t', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})

                            # Проходим по каждой формуле в параграфе
                            for formula in formulas:
                                formula_text = formula.text
                                if formula_text and "=" in formula_text:
                                    

                                    # Разделяем текст на части до и после "="
                                    parts = formula_text.split("=", 1)
                                    if len(parts) == 2:
                                        # Находим маркер перед "="
                                        marker = parts[0].strip()
                                        print(f"Найден маркер: {marker}")  # Отладочное сообщение

                                        # Если маркер есть в replacements, добавляем данные после "=" без пробела
                                        if key in replacements:
                                            # Сохраняем исходный текст формулы до "="
                                            original_text = parts[0].strip()
                                            # Добавляем новые данные после "=" без пробела
                                            new_text = f"{original_text}={replacements[key]}"
                                           

                                            # Обновляем текст формулы
                                            formula.text = new_text

                            # Обновляем XML-структуру параграфа
                            paragraph._element.getparent().replace(paragraph._element, p_tree)

                # Сохраняем измененный документ
                doc.save('отчет_без_регенерации.docx')
              
            
        
        elif lang.get() == 'c регенерацией':
            
                # Открываем шаблон
                doc = Document('шаблон_регенерация.docx')

                # Словарь для замены значений после "=" в формулах
                replacements = {
                    "Температура окружающей среды": replace_dot_with_comma(inp1.get()),  # Температура окружающей среды
                    "Давление окружающей среды": replace_dot_with_comma(inp2.get()),      # Давление окружающей среды
                    "Начальная температура газа перед турбиной": replace_dot_with_comma(inp3.get()),  # Начальная температура газа перед турбиной
                    "Газовая постоянная воздуха": replace_dot_with_comma(inp4.get()),      # Газовая постоянная воздуха
                    "Показатель изоэнтропы для воздуха": replace_dot_with_comma(inp15.get()),  # Показатель изоэнтропы для воздуха
                    "Показатель изоэнтропы для газа": replace_dot_with_comma(inp16.get()),     # Показатель изоэнтропы для газа
                    "Теплоемкость воздуха": replace_dot_with_comma(inp17.get()),              # Теплоемкость воздуха
                    "Теплоемкость газа": replace_dot_with_comma(inp18.get()),                # Теплоемкость газа
                    "Коэффициент потерь на входе в компрессор": replace_dot_with_comma(inp5.get()),  # Коэффициент потерь на входе в компрессор (делим на 100)
                    "Коэффициент потерь давления воздуха в вых. устройстве": replace_dot_with_comma(inp7.get()),  # Коэффициент потерь давления воздуха в вых. устройстве (делим на 100)
                    "Коэффициент потерь давления воздуха перед КС": replace_dot_with_comma(inp8.get()),  # Коэффициент потерь давления воздуха перед КС (делим на 100)
                    "Коэффициент потерь давления воздуха в КС": replace_dot_with_comma(inp9.get()),      # Коэффициент потерь давления воздуха в КС (делим на 100)
                    "Политропный КПД турбины": divide_by_100(inp10.get()),                      # Политропный КПД турбины (делим на 100)
                    "Механический КПД турбины": divide_by_100(inp11.get()),                     # Механический КПД турбины (делим на 100)
                    "Механический КПД компрессора": divide_by_100(inp12.get()),                 # Механический КПД компрессора (делим на 100)
                    "Адиабатический КПД компрессора": divide_by_100(inp13.get()),               # Адиабатический КПД компрессора (делим на 100)
                    "Степень повышения давления в компрессоре": replace_dot_with_comma(inp6.get()),      # Степень повышения давления в компрессоре
                    "Эффективная мощность": replace_dot_with_comma(inp100.get()),                        # Эффективная мощность
                    "Степень рекуперации": replace_dot_with_comma(inp14.get()),                         # Степень рекуперации
                }

                # Проходим по всем параграфам в документе
                for paragraph in doc.paragraphs:
                    # Проверяем, содержит ли параграф текст, который есть в словаре replacements
                    for key in replacements:
                        if key in paragraph.text:
                      

                            # Получаем XML-структуру параграфа
                            p_xml = paragraph._element.xml

                            # Парсим XML с помощью lxml
                            p_tree = etree.fromstring(p_xml)

                            # Ищем все формулы (элементы <m:t>)
                            formulas = p_tree.xpath('.//m:t', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})

                            # Проходим по каждой формуле в параграфе
                            for formula in formulas:
                                formula_text = formula.text
                                if formula_text and "=" in formula_text:
                         

                                    # Разделяем текст на части до и после "="
                                    parts = formula_text.split("=", 1)
                                    if len(parts) == 2:
                                        # Находим маркер перед "="
                                        marker = parts[0].strip()
                           

                                        # Если маркер есть в replacements, добавляем данные после "=" без пробела
                                        if key in replacements:
                                            # Сохраняем исходный текст формулы до "="
                                            original_text = parts[0].strip()
                                            # Добавляем новые данные после "=" без пробела
                                            new_text = f"{original_text}={replacements[key]}"
                       

                                            # Обновляем текст формулы
                                            formula.text = new_text

                            # Обновляем XML-структуру параграфа
                            paragraph._element.getparent().replace(paragraph._element, p_tree)

                # Сохраняем измененный документ
                doc.save('отчет_с_регенерацией.docx')
  
            
           

        else:
            print("Режим расчета не выбран!")
    except Exception as e:
                print(f"Ошибка при заполнении шаблона: {e}")
    return

#==================================================================================================================================================
#==================================================================================================================================================
def fill_calculated_values():
    try:
        import re
        from lxml import etree

        # Определяем режим расчета
        if lang.get() == 'без регенерации':
            template_path = 'отчет_без_регенерации.docx'
        elif lang.get() == 'c регенерацией':
            template_path = 'отчет_с_регенерацией.docx'
        else:
            print("Режим расчета не выбран!")
            return

        # Функция для извлечения числа после знака '=' в тексте лейбла
        def extract_value(label_text):
            # Ищем значение после знака '=', включая отрицательные числа
            match = re.search(r"=\s*(-?[\d.,]+)", label_text)
            if match:
                # Возвращаем значение, заменяя точку на запятую
                return match.group(1).replace('.', ',')
            return ""  # Если значение не найдено, возвращаем пустую строку

        # Собираем значения из интерфейса
        calculated_values = {
            "Давление воздуха перед компрессором": extract_value(label6.cget("text")),
            "Температура воздуха перед компрессором": extract_value(label7.cget("text")),
            "Давление воздуха за компрессором": extract_value(label21.cget("text")),
            "Температура воздуха за компрессором": extract_value(label22.cget("text")),
            "Работа, соответствующая изоэнтропийному перепаду в компрессоре:": extract_value(label23.cget("text")),
            "Полезная работа в компрессоре": extract_value(label24.cget("text")),
            "Давление воздуха перед РВ": extract_value(label36.cget("text")),
            "Давление газа перед турбиной": extract_value(label50.cget("text")),
            "Давление газа за турбиной": extract_value(label25.cget("text")),
            "Степень расширения газа в турбине": extract_value(label26.cget("text")),
            "Работа, соответствующая изоэнтропийному перепаду в турбине:": extract_value(label27.cget("text")),
            "Полезная работа в турбине": extract_value(label28.cget("text")),
            "Температура газа за турбиной": extract_value(label29.cget("text")),
            "Расход воздуха через компрессор": extract_value(label30.cget("text")),
            "Температура воздуха перед камерой сгорания": extract_value(label31.cget("text")),
            "Теплота с учетом потерь в камере сгорания": extract_value(label32.cget("text")),
            "Расход теплоты": extract_value(label33.cget("text")),
            "Эффективная удельная работа": extract_value(label200.cget("text")),
            "Эффективный КПД установки": extract_value(label34.cget("text")),
            "Коэффициент полезной работы": extract_value(label35.cget("text")),
        }

        doc = Document(template_path)
        paragraphs = list(doc.paragraphs)  # Получаем список всех параграфов

        # Проходим по всем параграфам
        for i, paragraph in enumerate(paragraphs):
            # Проверяем каждый параметр из списка
            for param_name in calculated_values:
                if param_name in paragraph.text:
                    # Проверяем, есть ли следующий параграф
                    if i + 1 < len(paragraphs):
                        next_paragraph = paragraphs[i + 1]
                    

                        # Получаем XML следующего параграфа
                        next_p_xml = next_paragraph._element.xml
                        next_p_tree = etree.fromstring(next_p_xml)

                        # Ищем формулы с '@'
                        formulas = next_p_tree.xpath(
                            './/m:t[contains(text(), "@")]',
                            namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
                        )

                        # Заменяем '@' на '= значение'
                        value = calculated_values[param_name]
                        for formula in formulas:
                            new_text = formula.text.replace('@', f'= {value}')
                            formula.text = new_text
                    

                        # Обновляем XML следующего параграфа
                        next_paragraph._element.getparent().replace(next_paragraph._element, next_p_tree)

        # Сохраняем изменения
        doc.save(template_path)


    except Exception as e:
        print(f"Ошибка при сохранении расчетных значений: {e}")
#==================================================================================================================================================
def combined_save():
    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()
    
    # Запрашиваем путь сохранения у пользователя
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
        title="Сохранить файл как",
        initialfile="Расчет.docx"
    )
    
    # Закрываем Tkinter окно
    root.destroy()

    if not file_path:
        messagebox.showinfo("Информация", "Сохранение отменено пользователем.")
        print("Сохранение отменено пользователем.")
        return

    try:
        # Заполняем документ
        fill_template()
        fill_calculated_values()
        
        # Определяем шаблон в зависимости от выбранного языка
        template = 'отчет_без_регенерации.docx' if lang.get() == 'без регенерации' else 'отчет_с_регенерацией.docx'
        
        # Сохраняем DOCX
        doc = Document(template)
        doc.save(file_path)
        messagebox.showinfo("Успех", f"Документ успешно сохранен:\n{file_path}")
        print(f"Файл сохранен как {file_path}")
            
    except Exception as e:
        messagebox.showerror("Ошибка", 
            f"Не удалось сохранить документ:\n{str(e)}\n\n"
            "Проверьте доступ к файлам и попробуйте еще раз.")
        print(f"Ошибка при сохранении документа: {e}")
#======================================================================================================

import os
import tempfile
from tkinter import Tk, filedialog, messagebox
from customtkinter import CTkProgressBar

def combined_save_pdf():
    # Создаем прогресс-бар
    progress_bar = CTkProgressBar(save_right_frame, mode='indeterminate')
    progress_bar.grid(row=8, column=0, columnspan=3, padx=20, pady=10, sticky="ew")
    progress_bar.start()
    
    try:
        # Создаем скрытое окно Tkinter
        root = Tk()
        root.withdraw()
        
        # Запрашиваем путь сохранения у пользователя
        file_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
            title="Сохранить отчет как",
            initialfile="Расчет.pdf"
        )
        
        # Закрываем Tkinter окно
        root.destroy()

        if not file_path:
            messagebox.showinfo("Информация", "Сохранение отменено пользователем.")
            progress_bar.stop()
            progress_bar.destroy()
            return

        try:
            # Заполняем документ
            fill_template()
            fill_calculated_values()
            
            # Определяем шаблон в зависимости от выбранного языка
            template_name = 'отчет_без_регенерации.docx' if lang.get() == 'без регенерации' else 'отчет_с_регенерацией.docx'
            
            # Создаем имя для временной копии DOCX в директории программы
            temp_docx_path = os.path.join(os.path.dirname(__file__), 'temp_report.docx')
            
            try:
                # Загружаем шаблон и сохраняем как временную копию
                doc = Document(template_name)
                doc.save(temp_docx_path)
                
                # Конвертируем в PDF (с подавлением вывода в консоль)
                try:
                    # Перенаправляем stdout и stderr временно
                    import sys
                    from io import StringIO
                    
                    old_stdout = sys.stdout
                    old_stderr = sys.stderr
                    sys.stdout = StringIO()
                    sys.stderr = StringIO()
                    
                    try:
                        convert(temp_docx_path, file_path)
                        messagebox.showinfo("Успех", 
                            f"PDF версия отчета успешно создана:\n{file_path}")
                    except Exception as e:
                        messagebox.showwarning("Предупреждение", 
                            f"Документ сохранен, но конвертация в PDF не удалась:\n{str(e)}")
                    finally:
                        # Восстанавливаем stdout и stderr
                        sys.stdout = old_stdout
                        sys.stderr = old_stderr
                        
                except Exception as e:
                    messagebox.showerror("Ошибка", 
                        f"Ошибка при конвертации в PDF:\n{str(e)}")
                    
            finally:
                # Удаляем временную копию DOCX файла
                try:
                    if os.path.exists(temp_docx_path):
                        os.unlink(temp_docx_path)
                except Exception as e:
                    print(f"Не удалось удалить временный файл: {e}")
                
        except Exception as e:
            messagebox.showerror("Ошибка", 
                f"Не удалось сохранить документ:\n{str(e)}\n\n"
                "Проверьте доступ к файлам и попробуйте еще раз.")
        finally:
            # Останавливаем и удаляем прогресс-бар в любом случае
            progress_bar.stop()
            progress_bar.destroy()
            
    finally:
        # Останавливаем и удаляем прогресс-бар
        progress_bar.stop()
        progress_bar.destroy()

# Создание кнопок навигации
calc_btn = create_nav_button(1, icon=calc_icon, text="Вариантный расчет", command=show_calc)
grafik_btn = create_nav_button(2, icon=grafik_icon, text="Построение графиков", command=show_grafik)
save_btn = create_nav_button(3, icon=save_menu_icon, text="Сохранение", command=show_save)
settings_btn = create_nav_button(5, icon=settings_icon, text="Настройки", command=show_settings) 


# Убедимся, что при запуске меню свернуто
left_frame.configure(width=70)
calc_btn.configure(text=None, anchor="c")
grafik_btn.configure(text=None, anchor="c")
save_btn.configure(text=None, anchor="c")
settings_btn.configure(text=None, anchor="c")

# Пустое пространство для растяжения
empty_space = ctk.CTkFrame(left_frame, fg_color="transparent")
empty_space.grid(row=4, column=0, sticky="nsew")

# Начальное состояние
toggle_menu()
toggle_menu()
show_calc()  # По умолчанию показываем расчет, иконка будет белой

# Контент для фрейма расчетов (основной контейнер)
calc_content = ctk.CTkFrame(frame_calc, fg_color="#1f1f1f")
calc_content.pack(fill="both", expand=True, padx=(0,0), pady=(0,0))

# Контент для фрейма графиков (основной контейнер)
grafik_content = ctk.CTkFrame(frame_grafik, fg_color="#1f1f1f")
grafik_content.pack(fill="both", expand=True, padx=(0,0), pady=(0,0))

# Контент для фрейма settings (основной контейнер)
settings_content = ctk.CTkFrame(frame_settings, fg_color="#1f1f1f")
settings_content.pack(fill="both", expand=True, padx=(0,0), pady=(0,0))

save_content = ctk.CTkFrame(frame_save, fg_color="#1f1f1f")
save_content.pack(fill="both", expand=True, padx=(0,0), pady=(0,0))

# Правый фрейм для calc_content (заполняет всю область)
calc_right_frame = ctk.CTkFrame(calc_content, 
                              fg_color="#1f1f1f",
                              corner_radius=10)
calc_right_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)

# Правый фрейм для calc_content (заполняет всю область)
save_right_frame = ctk.CTkFrame(save_content, 
                              fg_color="#151515",
                              corner_radius=10, border_width=1, border_color="#454545")
save_right_frame.grid(row=0, column=0, sticky="nsew", padx=(40,0), pady=(12,0))
# Настройка весов для растягивания
save_content.grid_rowconfigure(0, weight=1)
save_content.grid_columnconfigure(0, weight=1)

# Настройка весов для растягивания
calc_content.grid_rowconfigure(0, weight=1)
calc_content.grid_columnconfigure(0, weight=1)
calc_right_frame.grid_rowconfigure(0, weight=1)
calc_right_frame.grid_columnconfigure(0, weight=1)

# Правый фрейм для grafik_content (заполняет всю область)
grafik_right_frame = ctk.CTkFrame(grafik_content, 
                              fg_color="#1f1f1f",
                              corner_radius=10)
grafik_right_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)

# Настройка весов для растягивания
grafik_content.grid_rowconfigure(0, weight=1)
grafik_content.grid_columnconfigure(0, weight=1)
grafik_right_frame.grid_rowconfigure(0, weight=1)
grafik_right_frame.grid_columnconfigure(0, weight=1)

# Правый фрейм для settings_content (заполняет всю область)
settings_right_frame = ctk.CTkFrame(settings_content, 
                              fg_color="#151515",
                              corner_radius=10, border_width=1, border_color="#454545")
settings_right_frame.grid(row=0, column=0, sticky="nsew", padx=(40,0), pady=(12,0))

# Настройка весов для растягивания
settings_content.grid_rowconfigure(0, weight=1)
settings_content.grid_columnconfigure(0, weight=1)

tk_font2 = ctk.CTkFont(size=16, weight='bold')
tk_font3 = ctk.CTkFont(size=13, weight="normal")

# Заполнение для save_right_frame
save_label1 = ctk.CTkLabel(save_right_frame, text="Параметры сохранения", font=tk_font1, text_color="white")
save_label1.grid(row=0, column=0, padx=20, pady=(15, 10), sticky="w")

save_label2 = ctk.CTkLabel(save_right_frame, text="Отчёт", font=tk_font2, text_color="#e6e6e6")
save_label2.grid(row=1, column=0, padx=20, pady=(15,0), sticky="w")

save_label3 = ctk.CTkLabel(save_right_frame, text="Выберите содержание отчёта и формат сохранения", font=tk_font3, text_color="#6a6a6a")
save_label3.grid(row=2, column=0, columnspan=5, padx=20, pady=0, sticky="w")

word_1 = ctk.CTkButton(save_right_frame, text="", width=200, height=200, corner_radius=10, fg_color="#151515", hover_color="#252525", image=word_save_icon)#73777d
word_1.grid(row=4, column=1, rowspan=4, padx=0, pady=0, sticky="w")

pdf_1 = ctk.CTkButton(save_right_frame, text="", width=200, height=200, corner_radius=10, fg_color="#151515", hover_color="#252525", image=pdf_save_icon)#73777d
pdf_1.grid(row=4, column=2, rowspan=4, padx=0, pady=0, sticky="w")


# Переменная для хранения выбранного режима
lang = StringVar(value='без регенерации')  # По умолчанию выбран режим "без регенерации"
# Variables for checkboxes
check_vars = [ctk.BooleanVar() for _ in range(4)]

# Checkboxes creation
# Column 0
ctk.CTkCheckBox(
    save_right_frame,
    text="Расчет",
    variable=check_vars[0],
    font=tk_font2,
    text_color="#e6e6e6", hover_color="#7fff00", border_color="#e6e6e6", fg_color="#7fff00", checkmark_color="#000000"
).grid(row=4, column=0, padx=20, pady=(10,5), sticky="w")

ctk.CTkCheckBox(
    save_right_frame,
    text="Графики 2D",
    variable=check_vars[1],
    font=tk_font2,
    text_color="#e6e6e6", hover_color="#7fff00", border_color="#e6e6e6", fg_color="#7fff00", checkmark_color="#000000"
).grid(row=5, column=0, padx=20, pady=5, sticky="w")

# Создаем чекбокс для 3D графиков, который будем блокировать
graph_3d_checkbox = ctk.CTkCheckBox(
    save_right_frame,
    text="Графики 3D",
    variable=check_vars[2],
    font=tk_font2,
    text_color="#e6e6e6", hover_color="#7fff00", border_color="#e6e6e6", fg_color="#7fff00", checkmark_color="#000000"
)
graph_3d_checkbox.grid(row=6, column=0, padx=20, pady=5, sticky="w")

# Чекбокс "Полный отчет"
full_report_checkbox = ctk.CTkCheckBox(
    save_right_frame,
    text="Полный отчет",
    variable=check_vars[3],
    font=tk_font2,
    text_color="#e6e6e6", hover_color="#7fff00", border_color="#e6e6e6", fg_color="#7fff00", checkmark_color="#000000",
    command=lambda: select_all(check_vars)
)
full_report_checkbox.grid(row=7, column=0, padx=20, pady=(15,5), sticky="we")

def select_all(vars_list):
    if lang.get() == 'без регенерации':
        # В режиме без регенерации включаем только расчет и 2D графики
        state = vars_list[3].get()
        vars_list[0].set(state)
        vars_list[1].set(state)
        vars_list[2].set(False)  # 3D всегда выключены в этом режиме
    else:
        # В других режимах включаем все
        state = vars_list[3].get()
        for i in range(3):
            vars_list[i].set(state)

def update_full_report(*args):
    if lang.get() == 'без регенерации':
        # В режиме без регенерации НЕ ставим галочку автоматически
        pass
    else:
        # В других режимах ставим галочку, если все выбрано
        all_selected = all(var.get() for var in check_vars[:3])
        check_vars[3].set(all_selected)
        
        # Если включен полный отчет, убедимся что все чекбоксы (включая 3D) включены
        if check_vars[3].get():
            for i in range(3):
                check_vars[i].set(True)

def update_checkboxes_state(*args):
    if lang.get() == 'без регенерации':
        graph_3d_checkbox.configure(state="disabled")
        check_vars[2].set(False)
        # Полный отчет доступен, но не ставится автоматически
        full_report_checkbox.configure(state="normal") 
    else:
        graph_3d_checkbox.configure(state="normal")
        full_report_checkbox.configure(state="normal")
        
        # Если был выбран полный отчет в предыдущем режиме, убедимся что 3D тоже выбраны
        if check_vars[3].get():
            check_vars[2].set(True)

# Привязываем изменение режима расчета
lang.trace_add("write", update_checkboxes_state)

# Инициализация состояния
update_checkboxes_state()

def update_button_params(*args):
    calculation = check_vars[0].get()
    graph2ddd = check_vars[1].get()
    graph3ddd = check_vars[2].get() if lang.get() != 'без регенерации' else False
    fulll_report = check_vars[3].get()

    # Определяем режим экспорта
    if lang.get() == 'без регенерации':
        if fulll_report:
            mode = "calc_2d"  # В этом режиме полный отчет = расчет + 2D
        elif calculation and graph2ddd:
            mode = "calc_2d"
        elif calculation:
            mode = "calc"
        elif graph2ddd:
            mode = "2d"
        else:
            mode = "none"
    else:
        if fulll_report or (calculation and graph2ddd and graph3ddd):
            mode = "full"
        elif calculation and graph2ddd:
            mode = "calc_2d"
        elif calculation and graph3ddd:
            mode = "calc_3d"
        elif graph2ddd and graph3ddd:
            mode = "graphs_2d_3d"
        elif calculation:
            mode = "calc"
        elif graph2ddd:
            mode = "2d"
        elif graph3ddd:
            mode = "3d"
        else:
            mode = "none"

    configure_export_buttons(mode)

def configure_export_buttons(mode):
    # Настройки для кнопок Word и PDF
    if mode == "full":
        word_1.configure(command=save_word_report_full, state="normal")
        pdf_1.configure(command=save_pdf_report_full, state="normal")
    elif mode == "calc_2d":
        word_1.configure(command=save_word_report_2d, state="normal")
        pdf_1.configure(command=save_pdf_report_2d, state="normal")
    elif mode == "calc_3d":
        word_1.configure(command=save_word_report_3d, state="normal")
        pdf_1.configure(command=save_pdf_report_3d, state="normal")
    elif mode == "graphs_2d_3d":
        word_1.configure(command=save_plots_as_word_combined, state="normal")
        pdf_1.configure(command=save_plot_all_plots_pdf, state="normal")
    elif mode == "calc":
        word_1.configure(command=combined_save, state="normal")
        pdf_1.configure(command=combined_save_pdf, state="normal")
    elif mode == "2d":
        word_1.configure(command=save_plot_as_word_2d_all, state="normal")
        pdf_1.configure(command=save_plot_as_pdf_2d_all, state="normal")
    elif mode == "3d":
        word_1.configure(command=save_plot_as_word_3d_all, state="normal")
        pdf_1.configure(command=save_plot_as_pdf_3d_all, state="normal")
    else:
        word_1.configure(command=None, state="disabled")
        pdf_1.configure(command=None, state="disabled")

# Привязка функции обновления кнопок ко всем чекбоксам
for var in check_vars:
    var.trace_add("write", update_button_params)


save_right_frame.grid_rowconfigure(0, weight=0)
save_right_frame.grid_rowconfigure(4, weight=0)
save_right_frame.grid_rowconfigure(5, weight=0)
save_right_frame.grid_rowconfigure(6, weight=0)
save_right_frame.grid_rowconfigure(7, weight=0)
save_right_frame.grid_rowconfigure(8, weight=0)
save_right_frame.grid_rowconfigure(9, weight=1)
save_right_frame.grid_columnconfigure(0, weight=0)
save_right_frame.grid_columnconfigure(1, weight=0)
save_right_frame.grid_columnconfigure(2, weight=0)
save_right_frame.grid_columnconfigure(3, weight=1)
# Линия separator
separator = ctk.CTkFrame(save_right_frame, height=2, fg_color="gray")
separator.grid(row=3, column=0, columnspan=4, sticky="ew", padx=20, pady=10)

# Заполнение для settings_right_frame
settings_label1 = ctk.CTkLabel(settings_right_frame, text="Настройка приложения", font=tk_font1, text_color="white")
settings_label1.grid(row=0, column=0, padx=20, pady=(15, 10), sticky="w")

settings_label2 = ctk.CTkLabel(settings_right_frame, text="Тема", font=tk_font2, text_color="#e6e6e6")
settings_label2.grid(row=1, column=0, padx=20, pady=(15,0), sticky="w")

settings_label3 = ctk.CTkLabel(settings_right_frame, text="Изменить внешний вид и восприятие вашей программы", font=tk_font3, text_color="#6a6a6a")
settings_label3.grid(row=2, column=0, padx=20, pady=0, sticky="w")

# Линия separator
separator = ctk.CTkFrame(settings_right_frame, height=2, fg_color="gray")
separator.grid(row=3, column=0, columnspan=4, sticky="ew", padx=20, pady=10)

settings_label4 = ctk.CTkLabel(settings_right_frame, text="Тема интерфейса", font=tk_font3, text_color="#e6e6e6")
settings_label4.grid(row=4, column=0, padx=20, pady=(5,0), sticky="w")

tema_button = ctk.CTkButton(settings_right_frame, height=60, width=240, text="", image = tema_icon, corner_radius=10, fg_color="#151515", hover_color="#252525", anchor="center", command=toggle_color_theme)
tema_button.grid(row=4, column=1, rowspan=3, sticky="nsew", padx=5)

tema_button1 = ctk.CTkButton(settings_right_frame, height=60, width=240, text="", image = tema1_icon, corner_radius=10, fg_color="#151515", hover_color="#252525", anchor="center", command=toggle_color_theme)
tema_button1.grid(row=4, column=2, rowspan=3, sticky="nsew", padx=5)

settings_label5 = ctk.CTkLabel(settings_right_frame, text="Выберите или кастомизируйте тему приложения", font=tk_font3, text_color="#6a6a6a")
settings_label5.grid(row=5, column=0, padx=20, pady=(0,0), sticky="w")

# Пустая строка (добавлен пустой фрейм)
empty_row = ctk.CTkFrame(settings_right_frame, height=100, fg_color="transparent")
empty_row.grid(row=6, column=0, pady=10)

# Пустая строка (добавлен пустой фрейм)
empty_row1 = ctk.CTkFrame(settings_right_frame, height=20, fg_color="transparent")
empty_row1.grid(row=13, column=0, pady=10)

# Настройка grid для правильного отображения
settings_right_frame.grid_rowconfigure(4, weight=0) 
settings_right_frame.grid_rowconfigure(5, weight=0) 
settings_right_frame.grid_rowconfigure(6, weight=0) 
settings_right_frame.grid_rowconfigure(7, weight=0) 
settings_right_frame.grid_rowconfigure(14, weight=1)  # Пустое пространство внизу
settings_right_frame.grid_columnconfigure(0, weight=0)
settings_right_frame.grid_columnconfigure(1, weight=0)
settings_right_frame.grid_columnconfigure(2, weight=0)
settings_right_frame.grid_columnconfigure(3, weight=1)

# Контейнер для двух фреймов (горизонтальное расположение)
dual_frame_container = ctk.CTkFrame(calc_right_frame, 
                                  fg_color="transparent")
dual_frame_container.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)

# Настройка колонок для dual_frame_container
dual_frame_container.grid_rowconfigure(0, weight=1)
dual_frame_container.grid_columnconfigure(0, weight=0)  # Левый фрейм - фиксированная ширина
dual_frame_container.grid_columnconfigure(1, weight=1)  # Правый фрейм - растягивается

# Контейнер для двух фреймов (горизонтальное расположение)
dual_frame_container1 = ctk.CTkFrame(grafik_right_frame, 
                                  fg_color="transparent")
dual_frame_container1.grid(row=0, column=0, sticky="nsew", padx=0, pady=10)

# Настройка колонок для dual_frame_container
dual_frame_container1.grid_rowconfigure(0, weight=1)
dual_frame_container1.grid_columnconfigure(0, weight=0)  # Левый фрейм - фиксированная ширина
dual_frame_container1.grid_columnconfigure(1, weight=1)  # Правый фрейм - растягивается

# Левый узкий фрейм (Методика расчета)
left_method_frame = ctk.CTkFrame(dual_frame_container, 
                               width=200,
                               fg_color="#1f1f1f")
left_method_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
left_method_frame.grid_propagate(False)

# Настройка содержимого левого фрейма
left_method_frame.grid_rowconfigure(0, weight=0)  # Метка
left_method_frame.grid_rowconfigure(1, weight=0)  # Радио-кнопка 1
left_method_frame.grid_rowconfigure(2, weight=0)  # Радио-кнопка 2
left_method_frame.grid_rowconfigure(3, weight=0)  # Кнопка старт
left_method_frame.grid_rowconfigure(4, weight=0)  # Кнопка clear
left_method_frame.grid_rowconfigure(5, weight=0)  # Кнопка сохранения
left_method_frame.grid_rowconfigure(6, weight=0)  
left_method_frame.grid_rowconfigure(7, weight=1)  # Пустое пространство

left_method_frame.grid_columnconfigure(0, weight=1)  
left_method_frame.grid_columnconfigure(1, weight=0)  

# Левый узкий фрейм (Методика расчета)
left_method_frame1 = ctk.CTkFrame(dual_frame_container1, 
                               width=360,
                               fg_color="#1f1f1f")
left_method_frame1.grid(row=0, column=0, sticky="nsew", padx=0)
left_method_frame1.grid_propagate(False)

# В настройках left_method_frame1:
left_method_frame1.grid_rowconfigure(1, weight=1)  # Кнопка 1
left_method_frame1.grid_rowconfigure(3, weight=1)  # Кнопка 2
left_method_frame1.grid_rowconfigure(5, weight=1)  # Кнопка 3
left_method_frame1.grid_rowconfigure(7, weight=1)  # Кнопка 4
left_method_frame1.grid_columnconfigure(0, weight=1)  # Единственная колонка
left_method_frame1.grid_columnconfigure(1, weight=0)

tk_font4 = ctk.CTkFont(size=16, weight="bold")
method_label = ctk.CTkLabel(left_method_frame, 
                          text="Методика расчета", anchor="w", font=tk_font4)
method_label.grid(row=0, column=0, pady=(10, 20), padx=10, sticky="w")

method_var = StringVar(value="без регенерации")

# # Переменная для хранения выбранного режима
# lang = StringVar(value='без регенерации')  # По умолчанию выбран режим "без регенерации"

# Радиокнопки для выбора режима
withoutreg = 'без регенерации'
regeneration = 'c регенерацией'

without_btn = ctk.CTkRadioButton(left_method_frame, text=withoutreg, value=withoutreg, variable=lang,
                                 command=lambda: vvod('без регенерации'))
without_btn.grid(column=0, row=1, padx=(10,10), pady=2, sticky="ew")

reg_btn = ctk.CTkRadioButton(left_method_frame, text=regeneration, value=regeneration, variable=lang,
                             command=lambda: vvod('c регенерацией'))
reg_btn.grid(row=2, column=0, pady=(2, 10), padx=(10,10), sticky="ew")

btn11 = ctk.CTkButton(left_method_frame, text='', command=raschet, image=get_icon, fg_color="#a2f200", hover_color="#7fff00", height=40)
btn11.grid(row=3, column=0, pady=(20, 10), padx=(10,10), sticky="ew")

btn1_1 = ctk.CTkButton(left_method_frame, text='', image=cs_icon, command=clear, fg_color="#505052", hover_color="#73777d", height=30)
btn1_1.grid(row=4, column=0, pady=(2, 10), padx=(10,10), sticky="ew")

save_button = ctk.CTkButton(left_method_frame, text='', command=combined_save, image=word_icon, width=140, height=30, fg_color="#009dda", hover_color="#2dcbff", corner_radius=5)
save_button.grid(column=0, row=5, padx=(10,10), pady=(5, 10), sticky="ew")

save_button_pdf = ctk.CTkButton(left_method_frame, text='', command=combined_save_pdf, image=pdf_icon, width=140, height=30, fg_color="#f97666", hover_color="#fba297", corner_radius=5)
save_button_pdf.grid(column=0, row=6, padx=(10,10), pady=(0, 10), sticky="ew")


# Кнопка-переключатель (switch) для смены темы
# theme_switch = ctk.CTkSwitch(
#     left_method_frame,
#     text="Темная тема" if ctk.get_appearance_mode() == "Light" else "Светлая тема",
#     command=toggle_color_theme,
#     progress_color="#00ff00",  # Зеленый цвет переключателя
#     button_color="#ffffff",  # Цвет кнопки переключателя
#     button_hover_color="#cccccc",  # Цвет кнопки при наведении
# )
# theme_switch.grid(column=0, row=6, padx=10, pady=(2, 10), sticky="w")

right_content_frame = ctk.CTkScrollableFrame(dual_frame_container, 
                                 fg_color="#252525",
                                 orientation="vertical",
                                 corner_radius=15,
                                 border_width=1,
                                 scrollbar_button_color="#606060", 
                                 scrollbar_button_hover_color="#808080",
                                 border_color="#454545")
right_content_frame.grid(row=0, column=1, sticky="nsew")

# Настройка правого фрейма Grafik
right_content_frame.grid_rowconfigure(0, weight=1)

right_content_frame.grid_columnconfigure(0, weight=1)  # Первая колонка - растягивается
right_content_frame.grid_columnconfigure(1, weight=0)  # Вторая колонка - фиксированная (для полей ввода)
right_content_frame.grid_columnconfigure(2, weight=0)  # Третья колонка - фиксированная (для кнопок ввода)
right_content_frame.grid_columnconfigure(3, weight=0)  # Четвертая колонка - фиксированная (для кнопок удаления)

right_content_frame1 = ctk.CTkFrame(dual_frame_container1, 
                                 fg_color="#1f1f1f",
                                #  orientation="vertical",
                                 corner_radius=15,
                                 border_width=1,
                                #  scrollbar_button_color="#606060", 
                                #  scrollbar_button_hover_color="#808080",
                                 border_color="#1f1f1f")
right_content_frame1.grid(row=0, column=1, sticky="nsew")




# Линии
separator2 = ctk.CTkFrame(right_content_frame, height=2, fg_color="gray")
separator2.grid(column=0, row=1, columnspan=4, sticky="ew", pady=2, padx=5)

separator2 = ctk.CTkFrame(right_content_frame, height=2, fg_color="gray")
separator2.grid(column=0, row=7, columnspan=4, sticky="ew", pady=2, padx=5)

separator2 = ctk.CTkFrame(right_content_frame, height=2, fg_color="gray")
separator2.grid(column=0, row=11, columnspan=4, sticky="ew", pady=2, padx=5)

separator2 = ctk.CTkFrame(right_content_frame, height=2, fg_color="gray")
separator2.grid(column=0, row=15, columnspan=4, sticky="ew", pady=2, padx=5)

separator2 = ctk.CTkFrame(right_content_frame, height=2, fg_color="gray")
separator2.grid(column=0, row=21, columnspan=4, sticky="ew", pady=2, padx=5)

separator2 = ctk.CTkFrame(right_content_frame, height=2, fg_color="gray")
separator2.grid(column=0, row=27, columnspan=4, sticky="ew", pady=2, padx=5)

# separator2 = ctk.CTkFrame(right_content_frame, width=2, fg_color="gray")
# separator2.grid(column=4, row=0, rowspan=31, sticky="ns", padx=(100,10), pady=0)
# Создаем фрейм сразу (но скрываем его)
# Создаем фрейм с проверкой текущей темы
separ_frame = None

def create_separ_frame():
    global separ_frame
    
    # Определяем текущую тему
    current_theme = ctk.get_appearance_mode()
    fg_color = "#252525" if current_theme == "Dark" else "#414142"
    border_color = "#454545" if current_theme == "Dark" else "#b0b0b0"
    
    # Создаем фрейм с текущими настройками темы
    separ_frame = ctk.CTkFrame(
        right_content_frame,
        fg_color=fg_color,
        border_color=border_color,
        width=200,
        height=100
    )
    separ_frame.grid(row=0, column=4, rowspan=31, sticky="ns", padx=(10,10), pady=0)
    separ_frame.grid_remove()  # Сразу скрываем

def toggle_separ_frame():
    if separ_frame.winfo_ismapped():
        separ_frame.grid_remove()
        separ.configure(image=galka_icon)
    else:
        # Обновляем тему перед показом
        update_separ_frame_theme()
        separ_frame.grid()
        separ.configure(image=galka_active_icon)

def update_separ_frame_theme():
    if not separ_frame or not separ_frame.winfo_exists():
        return
        
    current_theme = ctk.get_appearance_mode()
    if current_theme == "Dark":
        separ_frame.configure(fg_color="#252525", border_color="#454545")
    else:
        separ_frame.configure(fg_color="#414142", border_color="#b0b0b0")

# Создаем фрейм при запуске
create_separ_frame()

# Кнопка separ
separ = ctk.CTkButton(
    right_content_frame, 
    text='', 
    image=galka_icon,
    fg_color="#252525", 
    hover_color="#505052", 
    width=30,
    command=toggle_separ_frame
)
separ.grid(column=3, row=0, rowspan=1, sticky="e", padx=(10,10), pady=0)

# Создание жирного шрифта
bold_font = ctk.CTkFont(size=15, weight='bold')


# Начальные параметры
lb1 = ctk.CTkLabel(right_content_frame, text='Начальные параметры', font=bold_font, width=500, anchor="w", text_color="white")
lb1.grid(column=0, row=0, padx=5, pady=(5,0), sticky="w")

label1 = ctk.CTkLabel(right_content_frame, text='Температура окружающей среды Tн, K', width=500, anchor="w", text_color="#bfbfbf")
label1.grid(column=0, row=2, padx=5, pady=5, sticky="w")

inp1 = ctk.CTkEntry(right_content_frame)
inp1.grid(column=1, row=2, padx=5, pady=5, sticky="e")

btn1 = ctk.CTkButton(right_content_frame, text='', command=get1, image=inp_icon)
btn1.grid(column=2, row=2, padx=5, pady=5, sticky="e")

btn2 = ctk.CTkButton(
    right_content_frame,
    text='',  # Текст пустой, чтобы отображалось только изображение
    image=trash_icon,  # Изображение корзины
    command=delete1,
    width=140,
    height=25,
    fg_color="#505052",
    hover_color="#9b2d30",  # Цвет при наведении
    corner_radius=5  # Закругленные углы
)

btn2.grid(column=3, row=2, padx=(5,5), pady=5, sticky="e")

# Давление окружающей среды
label2 = ctk.CTkLabel(right_content_frame, text='Давление окружающей среды Pн, Па', width=500, anchor="w", text_color="#bfbfbf")
label2.grid(column=0, row=3, padx=5, pady=5, sticky="w")

inp2 = ctk.CTkEntry(right_content_frame)
inp2.grid(column=1, row=3, padx=5, pady=5, sticky="e")

btn3 = ctk.CTkButton(right_content_frame, text='', command=get2, image=inp_icon)
btn3.grid(column=2, row=3, padx=5, pady=5, sticky="e")

btn4 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete2, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5)
btn4.grid(column=3, row=3, padx=5, pady=5, sticky="e")


# Начальная температура газа
label3 = ctk.CTkLabel(right_content_frame, text='Начальная температура газа перед турбиной T3, K', width=500, anchor="w", text_color="#bfbfbf")
label3.grid(column=0, row=4, padx=5, pady=5, sticky="w")

inp3 = ctk.CTkEntry(right_content_frame)
inp3.grid(column=1, row=4, padx=5, pady=5, sticky="e")

btn5 = ctk.CTkButton(right_content_frame, text='', command=get3, image=inp_icon)
btn5.grid(column=2, row=4, padx=5, pady=5, sticky="e")

btn6 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete3, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5)
btn6.grid(column=3, row=4, padx=5, pady=5, sticky="e")

label4 = ctk.CTkLabel(right_content_frame, text='Газовая постоянная для воздуха R, Дж/(кг*К)', width=500, anchor="w", text_color="#bfbfbf")
label4.grid(column=0,row=5, padx=5, pady=5, sticky="w")

inp4 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp4.grid(column=1,row=5, padx=5, pady=5, sticky="e")

btn7 = ctk.CTkButton(right_content_frame, text='', command=get4, image=inp_icon) #Кнопка
btn7.grid(column=2,row=5, padx=5, pady=5, sticky="e")

btn8 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete4, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn8.grid(column=3,row=5, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
lb2 = ctk.CTkLabel(right_content_frame, text='Коэффициенты потерь', font=bold_font, width=500, anchor="w", text_color="white")
lb2.grid(column=0, row=14, padx=5, pady=(15,5), sticky="w")

label5 = ctk.CTkLabel(right_content_frame, text='Коэффициент потерь на входе в компрессор σвх', width=500, anchor="w", text_color="#bfbfbf")
label5.grid(column=0,row=16, padx=5, pady=5, sticky="w")

inp5 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp5.grid(column=1,row=16, padx=5, pady=5, sticky="e")

btn9 = ctk.CTkButton(right_content_frame, text='', command=get5, image=inp_icon) #Кнопка
btn9.grid(column=2,row=16, padx=5, pady=5, sticky="e")

btn10 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete5, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn10.grid(column=3,row=16, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
label8 = ctk.CTkLabel(right_content_frame, text='πк - степень повышения давления', width=500, anchor="w", text_color="#bfbfbf")
label8.grid(column=0,row=28, padx=5, pady=5, sticky="w")

inp6 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp6.grid(column=1,row=28, padx=5, pady=5, sticky="e")

btn12 = ctk.CTkButton(right_content_frame, text='', command=get6, image=inp_icon) #Кнопка
btn12.grid(column=2,row=28, padx=5, pady=5, sticky="e")

btn13 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete6, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn13.grid(column=3,row=28, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
label9 = ctk.CTkLabel(right_content_frame, text='Коэффициент потерь давления воздуха в выходном устройстве σвых*', width=500, anchor="w", text_color="#bfbfbf")
label9.grid(column=0,row=17, padx=5, pady=5, sticky="w")

inp7 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp7.grid(column=1,row=17, padx=5, pady=5, sticky="e")

btn14 = ctk.CTkButton(right_content_frame, text='', command=get7, image=inp_icon) #Кнопка
btn14.grid(column=2,row=17, padx=5, pady=5, sticky="e")

btn15 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete7, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn15.grid(column=3,row=17, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
label10 = ctk.CTkLabel(right_content_frame, text='Коэффициент потерь давления воздуха перед кс σvtepl*', width=500, anchor="w", text_color="#bfbfbf")
label10.grid(column=0,row=18, padx=5, pady=5, sticky="w")

inp8 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp8.grid(column=1,row=18, padx=5, pady=5, sticky="e")

btn16 = ctk.CTkButton(right_content_frame, text='', command=get8, image=inp_icon) #Кнопка
btn16.grid(column=2,row=18, padx=5, pady=5, sticky="e")

btn17 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete8, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn17.grid(column=3,row=18, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
label11 = ctk.CTkLabel(right_content_frame, text='Коэффициент потерь давления воздуха в кс σкс', width=500, anchor="w", text_color="#bfbfbf")
label11.grid(column=0,row=19, padx=5, pady=5, sticky="w")

inp9 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp9.grid(column=1,row=19, padx=5, pady=5, sticky="e")

btn18 = ctk.CTkButton(right_content_frame, text='', command=get9, image=inp_icon) #Кнопка
btn18.grid(column=2,row=19, padx=5, pady=5, sticky="e")

btn19 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete9, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn19.grid(column=3,row=19, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
lb2 = ctk.CTkLabel(right_content_frame, text='Коэффициенты полезного действия', width=500, anchor="w", font=bold_font, text_color="white")
lb2.grid(column=0, row=20, padx=5, pady=(15,5), sticky="w")

label12 = ctk.CTkLabel(right_content_frame, text='Политропный КПД турбины ηпол %', width=500, anchor="w", text_color="#bfbfbf")
label12.grid(column=0,row=22, padx=5, pady=5, sticky="w")

inp10 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp10.grid(column=1,row=22, padx=5, pady=5, sticky="e")

btn20 = ctk.CTkButton(right_content_frame, text='', command=get10, image=inp_icon) #Кнопка
btn20.grid(column=2,row=22, padx=5, pady=5, sticky="e")

btn21 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete10, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn21.grid(column=3,row=22, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
label13 = ctk.CTkLabel(right_content_frame, text='Механический КПД турбины ηмт %', width=500, anchor="w", text_color="#bfbfbf")
label13.grid(column=0,row=23, padx=5, pady=5, sticky="w")

inp11 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp11.grid(column=1,row=23, padx=5, pady=5, sticky="e")

btn22 = ctk.CTkButton(right_content_frame, text='', command=get11, image=inp_icon) #Кнопка
btn22.grid(column=2,row=23, padx=5, pady=5, sticky="e")

btn23 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete11, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn23.grid(column=3,row=23, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
label14 = ctk.CTkLabel(right_content_frame, text='Механический КПД компрессора ηмк %', width=500, anchor="w", text_color="#bfbfbf")
label14.grid(column=0,row=24, padx=5, pady=5, sticky="w")

inp12 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp12.grid(column=1,row=24, padx=5, pady=5, sticky="e")

btn24 = ctk.CTkButton(right_content_frame, text='', command=get12, image=inp_icon) #Кнопка
btn24.grid(column=2,row=24, padx=5, pady=5, sticky="e")

btn25 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete12, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn25.grid(column=3,row=24, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
label15 = ctk.CTkLabel(right_content_frame, text='Адиабатический КПД компрессора ηкад %', width=500, anchor="w", text_color="#bfbfbf")
label15.grid(column=0,row=25, padx=5, pady=5, sticky="w")

inp13 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp13.grid(column=1,row=25, padx=5, pady=5, sticky="e")

btn26 = ctk.CTkButton(right_content_frame, text='', command=get13, image=inp_icon) #Кнопка
btn26.grid(column=2,row=25, padx=5, pady=5, sticky="e")

btn27 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete13, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn27.grid(column=3,row=25, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
lb3 = ctk.CTkLabel(right_content_frame, text='Прочее', font=bold_font, width=500, anchor="w", text_color="white")
lb3.grid(column=0, row=26, padx=5, pady=(15,5), sticky="w")

label16 = ctk.CTkLabel(right_content_frame, text='Степень рекуперации μ %', text_color="#616161", width=500, anchor="w")
label16.grid(column=0,row=30, padx=5, pady=5, sticky="w")

inp14 = ctk.CTkEntry(right_content_frame, state="readonly") #Модуль добавления текстового поля ввода
inp14.grid(column=1,row=30, padx=5, pady=5, sticky="e")

btn28 = ctk.CTkButton(right_content_frame, text='', command=get14, image=inp_icon) #Кнопка
btn28.grid(column=2,row=30, padx=5, pady=5, sticky="e")

btn29 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete14, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn29.grid(column=3,row=30, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
label100 = ctk.CTkLabel(right_content_frame, text='Эффективная мощность Ne, кВт', width=500, anchor="w", text_color="#bfbfbf")
label100.grid(column=0,row=29, padx=5, pady=5, sticky="w")

inp100 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp100.grid(column=1,row=29, padx=5, pady=5, sticky="e")

btn100 = ctk.CTkButton(right_content_frame, text='', command=get100, image=inp_icon) #Кнопка
btn100.grid(column=2,row=29, padx=5, pady=5, sticky="e")

btn101 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete100, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn101.grid(column=3,row=29, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
lb7 = ctk.CTkLabel(right_content_frame, text='Показатели изоэнтропы', width=500, anchor="w", font=bold_font, text_color="white")
lb7.grid(column=0, row=6, padx=5, pady=(15,5), sticky="w")

label17 = ctk.CTkLabel(right_content_frame, text='Показатель изоэнтропы для воздуха Kв', width=500, anchor="w", text_color="#bfbfbf")
label17.grid(column=0,row=8, padx=5, pady=5, sticky="w")

inp15 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp15.grid(column=1,row=8, padx=5, pady=5, sticky="e")

btn30 = ctk.CTkButton(right_content_frame, text='', command=get15, image=inp_icon) #Кнопка
btn30.grid(column=2,row=8, padx=5, pady=5, sticky="e")

btn31 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete15, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn31.grid(column=3,row=8, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
label18 = ctk.CTkLabel(right_content_frame, text='Показатель изоэнтропы для газа Kг', width=500, anchor="w", text_color="#bfbfbf")
label18.grid(column=0,row=9, padx=5, pady=5, sticky="w")

inp16 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp16.grid(column=1,row=9, padx=5, pady=5, sticky="e")

btn32 = ctk.CTkButton(right_content_frame, text='', command=get16, image=inp_icon) #Кнопка
btn32.grid(column=2,row=9, padx=5, pady=5, sticky="e")

btn33 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete16, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn33.grid(column=3,row=9, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
lb6 = ctk.CTkLabel(right_content_frame, text='Изобарные удельные теплоемкости', font=bold_font, width=500, anchor="w", text_color="white")
lb6.grid(column=0, row=10, padx=5, pady=(15,5), sticky="w")

label19 = ctk.CTkLabel(right_content_frame, text='Теплоёмкость воздуха Cв, Дж/(кг*К)', width=500, anchor="w", text_color="#bfbfbf")
label19.grid(column=0,row=12, padx=5, pady=5, sticky="w")

inp17 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp17.grid(column=1,row=12, padx=5, pady=5, sticky="e")

btn34 = ctk.CTkButton(right_content_frame, text='', command=get17, image=inp_icon) #Кнопка
btn34.grid(column=2,row=12, padx=5, pady=5, sticky="e")

btn35 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete17, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn35.grid(column=3,row=12, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
label20 = ctk.CTkLabel(right_content_frame, text='Теплоёмкость газа Cг, Дж/(кг*К)', width=500, anchor="w", text_color="#bfbfbf")
label20.grid(column=0,row=13, padx=5, pady=5, sticky="w")

inp18 = ctk.CTkEntry(right_content_frame) #Модуль добавления текстового поля ввода
inp18.grid(column=1,row=13, padx=5, pady=5, sticky="e")

btn36 = ctk.CTkButton(right_content_frame, text='', command=get18, image=inp_icon) #Кнопка
btn36.grid(column=2,row=13, padx=5, pady=5, sticky="e")

btn37 = ctk.CTkButton(right_content_frame, text='', image=trash_icon, command=delete18, width=140, height=25, fg_color="#505052", hover_color="#9b2d30", corner_radius=5) #Кнопка
btn37.grid(column=3,row=13, padx=5, pady=5, sticky="e")
#==================================================================================================================================================
#Блок расчета параметров
#==================================================================================================================================================

lb4 = ctk.CTkLabel(separ_frame, text='Методика расчета',text_color="white", anchor="e", font=bold_font)
lb4.grid(column=5, row=0, padx=(20,5), pady=(5,0), sticky="w")

separator2 = ctk.CTkFrame(separ_frame, height=2, fg_color="gray")
separator2.grid(column=5, row=1, columnspan=1, sticky="ew", pady=2, padx=(20,0))

label6 = ctk.CTkLabel(separ_frame, text='Давление воздуха перед компрессором P1*, Па', text_color="#bfbfbf")
label6.grid(column=5,row=2, padx=(20,5), pady=5, sticky="w")

#==================================================================================================================================================
label7 = ctk.CTkLabel(separ_frame, text='Температура воздуха перед компрессором T1*, К', text_color="#bfbfbf")
label7.grid(column=5,row=3, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label21 = ctk.CTkLabel(separ_frame, text='Давление воздуха за компрессором P2*, Па', text_color="#bfbfbf")
label21.grid(column=5,row=4, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label22 = ctk.CTkLabel(separ_frame, text='Температура воздуха за компрессором T2*, К', text_color="#bfbfbf")
label22.grid(column=5,row=5, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label23 = ctk.CTkLabel(separ_frame, text='Работа изоэнтропийного перепада в компрессоре Hok*, кДж/кг', text_color="#bfbfbf")
label23.grid(column=5,row=6, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label24 = ctk.CTkLabel(separ_frame, text='Полезная работа в копрессоре Hk, кДж/кг', text_color="#bfbfbf")
label24.grid(column=5,row=7, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label36 = ctk.CTkLabel(separ_frame, text='', text_color="#bfbfbf")
label36.grid(column=5,row=8, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label50 = ctk.CTkLabel(separ_frame, text='Давление газа перед турбиной P3*, Па', text_color="#bfbfbf")
label50.grid(column=5,row=9, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label25 = ctk.CTkLabel(separ_frame, text='Давление газа за турбиной P4*, Па', text_color="#bfbfbf")
label25.grid(column=5,row=10, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label26 = ctk.CTkLabel(separ_frame, text='Степень расширения газа в турбине πт*', text_color="#bfbfbf")
label26.grid(column=5,row=11, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label27 = ctk.CTkLabel(separ_frame, text='Работа изоэнтропийного перепада в турбине Hот*, кДж/кг', text_color="#bfbfbf")
label27.grid(column=5,row=12, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label28 = ctk.CTkLabel(separ_frame, text='Полезная работа в турбине Hт, кДж/кг', text_color="#bfbfbf")
label28.grid(column=5,row=13, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label29 = ctk.CTkLabel(separ_frame, text='Температура газа за турбиной T4*, К', text_color="#bfbfbf")
label29.grid(column=5,row=14, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label30 = ctk.CTkLabel(separ_frame, text='Расход воздуха через компрессор Gв, кг/с', text_color="#bfbfbf")
label30.grid(column=5,row=15, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label31 = ctk.CTkLabel(separ_frame, text='Температура воздуха перед кс, К', text_color="#bfbfbf")
label31.grid(column=5,row=16, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label32 = ctk.CTkLabel(separ_frame, text="Расход теплоты с учетом потерь тепла в кс Q1', кДж/кг", text_color="#bfbfbf")
label32.grid(column=5,row=17, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label33 = ctk.CTkLabel(separ_frame, text='Расход теплоты Q1, кДж/кг', text_color="#bfbfbf")
label33.grid(column=5,row=18, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
lb8 = ctk.CTkLabel(separ_frame, text='КПД и КПР', font=bold_font, text_color="white")
lb8.grid(column=5, row=19, padx=(20,5), pady=(15,5), sticky="w")

separator2 = ctk.CTkFrame(separ_frame, height=2, fg_color="gray")
separator2.grid(column=5, row=20, columnspan=1, sticky="ew", pady=2, padx=(20,0))

label34 = ctk.CTkLabel(separ_frame, text='Эффективный КПД установки ηе', text_color="#bfbfbf")
label34.grid(column=5,row=22, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label35 = ctk.CTkLabel(separ_frame, text='Коэффициент полезной работы φ', text_color="#bfbfbf")
label35.grid(column=5,row=23, padx=(20,5), pady=5, sticky="w")
#==================================================================================================================================================
label200 = ctk.CTkLabel(separ_frame, text='Эффективная удельная работа He, кДж/кг', text_color="#bfbfbf")
label200.grid(column=5,row=21, padx=(20,5), pady=5, sticky="w")

# lb2 = ctk.CTkLabel(scrollable_frame, text="")
# lb2.grid(column=0, row=0, padx=5, pady=5)
#==================================================================================================================================================

def make_graph1():
    if widget.active:
        # Если виджет активен, вызываем toggle_widget() два раза
        toggle_widget()
        toggle_widget()
    else:
        # Если виджет не активен, строим график зависимости φ от πк
        plot_phi_vs_pk()

def make_graph2():
    if widget.active:
        # Если виджет активен, вызываем toggle_widget() два раза
        toggle_widget()
        toggle_widget()
    else:
        # Если виджет не активен, строим график зависимости φ от πк
        plot_eta_vs_pk()

def make_graph3():
    if widget.active:
        # Если виджет активен, вызываем toggle_widget() два раза
        toggle_widget()
        toggle_widget()
    else:
        # Если виджет не активен, строим график зависимости φ от πк
        plot_he_vs_pk()

def make_graph4():
    if widget.active:
        # Если виджет активен, вызываем toggle_widget() два раза
        toggle_widget()
        toggle_widget()
    else:
        # Если виджет не активен, строим график зависимости φ от πк
        plot_eta_he_vs_pk()
        
# Объявляем глобальные переменные для кнопок в начале кода
btn_plot_phi = None
btn_plot_eta = None
btn_plot_he = None
btn_plot_eta_he = None
savegraf2_button = None
savegraf3_button = None
savegraf4_button = None
savegraf5_button = None
graph3d = None
graph3d3 = None
graph3d2 = None
graph3d1 = None
savegraf_pdf_button4 = None
savegraf_pdf_button3 = None
savegraf_pdf_button2 = None
savegraf_pdf_button1 = None

# Функции для создания/удаления кнопок в settings_frame
def create_btn_plot_phi():
    global btn_plot_phi, savegraf2_button, graph3d, graph3d1, savegraf_pdf_button1
    if btn_plot_phi is not None:
        btn_plot_phi.destroy()
    if savegraf2_button is not None:
        savegraf2_button.destroy()
    if graph3d1 is not None:
        graph3d1.destroy()
    if savegraf_pdf_button1 is not None:
        savegraf_pdf_button1.destroy()
        
    btn_plot_phi = ctk.CTkButton(
        settings_frame,
        text="Построить 2D график", text_color="black",
        command=make_graph1,
        width=170,
        height=40,
        corner_radius=5, fg_color="#a2f200", image=func_icon, hover_color="#7fff00"
    )
    btn_plot_phi.grid(row=100, column=100, padx=10, pady=(0,0), sticky="se")
    
    savegraf2_button = ctk.CTkButton(settings_frame, text='', image=pdf_icon, width=170, height=40, fg_color="#f97666", hover_color="#fba297", corner_radius=5, command=save_plot_as_pdf1)
    savegraf2_button.grid(row=101, column=100, padx=10, pady=(5,10), sticky="se")

    # Создаём 3D-кнопку только если lang.get() == 'c регенерацией'
    if lang.get() == 'c регенерацией':
        graph3d1 = ctk.CTkButton(
            settings_frame, 
            text="Построить 3D график", 
            text_color="black", 
            width=170, 
            height=40, 
            corner_radius=5, 
            fg_color="#a2f200", 
            image=d_curve_icon, 
            hover_color="#7fff00",
            command=plot_phi_vs_pk_mu
        )
        graph3d1.grid(row=100, column=99, padx=10, pady=(0,0), sticky="se")

        savegraf_pdf_button1 = ctk.CTkButton(settings_frame, text='', image=pdf_icon, width=170, height=40, fg_color="#f97666", hover_color="#fba297", corner_radius=5, command=save_plot_as_pdf_3d1)
        savegraf_pdf_button1.grid(row=101, column=99, padx=10, pady=(5,10), sticky="se")

def create_btn_plot_eta():
    global btn_plot_eta, savegraf3_button, graph3d2, savegraf_pdf_button2
    if btn_plot_eta is not None:
        btn_plot_eta.destroy()
    if savegraf3_button is not None:
        savegraf3_button.destroy()
    if graph3d2 is not None:
        graph3d2.destroy()
    if savegraf_pdf_button2 is not None:
        savegraf_pdf_button2.destroy()
        
    btn_plot_eta = ctk.CTkButton(
        settings_frame,
        text="Построить 2D график", text_color="black",
        command=make_graph2,
        width=170,
        height=40,
        corner_radius=5, fg_color="#a2f200", image=func_icon, hover_color="#7fff00"
    )
    btn_plot_eta.grid(row=100, column=100, padx=10, pady=(0,0), sticky="se")
    
    savegraf3_button = ctk.CTkButton(settings_frame, text='', image=pdf_icon, width=170, height=40, fg_color="#f97666", hover_color="#fba297", corner_radius=5, command=save_plot_as_pdf2)
    savegraf3_button.grid(row=101, column=100, padx=10, pady=(5,10), sticky="se")

    # Создаём 3D-кнопку только если lang.get() == 'c регенерацией'
    if lang.get() == 'c регенерацией':
        graph3d2 = ctk.CTkButton(
            settings_frame, 
            text="Построить 3D график", 
            text_color="black", 
            width=170, 
            height=40, 
            corner_radius=5, 
            fg_color="#a2f200", 
            image=d_curve_icon, 
            hover_color="#7fff00",
            command=plot_eta_vs_pk_mu
        )
        graph3d2.grid(row=100, column=99, padx=10, pady=(0,0), sticky="se")

        savegraf_pdf_button2 = ctk.CTkButton(settings_frame, text='', image=pdf_icon, width=170, height=40, fg_color="#f97666", hover_color="#fba297", corner_radius=5, command=save_plot_as_pdf_3d2)
        savegraf_pdf_button2.grid(row=101, column=99, padx=10, pady=(5,10), sticky="se")

def create_btn_plot_he():
    global btn_plot_he, savegraf4_button, graph3d3, savegraf_pdf_button3
    if btn_plot_he is not None:
        btn_plot_he.destroy()
    if savegraf4_button is not None:
        savegraf4_button.destroy()
    if graph3d3 is not None:
        graph3d3.destroy()
    if savegraf_pdf_button3 is not None:
        savegraf_pdf_button3.destroy()
        
    btn_plot_he = ctk.CTkButton(
        settings_frame,
        text="Построить 2D график", text_color="black",
        command=make_graph3,
        width=170,
        height=40,
        corner_radius=5, fg_color="#a2f200", image=func_icon, hover_color="#7fff00"
    )
    btn_plot_he.grid(row=100, column=100, padx=10, pady=(0,0), sticky="se")
    
    savegraf4_button = ctk.CTkButton(settings_frame, text='', image=pdf_icon, width=170, height=40, fg_color="#f97666", hover_color="#fba297", corner_radius=5, command=save_plot_as_pdf)
    savegraf4_button.grid(row=101, column=100, padx=10, pady=(5,10), sticky="se")

    # Создаём 3D-кнопку только если lang.get() == 'c регенерацией'
    if lang.get() == 'c регенерацией':
        graph3d3 = ctk.CTkButton(
            settings_frame, 
            text="Построить 3D график", 
            text_color="black", 
            width=170, 
            height=40, 
            corner_radius=5, 
            fg_color="#a2f200", 
            image=d_curve_icon, 
            hover_color="#7fff00",
            command=plot_he_vs_pk_mu
        )
        graph3d3.grid(row=100, column=99, padx=10, pady=(0,0), sticky="se")

        savegraf_pdf_button3 = ctk.CTkButton(settings_frame, text='', image=pdf_icon, width=170, height=40, fg_color="#f97666", hover_color="#fba297", corner_radius=5, command=save_plot_as_pdf_3d3)
        savegraf_pdf_button3.grid(row=101, column=99, padx=10, pady=(5,10), sticky="se")

def create_btn_plot_eta_he():
    global btn_plot_eta_he, savegraf5_button, graph3d, savegraf_pdf_button4
    
    # Уничтожаем старые кнопки, если они существуют
    if btn_plot_eta_he is not None:
        btn_plot_eta_he.destroy()
    if savegraf5_button is not None:
        savegraf5_button.destroy()
    if graph3d is not None:
        graph3d.destroy()
    if savegraf_pdf_button4 is not None:
        savegraf_pdf_button4.destroy()
    
    # Создаём основные кнопки
    btn_plot_eta_he = ctk.CTkButton(
        settings_frame,
        text="Построить 2D график", text_color="black",
        command=make_graph4,
        width=170,
        height=40,
        corner_radius=5, fg_color="#a2f200", image=func_icon, hover_color="#7fff00"
    )
    btn_plot_eta_he.grid(row=100, column=100, padx=10, pady=(0,0), sticky="se")
    
    savegraf5_button = ctk.CTkButton(settings_frame, text='', image=pdf_icon, width=170, height=40, fg_color="#f97666", hover_color="#fba297", corner_radius=5, command=save_plot_as_pdf4)
    savegraf5_button.grid(row=101, column=100, padx=10, pady=(5,10), sticky="se")
    
    # Создаём 3D-кнопку только если lang.get() == 'c регенерацией'
    if lang.get() == 'c регенерацией':
        graph3d = ctk.CTkButton(
            settings_frame, 
            text="Построить 3D график", 
            text_color="black", 
            width=170, 
            height=40, 
            corner_radius=5, 
            fg_color="#a2f200", 
            image=d_curve_icon, 
            hover_color="#7fff00",
            command=plot_eta_vs_mu_he
        )
        graph3d.grid(row=100, column=99, padx=10, pady=(0,0), sticky="se")

        savegraf_pdf_button4 = ctk.CTkButton(settings_frame, text='', image=pdf_icon, width=170, height=40, fg_color="#f97666", hover_color="#fba297", corner_radius=5, command=save_plot_as_pdf_3d4)
        savegraf_pdf_button4.grid(row=101, column=99, padx=10, pady=(5,10), sticky="se")

def remove_all_plot_buttons():
    global btn_plot_phi, btn_plot_eta, btn_plot_he, btn_plot_eta_he, savegraf2_button, savegraf3_button, savegraf4_button, savegraf5_button, graph3d, graph3d3, graph3d2, graph3d1, savegraf_pdf_button4, savegraf_pdf_button3, savegraf_pdf_button2, savegraf_pdf_button1 
    for btn in [btn_plot_phi, btn_plot_eta, btn_plot_he, btn_plot_eta_he, savegraf2_button, savegraf3_button, savegraf4_button, savegraf5_button, graph3d, graph3d3, graph3d2, graph3d1, savegraf_pdf_button4, savegraf_pdf_button3, savegraf_pdf_button2, savegraf_pdf_button1]:
        if btn is not None:
            btn.destroy()
    btn_plot_phi = None
    btn_plot_eta = None
    btn_plot_he = None
    btn_plot_eta_he = None
    savegraf2_button = None
    savegraf3_button = None
    savegraf4_button = None
    savegraf5_button = None
    graph3d = None
    graph3d3 = None
    graph3d2 = None
    graph3d1 = None
    savegraf_pdf_button4 = None
    savegraf_pdf_button3 = None
    savegraf_pdf_button2 = None
    savegraf_pdf_button1 = None

def deactivate_other_buttons(active_button):
    buttons = [toggle_graph_frame_btn, toggle_graph_frame_btn2, 
               toggle_graph_frame_btn3, toggle_graph_frame_btn4]
    
    for btn in buttons:
        if btn != active_button:
            # Скрываем соответствующий фрейм
            if btn == toggle_graph_frame_btn:
                graph_frame.grid_remove()
                btn.configure(text="", image=graph_icon)
            elif btn == toggle_graph_frame_btn2:
                graph_frame2.grid_remove()
                btn.configure(text="", image=graph2_icon)
            elif btn == toggle_graph_frame_btn3:
                graph_frame3.grid_remove()
                btn.configure(text="", image=graph3_icon)
            elif btn == toggle_graph_frame_btn4:
                graph_frame4.grid_remove()
                btn.configure(text="", image=graph4_icon)

# Модифицированные функции переключения фреймов остаются без изменений
def toggle_graph_frame1(frame):
    current_theme = get_current_theme()
    if frame.winfo_ismapped():
        frame.grid_remove()
        toggle_graph_frame_btn.configure(text="", image=graph_icon)
        remove_all_plot_buttons()
    else:
        # Деактивируем другие кнопки
        deactivate_other_buttons(toggle_graph_frame_btn)
        # Обновляем цвет перед показом
        frame.configure(fg_color="#252525" if current_theme == "dark" else "#777a80")
        frame.grid()
        toggle_graph_frame_btn.configure(text="", image=graph1_1_icon)
        remove_all_plot_buttons()
        create_btn_plot_phi()

def toggle_graph_frame2(frame2):
    current_theme = get_current_theme()
    if frame2.winfo_ismapped():
        frame2.grid_remove()
        toggle_graph_frame_btn2.configure(text="", image=graph2_icon)
        remove_all_plot_buttons()
    else:
        # Деактивируем другие кнопки
        deactivate_other_buttons(toggle_graph_frame_btn2)
        # Обновляем цвет перед показом
        frame2.configure(fg_color="#252525" if current_theme == "dark" else "#777a80")
        frame2.grid()
        toggle_graph_frame_btn2.configure(text="", image=graph2_2_icon)
        remove_all_plot_buttons()
        create_btn_plot_eta()

def toggle_graph_frame3(frame3):
    current_theme = get_current_theme()
    if frame3.winfo_ismapped():
        frame3.grid_remove()
        toggle_graph_frame_btn3.configure(text="", image=graph3_icon)
        remove_all_plot_buttons()
    else:
        # Деактивируем другие кнопки
        deactivate_other_buttons(toggle_graph_frame_btn3)
        # Обновляем цвет перед показом
        frame3.configure(fg_color="#252525" if current_theme == "dark" else "#777a80")
        frame3.grid()
        toggle_graph_frame_btn3.configure(text="", image=graph3_3_icon)
        remove_all_plot_buttons()
        create_btn_plot_he()

def toggle_graph_frame4(frame4):
    current_theme = get_current_theme()
    if frame4.winfo_ismapped():
        frame4.grid_remove()
        toggle_graph_frame_btn4.configure(text="", image=graph4_icon)
        remove_all_plot_buttons()
    else:
        # Деактивируем другие кнопки
        deactivate_other_buttons(toggle_graph_frame_btn4)
        # Обновляем цвет перед показом
        frame4.configure(fg_color="#252525" if current_theme == "dark" else "#777a80")
        frame4.grid()
        toggle_graph_frame_btn4.configure(text="", image=graph4_4_icon)
        remove_all_plot_buttons()
        create_btn_plot_eta_he()
#========================================================================================================================================
#========================================================================================================================================
def cleargraf():
    entry_min_phi.delete(0, ctk.END)
    entry_max_phi.delete(0, ctk.END)
    entry_step_phi.delete(0, ctk.END)
    
    # Определяем целевой фрейм
    target_frame = subframes[0] if widget.active else display_frame
    create_empty_plot(target_frame)

def cleargraf2():
    entry_min_eta.delete(0, ctk.END)
    entry_max_eta.delete(0, ctk.END)
    entry_step_eta.delete(0, ctk.END)
    
    target_frame = subframes[1] if widget.active else display_frame
    create_empty_plot(target_frame)

def cleargraf3():
    entry_min_he.delete(0, ctk.END)
    entry_max_he.delete(0, ctk.END)
    entry_step_he.delete(0, ctk.END)
    
    target_frame = subframes[2] if widget.active else display_frame
    create_empty_plot(target_frame)

def cleargraf4():
    entry_min_eta_he.delete(0, ctk.END)
    entry_max_eta_he.delete(0, ctk.END)
    entry_step_eta_he.delete(0, ctk.END)
    
    target_frame = subframes[3] if widget.active else display_frame
    create_empty_plot(target_frame)

#========================================================================================================================================
# В начале кода объявляем все глобальные переменные для холстов
global canvas, canvas_eta, canvas_he, canvas_eta_he, canvas_e, canvas_3d, canvas_3d2, canvas_3d4, canvas_3d_he
canvas = None
canvas_eta = None
canvas_he = None
canvas_eta_he = None
canvas_e = None 
canvas_3d = None
canvas_3d2 = None
canvas_3d4 = None
canvas_3d_he = None
canvas_e3 = None

def destroy_all_canvases():
    """Уничтожает все существующие холсты графиков"""
    global canvas, canvas_eta, canvas_he, canvas_eta_he, canvas_e, canvas_3d, canvas_3d2, canvas_3d4, canvas_3d_he, canvas_e3
    
    for canvas_item in [canvas, canvas_eta, canvas_he, canvas_eta_he, canvas_e, canvas_3d, canvas_3d2, canvas_3d4, canvas_3d_he, canvas_e3]:
        if canvas_item is not None:
            try:
                # Получаем Tkinter-виджет и проверяем его существование
                widget = canvas_item.get_tk_widget()
                if widget.winfo_exists():
                    widget.destroy()
            except Exception as e:
                print(f"Ошибка при уничтожении холста: {e}")
            canvas_item = None

def get_current_theme():
    """
    Возвращает текущую тему программы ("dark" или "light").
    """
    return ctk.get_appearance_mode().lower()  # Возвращает "dark" или "light"

global global_save_fig
def plot_phi_vs_pk(master=None):
    global canvas, global_save_fig  # Используем глобальную переменную для холста

    try:
        # Если указан master, уничтожаем графики только в нем
        if master:
            for child in master.winfo_children():
                child.destroy()
        else:
            destroy_all_canvases()  # Уничтожаем все предыдущие графики

        # Проверка заполнения полей для πк
        min_pk_text = entry_min_phi.get().strip()
        max_pk_text = entry_max_phi.get().strip()
        step_pk_text = entry_step_phi.get().strip()
        
        # Если не введены значения πк, строим пустой график
        if not min_pk_text or not max_pk_text or not step_pk_text:
            create_empty_plot(master if master else None)
            
            return
        
        # Проверка заполнения всех необходимых полей
        required_fields = [
            inp1, inp2, inp3, inp5, inp6, inp7, inp8, inp9,
            inp10, inp11, inp12, inp13, inp15, inp16,
            inp17, inp18, inp100
        ]

        # Если выбран режим с регенерацией, добавляем inp14 (степень рекуперации) в обязательные поля
        if lang.get() == 'c регенерацией':
            required_fields.append(inp14)

        for field in required_fields:
            value = field.get().strip()
            if not value:
                create_empty_plot(master if master else None)
                messagebox.showwarning(
                    "Ошибка ввода",
                    "Не все обязательные поля заполнены"
                )
                raise ValueError("Не все обязательные поля заполнены")

        # Получение диапазона πк (без дефолтных значений)
        try:
            min_pk = float(replace_comma_with_dot(entry_min_phi.get()))
            max_pk = float(replace_comma_with_dot(entry_max_phi.get()))
            step_pk = float(replace_comma_with_dot(entry_step_phi.get())) 
        except ValueError:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк должны быть числовыми"
            )
            return

        if min_pk <= 0 or max_pk <= 0 or step_pk <= 0:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк и шаг должны быть положительными"
            )
            raise ValueError("Значения πк и шаг должны быть положительными")
            
        if min_pk >= max_pk:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода", 
                "Минимальное πк должно быть меньше максимального"
            )
            raise ValueError("Минимальное πк должно быть меньше максимального")
        
        # Проверка, что шаг не слишком большой
        if step_pk > (max_pk - min_pk):
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Шаг слишком большой для заданного диапазона πк"
            )
            return

        # Генерация массива πк
        pk_values = np.arange(min_pk, max_pk + step_pk, step_pk)

        # Сохранение исходного значения πк
        original_pk = inp6.get()

        phi_values = []
        for pk in pk_values:
            # Установка текущего πк
            inp6.delete(0, ctk.END)
            inp6.insert(0, str(pk))

            # Вызов расчета в зависимости от режима
            if lang.get() == 'без регенерации':
                raschet()  # Расчет для режима без регенерации
            elif lang.get() == 'c регенерацией':
                raschet()  # Расчет для режима с регенерации
            else:
                raise ValueError("Режим расчета не выбран")

            # Получение φ
            phi_text = label35.cget("text")
            if 'φ=' in phi_text:
                phi_str = phi_text.split('=')[1].strip()
                phi = float(phi_str)
                phi_values.append(phi)
            else:
                phi_values.append(0.0)

        # Интерполяция для плавного графика
        x_new = np.linspace(min_pk, max_pk, 300)  # 300 точек для плавного графика
        spl = make_interp_spline(pk_values, phi_values, k=3)  # Кубическая интерполяция
        y_new = spl(x_new)


        # Восстановление исходного πк
        inp6.delete(0, ctk.END)
        inp6.insert(0, original_pk)
        raschet()

        # Создание ДВУХ фигур: для отображения и сохранения
        # Фигура для отображения (с учетом текущей темы)
        fig_display = plt.Figure(figsize=(4, 3), dpi=100)
        ax_display = fig_display.add_subplot(111)

        # Фигура для сохранения (всегда светлая)
        fig_save = plt.Figure(figsize=(6, 4), dpi=100)
        ax_save = fig_save.add_subplot(111)

        # Настройка стиля графика для отображения
        current_theme = get_current_theme()
        if current_theme == "dark":
            plt.style.use("dark_background")
            text_color = "white"
            line_color = "#00A0DC"
            bg_color = "#151515"
        else:
            plt.style.use("default")
            text_color = "black"
            line_color = "#00A0DC"
            bg_color = "#ffffff"

        # Настройка стиля для сохранения (всегда светлый)
        save_text_color = "black"
        save_line_color = "#00A0DC"
        save_bg_color = "#ffffff"

        # Общая функция для настройки графиков
        def configure_plot(ax, text_color, line_color, bg_color, is_save=False):
            ax.plot(x_new, y_new, linestyle='-', color=line_color)
            ax.scatter(pk_values, phi_values, color='#00A0DC', label='Значения πк', marker='o')
            
            ax.set_xlabel('Степень повышения давления (πк)', color=text_color)
            ax.set_ylabel('Коэффициент полезной работы (φ)', color=text_color)
            ax.set_title(f'Зависимость φ от πк при T3 = {inp3.get()} K ({lang.get()})', color=text_color)
            ax.legend()
            
            # Настройка сетки
            if step_pk < 1:
                x_ticks = np.arange(min_pk, max_pk + 2 * step_pk, 2 * step_pk)
            else:
                x_ticks = np.arange(min_pk, max_pk + step_pk, step_pk)
            ax.set_xticks(x_ticks)
            
            ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
            ax.grid(True, color="gray" if text_color == "white" else "black")
            ax.set_facecolor(bg_color)
            
            if is_save:
                ax.figure.patch.set_facecolor(save_bg_color)
            else:
                ax.figure.patch.set_facecolor(bg_color)
            
            for spine in ax.spines.values():
                spine.set_edgecolor(text_color)
            ax.tick_params(axis='x', colors=text_color)
            ax.tick_params(axis='y', colors=text_color)

        # Настройка графика для отображения
        configure_plot(ax_display, text_color, line_color, bg_color)

        # Настройка графика для сохранения
        plt.style.use("default")  # Принудительно светлый стиль
        configure_plot(ax_save, save_text_color, save_line_color, save_bg_color, is_save=True)
        global_save_fig = fig_save  # Сохраняем светлую фигуру

        # Удаление старого графика, если он существует
        if canvas:
            canvas.get_tk_widget().destroy()


        # Встраивание нового графика в указанный master или в display_frame
        if master:
            canvas = FigureCanvasTkAgg(fig_display, master=master)
        else:
            canvas = FigureCanvasTkAgg(fig_display, master=display_frame)
            if canvas_eta:
                canvas_eta.get_tk_widget().destroy()
                if canvas_he:
                    canvas_he.get_tk_widget().destroy()
                    if canvas_eta_he:
                        canvas_eta_he.get_tk_widget().destroy()
                        if canvas_e:
                            canvas_e.get_tk_widget().destroy()
            
        canvas.draw()
        canvas.get_tk_widget().pack(fill='both', expand=True, padx=5, pady=5)

    except ValueError as ve:
        print(f"Ошибка: {ve}")
        if master: 
            create_empty_plot_in_frame(master)
    except Exception as e:
        print(f"Неожиданная ошибка: {e}")
        if master:
            create_empty_plot_in_frame(master)

#==================================================================================================================================================
# Функция для построения графика зависимости ηe от πк
global saved_fig_eta
def plot_eta_vs_pk(master=None):
    global canvas_eta  # Глобальная переменная для холста
    global saved_fig_eta  # Глобальная переменная для сохранения графика в светлом стиле

    try:
        # Уничтожаем графики в переданном master, если он есть
        if master:
            for child in master.winfo_children():
                child.destroy()
        else:
            destroy_all_canvases()

        # Проверка заполнения полей для πк
        min_pk_text = entry_min_eta.get().strip()
        max_pk_text = entry_max_eta.get().strip()
        step_pk_text = entry_step_eta.get().strip()
        
        # Если не введены значения πк, строим пустой график
        if not min_pk_text or not max_pk_text or not step_pk_text:
            create_empty_plot(master if master else None)
            
            return
        
        # Проверка заполнения всех необходимых полей
        required_fields = [
            inp1, inp2, inp3, inp5, inp6, inp7, inp8, inp9,
            inp10, inp11, inp12, inp13, inp15, inp16,
            inp17, inp18, inp100
        ]

        # Если выбран режим с регенерацией, добавляем inp14 (степень рекуперации) в обязательные поля
        if lang.get() == 'c регенерацией':
            required_fields.append(inp14)

        for field in required_fields:
            value = field.get().strip()
            if not value:
                create_empty_plot(master if master else None)
                messagebox.showwarning(
                    "Ошибка ввода",
                    "Не все обязательные поля заполнены"
                )
                raise ValueError("Не все обязательные поля заполнены")

        # Получение диапазона πк (без дефолтных значений)
        try:
            min_pk = float(replace_comma_with_dot(entry_min_eta.get()))
            max_pk = float(replace_comma_with_dot(entry_max_eta.get()))
            step_pk = float(replace_comma_with_dot(entry_step_eta.get())) 
        except ValueError:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк должны быть числовыми"
            )
            return

        if min_pk <= 0 or max_pk <= 0 or step_pk <= 0:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк и шаг должны быть положительными"
            )
            raise ValueError("Значения πк и шаг должны быть положительными")
            
        if min_pk >= max_pk:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода", 
                "Минимальное πк должно быть меньше максимального"
            )
            raise ValueError("Минимальное πк должно быть меньше максимального")
        
        # Проверка, что шаг не слишком большой
        if step_pk > (max_pk - min_pk):
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Шаг слишком большой для заданного диапазона πк"
            )
            return

        # Генерация массива πк
        pk_values = np.arange(min_pk, max_pk + step_pk, step_pk)

        # Сохранение исходного значения πк
        original_pk = inp6.get()

        eta_values = []
        for pk in pk_values:
            # Установка текущего πк
            inp6.delete(0, ctk.END)
            inp6.insert(0, str(pk))

            # Вызов расчета в зависимости от режима
            if lang.get() == 'без регенерации':
                raschet()  # Расчет для режима без регенерации
            elif lang.get() == 'c регенерацией':
                raschet()  # Расчет для режима с регенерации
            else:
                raise ValueError("Режим расчета не выбран")

            # Получение ηe
            eta_text = label34.cget("text")
            if 'ηе=' in eta_text:
                eta_str = eta_text.split('=')[1].strip()
                eta = float(eta_str)
                eta_values.append(eta)
            else:
                eta_values.append(0.0)

        # Интерполяция для плавного графика
        x_new = np.linspace(min_pk, max_pk, 300)
        spl = make_interp_spline(pk_values, eta_values, k=3)
        y_new = spl(x_new)

        # Нахождение максимального значения ηe
        max_eta = np.max(y_new)
        max_eta_index = np.argmax(y_new)
        max_pk_value = x_new[max_eta_index]

        # Восстановление исходного πк
        inp6.delete(0, ctk.END)
        inp6.insert(0, original_pk)
        raschet()

        # Создание графиков
        fig_display = plt.Figure(figsize=(4, 3), dpi=100)
        ax_display = fig_display.add_subplot(111)

        fig_save = plt.Figure(figsize=(6, 4), dpi=100)
        ax_save = fig_save.add_subplot(111)

        # Настройка стилей
        current_theme = get_current_theme()
        if current_theme == "dark":
            plt.style.use("dark_background")  # Темный стиль
            text_color = "white"  # Белый цвет для текста
            line_color = "#00A0DC"  # Синий цвет для линии графика
            bg_color = "#151515"  # Цвет фона для темной темы
        else:
            plt.style.use("default")  # Светлый стиль
            text_color = "black"  # Черный цвет для текста
            line_color = "#00A0DC"  # Синий цвет для линии графика
            bg_color = "#ffffff"  # Цвет фона для светлой темы

        # Настройка стиля для сохранения (всегда светлый)
        save_text_color = "black"
        save_line_color = "#00A0DC"
        save_bg_color = "#ffffff"

        # Функция настройки графиков
        def configure_plot(ax, text_color, line_color, bg_color, is_save=False):
            ax.plot(x_new, y_new, linestyle='-', color=line_color)
            ax.scatter(pk_values, eta_values, color='#00A0DC', label='Значения πк', marker='o')
            ax.scatter(max_pk_value, max_eta, color='#c22b30', 
                    label=f'Максимум: ηe={max_eta:.3f} при πк={max_pk_value:.2f}')
    
            ax.set_xlabel('Степень повышения давления (πк)', color=text_color)
            ax.set_ylabel('Эффективный КПД (ηe)', color=text_color)
            ax.set_title(f'Зависимость ηe от πк при T3 = {inp3.get()} K ({lang.get()})', 
                    color=text_color)
    
            # Измененная секция легенды
            ax.legend()

            # Настройка сетки
            if step_pk < 1:
                x_ticks = np.arange(min_pk, max_pk + 2 * step_pk, 2 * step_pk)
            else:
                x_ticks = np.arange(min_pk, max_pk + step_pk, step_pk)
            ax.set_xticks(x_ticks)
            
            ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True)) 
            ax.grid(True, color="gray" if text_color == "white" else "black")
            ax.set_facecolor(bg_color)
            if is_save:
                ax.figure.patch.set_facecolor(save_bg_color)
            else:
                ax.figure.patch.set_facecolor(bg_color)
            # Настройка цвета осей и меток
            ax.tick_params(axis='x', colors=text_color)
            ax.tick_params(axis='y', colors=text_color)

            for spine in ax.spines.values():
                spine.set_edgecolor(text_color)
            

        # Настройка графика для отображения
        configure_plot(ax_display, text_color, line_color, bg_color)

        # Настройка графика для сохранения (всегда светлый)
        plt.style.use("default")  # Принудительно устанавливаем светлый стиль
        configure_plot(ax_save, save_text_color, save_line_color, save_bg_color, is_save=True)
        saved_fig_eta = fig_save  # Сохраняем график для экспорта

        # Удаление старого графика
        if canvas_eta:
            canvas_eta.get_tk_widget().destroy()

        # Встраивание графика
        if master:
            canvas_eta = FigureCanvasTkAgg(fig_display, master=master)
        else:
            canvas_eta = FigureCanvasTkAgg(fig_display, master=display_frame)
            # Очистка предыдущих графиков
            if canvas:
                canvas.get_tk_widget().destroy()
                if canvas_he:
                    canvas_he.get_tk_widget().destroy()
                    if canvas_eta_he:
                        canvas_eta_he.get_tk_widget().destroy()
                        if canvas_e:
                            canvas_e.get_tk_widget().destroy()

        canvas_eta.draw()
        canvas_eta.get_tk_widget().pack(fill='both', expand=True, padx=5, pady=5)
    
    except Exception as e:
        if master:
            create_empty_plot_in_frame(master)
        raise e

import time  # Импортируем модуль time для задержки
#==================================================================================================================================================
global saved_fig_he
def plot_he_vs_pk(master=None):
    global canvas_he  # Используем глобальную переменную для холста
    global saved_fig_he  # Глобальная переменная для сохранения графика в светлом стиле

    try:
        # Уничтожаем графики в переданном master, если он есть
        if master:
            for child in master.winfo_children():
                child.destroy()
        else:
            destroy_all_canvases()

        # Проверка заполнения полей для πк
        min_pk_text = entry_min_he.get().strip()
        max_pk_text = entry_max_he.get().strip()
        step_pk_text = entry_step_he.get().strip()
        
        # Если не введены значения πк, строим пустой график
        if not min_pk_text or not max_pk_text or not step_pk_text:
            create_empty_plot(master if master else None)
            
            return
        
        # Проверка заполнения всех необходимых полей
        required_fields = [
            inp1, inp2, inp3, inp5, inp6, inp7, inp8, inp9,
            inp10, inp11, inp12, inp13, inp15, inp16,
            inp17, inp18, inp100
        ]

        # Если выбран режим с регенерацией, добавляем inp14 (степень рекуперации) в обязательные поля
        if lang.get() == 'c регенерацией':
            required_fields.append(inp14)

        for field in required_fields:
            value = field.get().strip()
            if not value:
                create_empty_plot(master if master else None)
                messagebox.showwarning(
                    "Ошибка ввода",
                    "Не все обязательные поля заполнены"
                )
                raise ValueError("Не все обязательные поля заполнены")

        # Получение диапазона πк (без дефолтных значений)
        try:
            min_pk = float(replace_comma_with_dot(entry_min_he.get()))
            max_pk = float(replace_comma_with_dot(entry_max_he.get()))
            step_pk = float(replace_comma_with_dot(entry_step_he.get())) 
        except ValueError:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк должны быть числовыми"
            )
            return

        if min_pk <= 0 or max_pk <= 0 or step_pk <= 0:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк и шаг должны быть положительными"
            )
            raise ValueError("Значения πк и шаг должны быть положительными")
            
        if min_pk >= max_pk:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода", 
                "Минимальное πк должно быть меньше максимального"
            )
            raise ValueError("Минимальное πк должно быть меньше максимального")
        
        # Проверка, что шаг не слишком большой
        if step_pk > (max_pk - min_pk):
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Шаг слишком большой для заданного диапазона πк"
            )
            return

        # Генерация массива πк
        pk_values = np.arange(min_pk, max_pk + step_pk, step_pk)

        # Сохранение исходного значения πк
        original_pk = inp6.get()

        he_values = []
        for pk in pk_values:
            # Установка текущего πк
            inp6.delete(0, ctk.END)
            inp6.insert(0, str(pk))

            # Вызов расчета в зависимости от режима
            if lang.get() == 'без регенерации':
                raschet()  # Расчет для режима без регенерации
            elif lang.get() == 'c регенерацией':
                raschet()  # Расчет для режима с регенерации
            else:
                raise ValueError("Режим расчета не выбран")

            # Получение He
            he_text = label200.cget("text")
            if 'He=' in he_text:
                he_str = he_text.split('=')[1].strip().split()[0]  # Берем только число перед "кДж/кг"
                he = float(he_str.replace(',', '.')) 
                he_values.append(he)
            else:
                he_values.append(0.0)

        # Интерполяция для плавного графика
        x_new = np.linspace(min_pk, max_pk, 300)  # 300 точек для плавного графика
        spl = make_interp_spline(pk_values, he_values, k=3)  # Кубическая интерполяция
        y_new = spl(x_new)

        # Нахождение максимального значения He и соответствующего πк
        max_he = np.max(y_new)  # Максимальное значение ηe
        max_he_index = np.argmax(y_new)  # Индекс максимального значения
        max_pk_value = x_new[max_he_index]  # Соответствующее значение πк

        # Восстановление исходного πк
        inp6.delete(0, ctk.END)
        inp6.insert(0, original_pk)
        raschet()

        # Создание графика для отображения (с учетом текущей темы)
        fig_display = plt.Figure(figsize=(4, 3), dpi=100)
        ax_display = fig_display.add_subplot(111)

        # Создание графика для сохранения (всегда в светлом стиле)
        fig_save = plt.Figure(figsize=(6, 4), dpi=100)
        ax_save = fig_save.add_subplot(111)

        # Настройка стиля графика для отображения в зависимости от темы
        current_theme = get_current_theme()
        if current_theme == "dark":
            plt.style.use("dark_background")  # Темный стиль
            text_color = "white"  # Белый цвет для текста
            line_color = "#00A0DC"  # Синий цвет для линии графика
            bg_color = "#151515"  # Цвет фона для темной темы
        else:
            plt.style.use("default")  # Светлый стиль
            text_color = "black"  # Черный цвет для текста
            line_color = "#00A0DC"  # Синий цвет для линии графика
            bg_color = "#ffffff"  # Цвет фона для светлой темы

        # Настройка стиля для сохранения (всегда светлый)
        save_text_color = "black"
        save_line_color = "#00A0DC"
        save_bg_color = "#ffffff"

        # Функция для настройки графика
        def configure_plot(ax, text_color, line_color, bg_color, is_save=False):
            ax.plot(x_new, y_new, linestyle='-', color=line_color)
            
            # Отображение точек, соответствующих выбранному шагу πк
            ax.scatter(pk_values, he_values, color='#00A0DC', label='Значения πк', marker='o')

            # Добавление красной точки для максимального значения ηe
            label = f'Максимум: He={max_he:.3f} кДж/кг при πк={max_pk_value:.2f}'
            ax.scatter(max_pk_value, max_he, color='#c22b30', label=label)

            ax.set_xlabel('Степень повышения давления (πк)', color=text_color)
            ax.set_ylabel('Эффективная удельная работа (He, кДж/кг)', color=text_color)
            ax.set_title(f'Зависимость He от πк при T3 = {inp3.get()} K ({lang.get()})', color=text_color)
            ax.legend()
            
            # Настройка шага сетки по оси X
            if step_pk < 1:
                x_ticks = np.arange(min_pk, max_pk + 2 * step_pk, 2 * step_pk)
            else:
                x_ticks = np.arange(min_pk, max_pk + step_pk, step_pk)
            ax.set_xticks(x_ticks)

            # Настройка шага сетки по оси Y (автоматически)
            ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))

            ax.grid(True, color="gray" if text_color == "white" else "black")
            ax.set_facecolor(bg_color)
            if is_save:
                ax.figure.patch.set_facecolor(save_bg_color)
            else:
                ax.figure.patch.set_facecolor(bg_color)

            # Настройка цвета осей и меток
            ax.tick_params(axis='x', colors=text_color)
            ax.tick_params(axis='y', colors=text_color)

            # Настройка цвета рамки (границ) осей
            for spine in ax.spines.values():
                spine.set_edgecolor(text_color)

        # Настройка графика для отображения
        configure_plot(ax_display, text_color, line_color, bg_color)

        # Настройка графика для сохранения (всегда светлый)
        plt.style.use("default")  # Принудительно устанавливаем светлый стиль
        configure_plot(ax_save, save_text_color, save_line_color, save_bg_color, is_save=True)
        saved_fig_he = fig_save  # Сохраняем график в светлом стиле

        # Удаление старого графика, если он существует

        if canvas_he:
            canvas_he.get_tk_widget().destroy()

        # Встраивание графика в указанный master или в display_frame
        if master:
            canvas_he = FigureCanvasTkAgg(fig_display, master=master)
        else:
            canvas_he = FigureCanvasTkAgg(fig_display, master=display_frame)
            if canvas:
                canvas.get_tk_widget().destroy()
                if canvas_eta:
                    canvas_eta.get_tk_widget().destroy()
                    if canvas_eta_he:
                        canvas_eta_he.get_tk_widget().destroy()
                        if canvas_e:
                            canvas_e.get_tk_widget().destroy()
        
        canvas_he.draw()
        canvas_he.get_tk_widget().pack(fill='both', expand=True, padx=5, pady=5)
    
    except Exception as e:
        if master:
            create_empty_plot_in_frame(master)
        raise e

#==================================================================================================================================================
global saved_fig_eta_he
def plot_eta_he_vs_pk(master=None):
    global canvas_eta_he  # Используем глобальную переменную для холста
    global saved_fig_eta_he  # Глобальная переменная для сохранения графика в светлом стиле

    try:
        # Уничтожаем графики в переданном master, если он есть
        if master:
            for child in master.winfo_children():
                child.destroy()
        else:
            destroy_all_canvases()

        # Проверка заполнения полей для πк
        min_pk_text = entry_min_eta_he.get().strip()
        max_pk_text = entry_max_eta_he.get().strip()
        step_pk_text = entry_step_eta_he.get().strip()
        
        # Если не введены значения πк, строим пустой график
        if not min_pk_text or not max_pk_text or not step_pk_text:
            create_empty_plot(master if master else None)
            
            return
            
        # Проверка заполнения всех необходимых полей
        required_fields = [
            inp1, inp2, inp3, inp5, inp6, inp7, inp8, inp9,
            inp10, inp11, inp12, inp13, inp15, inp16,
            inp17, inp18, inp100
        ]

        # Если выбран режим с регенерацией, добавляем inp14 (степень рекуперации) в обязательные поля
        if lang.get() == 'c регенерацией':
            required_fields.append(inp14)

        for field in required_fields:
            value = field.get().strip()
            if not value:
                create_empty_plot(master if master else None)
                messagebox.showwarning(
                    "Ошибка ввода",
                    "Не все обязательные поля заполнены"
                )
                raise ValueError("Не все обязательные поля заполнены")

        # Получение диапазона πк (без дефолтных значений)
        try:
            min_pk = float(replace_comma_with_dot(entry_min_eta_he.get()))
            max_pk = float(replace_comma_with_dot(entry_max_eta_he.get()))
            step_pk = float(replace_comma_with_dot(entry_step_eta_he.get())) 
        except ValueError:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк должны быть числовыми"
            )
            return

        if min_pk <= 0 or max_pk <= 0 or step_pk <= 0:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк и шаг должны быть положительными"
            )
            raise ValueError("Значения πк и шаг должны быть положительными")
            
        if min_pk >= max_pk:
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода", 
                "Минимальное πк должно быть меньше максимального"
            )
            raise ValueError("Минимальное πк должно быть меньше максимального")
        
        # Проверка, что шаг не слишком большой
        if step_pk > (max_pk - min_pk):
            create_empty_plot(master if master else None)
            messagebox.showwarning(
                "Ошибка ввода",
                "Шаг слишком большой для заданного диапазона πк"
            )
            return

        # Генерация массива πк с шагом 1/4 от исходного шага для интерполяции
        interp_step = step_pk / 4  # Шаг интерполяции = 1/4 от шага πк
        pk_values_interp = np.arange(min_pk, max_pk + interp_step, interp_step)
        pk_values = np.arange(min_pk, max_pk + step_pk, step_pk)  # Оригинальные значения для точек

        # Сохранение исходного значения πк
        original_pk = inp6.get()

        he_values = []
        eta_values = []
        he_values_interp = []
        eta_values_interp = []
        
        # Расчет значений для интерполяции (с малым шагом)
        for pk in pk_values_interp:
            # Установка текущего πк
            inp6.delete(0, ctk.END)
            inp6.insert(0, str(pk))

            # Вызов расчета в зависимости от режима
            if lang.get() == 'без регенерации':
                raschet()  # Расчет для режима без регенерации
            elif lang.get() == 'c регенерацией':
                raschet()  # Расчет для режима с регенерации
            else:
                raise ValueError("Режим расчета не выбран")

            # Получение He
            he_text = label200.cget("text")
            if 'He=' in he_text:
                he_str = he_text.split('=')[1].strip().split()[0]  # Берем только число перед "кДж/кг"
                he = float(he_str.replace(',', '.')) 
                he_values_interp.append(he)
            else:
                he_values_interp.append(0.0)

            # Получение ηe
            eta_text = label34.cget("text")
            if 'ηе=' in eta_text:
                eta_str = eta_text.split('=')[1].strip()
                eta = float(eta_str)
                eta_values_interp.append(eta)
            else:
                eta_values_interp.append(0.0)

        # Расчет значений для отображения точек (с оригинальным шагом)
        for pk in pk_values:
            # Установка текущего πк
            inp6.delete(0, ctk.END)
            inp6.insert(0, str(pk))

            # Вызов расчета в зависимости от режима
            if lang.get() == 'без регенерации':
                raschet()  # Расчет для режима без регенерации
            elif lang.get() == 'c регенерацией':
                raschet()  # Расчет для режима с регенерации
            else:
                raise ValueError("Режим расчета не выбран")

            # Получение He
            he_text = label200.cget("text")
            if 'He=' in he_text:
                he_str = he_text.split('=')[1].strip().split()[0]
                he = float(he_str.replace(',', '.')) 
                he_values.append(he)
            else:
                he_values.append(0.0)

            # Получение ηe
            eta_text = label34.cget("text")
            if 'ηе=' in eta_text:
                eta_str = eta_text.split('=')[1].strip()
                eta = float(eta_str)
                eta_values.append(eta)
            else:
                eta_values.append(0.0)

        # Восстановление исходного πк
        inp6.delete(0, ctk.END)
        inp6.insert(0, original_pk)
        raschet()

        # Создание графика для отображения (с учетом текущей темы)
        fig_display = plt.Figure(figsize=(4, 3), dpi=100)
        ax_display = fig_display.add_subplot(111)

        # Создание графика для сохранения (всегда в светлом стиле)
        fig_eta_he_save = plt.Figure(figsize=(6, 4), dpi=100)
        ax_save = fig_eta_he_save.add_subplot(111)

        # Настройка стиля графика для отображения в зависимости от темы
        current_theme = get_current_theme()
        if current_theme == "dark":
            plt.style.use("dark_background")  # Темный стиль
            text_color = "white"  # Белый цвет для текста
            line_color = "#00A0DC"  # Синий цвет для линии графика
            bg_color = "#151515"  # Цвет фона для темной темы
        else:
            plt.style.use("default")  # Светлый стиль
            text_color = "black"  # Черный цвет для текста
            line_color = "#00A0DC"  # Синий цвет для линии графика
            bg_color = "#ffffff"  # Цвет фона для светлой темы

        # Настройка стиля для сохранения (всегда светлый)
        save_text_color = "black"
        save_line_color = "#00A0DC"
        save_bg_color = "#ffffff"

        # Функция для настройки графика
        def configure_plot(ax, text_color, line_color, bg_color, is_save=False):
            # Строим плавный график ηe от He с интерполяцией
            ax.plot(he_values_interp, eta_values_interp, linestyle='-', color=line_color)
            
            # Отображение точек с подписями πк (только первая и последняя)
            for i, (he, eta) in enumerate(zip(he_values, eta_values)):
                ax.scatter(he, eta, color='#00A0DC', marker='o')
    
                if i == 0 or i == len(he_values) - 1:
                    if i == 0:  # Первая точка
                        text_offset_x = 0.05 * (max(he_values_interp) - min(he_values_interp))  # Сдвиг вправо
                        text_offset_y = 0.00 * (max(eta_values_interp) - min(eta_values_interp))  # Сдвиг вверх
                        ha = 'left'  # Текст слева от точки (но сдвинут вправо)
                    else:  # Последняя точка
                        text_offset_x = -0.02 * (max(he_values_interp) - min(he_values_interp))  # Сдвиг влево
                        text_offset_y = -0.02 * (max(eta_values_interp) - min(eta_values_interp))  # Сдвиг вверх
                        ha = 'right'  # Текст справа от точки (но сдвинут влево)
                    ax.text(he + text_offset_x, eta + text_offset_y, f'πк={pk_values[i]:.2f}', 
                        color=text_color, 
                        fontsize=8,
                        ha=ha,
                        va='bottom')

            ax.set_xlabel('Эффективная удельная работа (He, кДж/кг)', color=text_color)
            ax.set_ylabel('Эффективный КПД (ηe)', color=text_color)
            ax.set_title(f'Зависимость ηe от He при T3 = {inp3.get()} K ({lang.get()})', color=text_color)
            
            # Настройка шага сетки
            ax.grid(True, color="gray" if text_color == "white" else "black")
            ax.set_facecolor(bg_color)
            if is_save:
                ax.figure.patch.set_facecolor(save_bg_color)
            else:
                ax.figure.patch.set_facecolor(bg_color)

            # Настройка цвета осей и меток
            ax.tick_params(axis='x', colors=text_color)
            ax.tick_params(axis='y', colors=text_color)

            # Настройка цвета рамки (границ) осей
            for spine in ax.spines.values():
                spine.set_edgecolor(text_color)

        # Настройка графика для отображения
        configure_plot(ax_display, text_color, line_color, bg_color)

        # Настройка графика для сохранения (всегда светлый)
        plt.style.use("default")  # Принудительно устанавливаем светлый стиль
        configure_plot(ax_save, save_text_color, save_line_color, save_bg_color, is_save=True)
        saved_fig_eta_he = fig_eta_he_save  # Сохраняем график в светлом стиле

        # Удаление старого графика, если он существует
        if canvas_eta_he:
            canvas_eta_he.get_tk_widget().destroy()
                        
        # Встраивание графика в указанный master или в display_frame
        if master:
            canvas_eta_he = FigureCanvasTkAgg(fig_display, master=master)
        else:
            canvas_eta_he = FigureCanvasTkAgg(fig_display, master=display_frame)
            # Удаление старого графика, если он существует
            if canvas:
                canvas.get_tk_widget().destroy()
                if canvas_eta:
                    canvas_eta.get_tk_widget().destroy()
                    if canvas_he:
                        canvas_he.get_tk_widget().destroy()
                        if canvas_e:
                            canvas_e.get_tk_widget().destroy()
        
        canvas_eta_he.draw()
        canvas_eta_he.get_tk_widget().pack(fill='both', expand=True, padx=5, pady=5)
    
    except Exception as e:
        if master:
            create_empty_plot_in_frame(master)
        raise e

        

def get_pk_from_label():
    try:
        # Получаем текст из label8
        label_text = label8.cget("text")
        # Извлекаем числовое значение после "πк="
        pk_str = label_text.split('=')[1].strip()
        # Заменяем запятую на точку, если есть
        pk = float(pk_str.replace(',', '.'))
        return pk
    except Exception as e:
        print(f"Ошибка при получении πк из label8: {e}")
        return None
#========================================================================================================================================
bbold_font = ctk.CTkFont(size=20, weight='bold')

# Основной контейнер для графиков
right_content_frame1.grid_columnconfigure(0, weight=1)  # единственная колонка расширяется
right_content_frame1.grid_rowconfigure(0, weight=1)     # display_frame займёт всё доступное пространство
right_content_frame1.grid_rowconfigure(1, weight=0)     # settings_frame будет прижат к низу

# Фрейм для вывода графика (верхний)
display_frame = ctk.CTkFrame(right_content_frame1, fg_color="#151515", height=500, corner_radius=15, border_width=1, border_color="#454545")
display_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)



# Фрейм настройки графика (нижний, скрывается/показывается)
settings_frame = ctk.CTkFrame(right_content_frame1, fg_color="#151515", corner_radius=15, border_width=1, border_color="#454545", height=300)
settings_frame.grid(row=1, column=0, sticky="ew", padx=0, pady=(10, 0))  # sticky="ew" — растягивается по ширине, но не по высоте

subframes = []
widget = None

def toggle_widget():
    global widget, subframes  # Добавляем subframes в глобальные переменные
    
    # Переключаем состояние
    widget.active = not widget.active
    
    if widget.active:
        # Активация - создаем 4 фрейма (2x2)
        widget.configure(image=widget_active_icon)
        
        # Очищаем display_frame перед созданием новых фреймов
        for child in display_frame.winfo_children():
            child.destroy()
        
        # Создаем 4 подфрейма в сетке 2x2
        subframes = []
        for i in range(2):
            for j in range(2):
                subframe = ctk.CTkFrame(
                    display_frame,
                    fg_color="#151515" if ctk.get_appearance_mode() == "Dark" else "white",
                    corner_radius=10,
                    border_width=1,
                    border_color="#454545" if ctk.get_appearance_mode() == "Dark" else "black"
                )
                subframe.grid(row=i, column=j, padx=5, pady=5, sticky="nsew")
                subframes.append(subframe)
                
                try:
                    # Определяем, какой график строить в каждом subframe
                    if i == 0 and j == 0:
                        plot_phi_vs_pk(master=subframe)
                    elif i == 0 and j == 1:
                        plot_eta_vs_pk(master=subframe)
                    elif i == 1 and j == 0:
                        plot_he_vs_pk(master=subframe)
                    elif i == 1 and j == 1:
                        plot_eta_he_vs_pk(master=subframe)
                except Exception as e:
                    print(f"Ошибка при построении графика: {e}")
                    create_empty_plot_in_frame(subframe)
        
        # Настраиваем веса для правильного растягивания
        display_frame.grid_rowconfigure(0, weight=1)
        display_frame.grid_rowconfigure(1, weight=1)
        display_frame.grid_columnconfigure(0, weight=1)
        display_frame.grid_columnconfigure(1, weight=1)
        
    else:
        # Деактивация - удаляем все подфреймы
        widget.configure(image=widget_icon)
        for child in display_frame.winfo_children():
            child.destroy()
        # Создаем один пустой график
        create_empty_plot()

widget = ctk.CTkButton(settings_frame, text="", text_color="black", width=160, height=160, corner_radius=10, fg_color="#151515", image=widget_icon, hover_color="#252525", command=toggle_widget)
widget.grid(row=99, column=98, rowspan=3, padx=5, pady=(5,5), sticky="w")
widget.active = False  # Добавляем атрибут для отслеживания состояния

def create_empty_3d_plot():
    # Очищаем display_frame
    for child in display_frame.winfo_children():
        child.destroy()
    
    # Создаем фигуру
    fig = plt.Figure(figsize=(6, 4), dpi=100)
    ax = fig.add_subplot(111, projection='3d')
    
    # Настройки темы
    current_theme = get_current_theme()
    text_color = "white" if current_theme == "dark" else "black"
    bg_color = "#151515" if current_theme == "dark" else "white"
    
    # Устанавливаем стандартную ориентацию осей (как на вашем изображении)
    ax.view_init(elev=20, azim=30)  # Стандартный ракурс
    
    # Настройка осей с подписями
    ax.set_xlabel('Z', color=text_color, fontsize=10)
    ax.set_ylabel('Y', color=text_color, fontsize=10)
    ax.set_zlabel('X', color=text_color, fontsize=10)
    
    # Устанавливаем пределы осей (от 0 до 1)
    ax.set_xlim(1, 0)
    ax.set_ylim(0, 1)
    ax.set_zlim(0, 1)
    
    # Добавляем деления на осях как на изображении
    ax.set_xticks([0.0, 0.2, 0.4, 0.6, 0.8, 1.0])
    ax.set_yticks([0.0, 0.2, 0.4, 0.6, 0.8, 1.0])
    ax.set_zticks([0.0, 0.2, 0.4, 0.6, 0.8, 1.0])
    
    # Делаем оси более заметными
    ax.xaxis.set_pane_color((0.95, 0.95, 0.95, 0.1))
    ax.yaxis.set_pane_color((0.95, 0.95, 0.95, 0.1))
    ax.zaxis.set_pane_color((0.95, 0.95, 0.95, 0.1))
    
    # Цвета элементов
    ax.xaxis.label.set_color(text_color)
    ax.yaxis.label.set_color(text_color)
    ax.zaxis.label.set_color(text_color)
    ax.tick_params(axis='x', colors=text_color)
    ax.tick_params(axis='y', colors=text_color)
    ax.tick_params(axis='z', colors=text_color)
    
    # Фон
    ax.set_facecolor(bg_color)
    fig.patch.set_facecolor(bg_color)
    
    # Включаем сетку
    ax.grid(True, linestyle='--', alpha=0.5)

    
    # Создаем и размещаем холст
    canvas_e3 = FigureCanvasTkAgg(fig, master=display_frame)
    canvas_e3.draw()
    canvas_e3.get_tk_widget().pack(fill='both', expand=True, padx=10, pady=10)
    
    return canvas_e3

global fig_phi_light
def plot_phi_vs_pk_mu():
    global canvas_3d, fig_phi_light
    
    try:
        # Очищаем display_frame
        for child in display_frame.winfo_children():
            child.destroy()
        
        # Проверка заполнения полей для πк
        min_pk_text = entry_min_phi.get().strip()
        max_pk_text = entry_max_phi.get().strip()
        step_pk_text = entry_step_phi.get().strip()
        
        # Если не введены значения πк
        if not min_pk_text or not max_pk_text or not step_pk_text:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Не введены значения для минимального, максимального πк или шага"
            )
            return
        
        # Проверка заполнения всех необходимых полей
        required_fields = [
            inp1, inp2, inp3, inp5, inp6, inp7, inp8, inp9,
            inp10, inp11, inp12, inp13, inp15, inp16,
            inp17, inp18, inp100
        ]
        
        for field in required_fields:
            value = field.get().strip()
            if not value:
                create_empty_3d_plot()
                messagebox.showwarning(
                    "Ошибка ввода",
                    "Не все обязательные поля заполнены"
                )
                return

        # Получение диапазона μ (от 0 до 100 с шагом 2)
        mu_values = np.arange(0, 101, 2)
        
        # Получение диапазона πк (без дефолтных значений)
        try:
            min_pk = float(replace_comma_with_dot(entry_min_phi.get())) 
            max_pk = float(replace_comma_with_dot(entry_max_phi.get())) 
            step_pk = float(replace_comma_with_dot(entry_step_phi.get()))
        except ValueError:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк должны быть числовыми"
            )
            return

        # Проверка положительности значений πк и шага
        if min_pk <= 0 or max_pk <= 0 or step_pk <= 0:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк и шаг должны быть положительными"
            )
            return
            
        # Проверка, что минимальное πк меньше максимального
        if min_pk >= max_pk:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода", 
                "Минимальное πк должно быть меньше максимального"
            )
            return
        
        # Проверка, что шаг не слишком большой
        if step_pk > (max_pk - min_pk):
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Шаг слишком большой для заданного диапазона πк"
            )
            return
        
        pk_values = np.arange(min_pk, max_pk + step_pk, step_pk)
        
        # Создаем сетку значений
        PK, MU = np.meshgrid(pk_values, mu_values)
        PHI = np.zeros_like(PK)
        
        # Сохранение исходных значений
        original_pk = inp6.get()
        original_mu = inp14.get()
        
        # Расчет φ для каждой комбинации πк и μ
        for i in range(len(mu_values)):
            for j in range(len(pk_values)):
                # Устанавливаем текущее μ
                inp14.delete(0, ctk.END)
                inp14.insert(0, str(mu_values[i]))
                
                # Устанавливаем текущее πк
                inp6.delete(0, ctk.END)
                inp6.insert(0, str(pk_values[j]))
                
                # Выполняем расчет
                raschet()
                
                # Получаем φ
                phi_text = label35.cget("text")
                if 'φ=' in phi_text:
                    phi_str = phi_text.split('=')[1].strip()
                    phi_value = min(100.0, max(0.0, float(phi_str) * 100))
                    PHI[i,j] = phi_value
                else:
                    PHI[i,j] = 0.0
        
        # Восстановление исходных значений
        inp6.delete(0, ctk.END)
        inp6.insert(0, original_pk)
        inp14.delete(0, ctk.END)
        inp14.insert(0, original_mu)
        raschet()
        
        # Создание 3D графика
        fig = plt.Figure(figsize=(14, 10), dpi=100)
        ax = fig.add_subplot(111, projection='3d')

        # Настройка стиля в зависимости от темы
        current_theme = get_current_theme()
        if current_theme == "dark":
            plt.style.use("dark_background")
            text_color = "white"
            bg_color = "#151515"
        else:
            plt.style.use("default")
            text_color = "black"
            bg_color = "white"
        
        # Построение поверхности
        surf = ax.plot_surface(PK, MU, PHI, cmap='viridis', edgecolor='none')
        fig.colorbar(surf, ax=ax, shrink=0.5, aspect=10, pad=0.1)
        
        # Настройка осей
        ax.set_xlabel('Степень повышения давления (πк)', color=text_color)
        ax.set_ylabel('Степень рекуперации (μ, %)', color=text_color)
        ax.set_zlabel('Коэффициент полезной работы (φ, %)', color=text_color)
        ax.set_title('Зависимость φ от πк и μ', color=text_color)
        
        # Настройка цветов
        ax.xaxis.label.set_color(text_color)
        ax.yaxis.label.set_color(text_color)
        ax.zaxis.label.set_color(text_color)
        ax.tick_params(axis='x', colors=text_color)
        ax.tick_params(axis='y', colors=text_color)
        ax.tick_params(axis='z', colors=text_color)
        
        # Фон
        ax.set_facecolor(bg_color)
        fig.patch.set_facecolor(bg_color)
        
        # Встраивание графика
        canvas_3d = FigureCanvasTkAgg(fig, master=display_frame)
        canvas_3d.draw()
        canvas_3d.get_tk_widget().pack(fill='both', expand=True, padx=10, pady=10)

        # Создание светлой версии
        fig_phi_light = plt.Figure(figsize=(14, 10), dpi=100)
        ax_light = fig_phi_light.add_subplot(111, projection='3d')
        plt.style.use("default")
        text_color_light = "black"
        bg_color_light = "white"

        surf_light = ax_light.plot_surface(PK, MU, PHI, cmap='viridis', edgecolor='none')
        fig_phi_light.colorbar(surf_light, ax=ax_light, shrink=0.5, aspect=10, pad=0.1)

        # 2. Установка черного цвета для линий осей (X, Y, Z)
        ax_light.xaxis.line.set_color("black")  # Ось X
        ax_light.yaxis.line.set_color("black")  # Ось Y
        ax_light.zaxis.line.set_color("black")  # Ось Z

        # 3. Установка черного цвета для сетки
        ax_light.xaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
        ax_light.yaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
        ax_light.zaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
        
        ax_light.set_xlabel('Степень повышения давления (πк)', color=text_color_light)
        ax_light.set_ylabel('Степень рекуперации (μ, %)', color=text_color_light)
        ax_light.set_zlabel('Коэффициент полезной работы (φ, %)', color=text_color_light)
        ax_light.set_title('Зависимость φ от πк и μ', color=text_color_light)
        
        ax_light.xaxis.label.set_color(text_color_light)
        ax_light.yaxis.label.set_color(text_color_light)
        ax_light.zaxis.label.set_color(text_color_light)
        ax_light.tick_params(axis='x', colors=text_color_light)
        ax_light.tick_params(axis='y', colors=text_color_light)
        ax_light.tick_params(axis='z', colors=text_color_light)
        
        ax_light.set_facecolor(bg_color_light)
        fig_phi_light.patch.set_facecolor(bg_color_light)
        
    except Exception as e:
        print(f"Ошибка при построении 3D графика: {e}")
        create_empty_3d_plot()

global fig_eta_light
def plot_eta_vs_pk_mu():
    global canvas_3d2, fig_eta_light
    
    try:
        # Очищаем display_frame
        for child in display_frame.winfo_children():
            child.destroy()

        # Проверка заполнения полей для πк
        min_pk_text = entry_min_eta.get().strip()
        max_pk_text = entry_max_eta.get().strip()
        step_pk_text = entry_step_eta.get().strip()
        
        # Если не введены значения πк
        if not min_pk_text or not max_pk_text or not step_pk_text:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Не введены значения для минимального, максимального πк или шага"
            )
            return
        
        # Проверка заполнения всех необходимых полей
        required_fields = [
            inp1, inp2, inp3, inp5, inp6, inp7, inp8, inp9,
            inp10, inp11, inp12, inp13, inp15, inp16,
            inp17, inp18, inp100
        ]
        
        for field in required_fields:
            value = field.get().strip()
            if not value:
                create_empty_3d_plot()
                messagebox.showwarning(
                    "Ошибка ввода",
                    "Не все обязательные поля заполнены"
                )
                return

        # Получение диапазона μ (аналогично plot_he_vs_pk_mu)
        mu_values = np.arange(0, 101, 2)
        
        # Получение диапазона πк (без дефолтных значений)
        try:
            min_pk = float(replace_comma_with_dot(entry_min_eta.get())) 
            max_pk = float(replace_comma_with_dot(entry_max_eta.get())) 
            step_pk = float(replace_comma_with_dot(entry_step_eta.get()))
        except ValueError:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк должны быть числовыми"
            )
            return

        # Проверка положительности значений πк и шага
        if min_pk <= 0 or max_pk <= 0 or step_pk <= 0:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк и шаг должны быть положительными"
            )
            return
            
        # Проверка, что минимальное πк меньше максимального
        if min_pk >= max_pk:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода", 
                "Минимальное πк должно быть меньше максимального"
            )
            return
        
        # Проверка, что шаг не слишком большой
        if step_pk > (max_pk - min_pk):
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Шаг слишком большой для заданного диапазона πк"
            )
            return
        
        pk_values = np.arange(min_pk, max_pk + step_pk, step_pk)
        
        # Создаем сетку значений
        PK, MU = np.meshgrid(pk_values, mu_values)
        ETA = np.zeros_like(PK)
        
        # Сохранение исходных значений
        original_pk = inp6.get()
        original_mu = inp14.get()
        
        # Расчет η для каждой комбинации πк и μ
        for i in range(len(mu_values)):
            for j in range(len(pk_values)):
                # Устанавливаем текущее μ
                inp14.delete(0, ctk.END)
                inp14.insert(0, str(mu_values[i]))
                
                # Устанавливаем текущее πк
                inp6.delete(0, ctk.END)
                inp6.insert(0, str(pk_values[j]))
                
                # Выполняем расчет
                raschet()
                
                # Получаем η
                eta_text = label34.cget("text")
                if 'ηе=' in eta_text:
                    eta_str = eta_text.split('=')[1].strip()
                    eta_value = min(100.0, max(0.0, float(eta_str) * 100))
                    ETA[i,j] = eta_value
                else:
                    ETA[i,j] = 0.0
        
        # Восстановление исходных значений
        inp6.delete(0, ctk.END)
        inp6.insert(0, original_pk)
        inp14.delete(0, ctk.END)
        inp14.insert(0, original_mu)
        raschet()
        
        # Создание 3D графика
        fig = plt.Figure(figsize=(14, 10), dpi=100)
        ax = fig.add_subplot(111, projection='3d')

        # Настройка стиля в зависимости от темы
        current_theme = get_current_theme()
        if current_theme == "dark":
            plt.style.use("dark_background")
            text_color = "white"
            bg_color = "#151515"
        else:
            plt.style.use("default")
            text_color = "black"
            bg_color = "white"
        
        # Построение поверхности
        surf = ax.plot_surface(PK, MU, ETA, cmap='viridis', edgecolor='none')
        fig.colorbar(surf, ax=ax, shrink=0.5, aspect=10, pad=0.1)
        
        # Настройка осей (теперь X: μ, Y: πк, Z: η)
        ax.set_xlabel('Степень повышения давления (πк)', color=text_color)
        ax.set_ylabel('Степень рекуперации (μ, %)', color=text_color)
        ax.set_zlabel('Эффективный КПД (ηe, %)', color=text_color)
        ax.set_title('Зависимость ηe от πк и μ', color=text_color)
        
        # Настройка цветов
        ax.xaxis.label.set_color(text_color)
        ax.yaxis.label.set_color(text_color)
        ax.zaxis.label.set_color(text_color)
        ax.tick_params(axis='x', colors=text_color)
        ax.tick_params(axis='y', colors=text_color)
        ax.tick_params(axis='z', colors=text_color)
        
        # Фон
        ax.set_facecolor(bg_color)
        fig.patch.set_facecolor(bg_color)
        
        # Встраивание графика
        canvas_3d2 = FigureCanvasTkAgg(fig, master=display_frame)
        canvas_3d2.draw()
        canvas_3d2.get_tk_widget().pack(fill='both', expand=True, padx=10, pady=10)

        # Создание светлой версии
        fig_eta_light = plt.Figure(figsize=(14, 10), dpi=100)
        ax_light = fig_eta_light.add_subplot(111, projection='3d')
        plt.style.use("default")
        text_color_light = "black"
        bg_color_light = "white"

        surf_light = ax_light.plot_surface(PK, MU, ETA, cmap='viridis', edgecolor='none')
        fig_eta_light.colorbar(surf_light, ax=ax_light, shrink=0.5, aspect=10, pad=0.1)

        # 2. Установка черного цвета для линий осей (X, Y, Z)
        ax_light.xaxis.line.set_color("black")  # Ось X
        ax_light.yaxis.line.set_color("black")  # Ось Y
        ax_light.zaxis.line.set_color("black")  # Ось Z

        # 3. Установка черного цвета для сетки
        ax_light.xaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
        ax_light.yaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
        ax_light.zaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
        
            
        ax_light.set_xlabel('Степень повышения давления (πк)', color=text_color_light)
        ax_light.set_ylabel('Степень рекуперации (μ, %)', color=text_color_light)
        ax_light.set_zlabel('Эффективный КПД (ηe, %)', color=text_color_light)
        ax_light.set_title('Зависимость ηe от πк и μ', color=text_color_light)

        ax_light.xaxis.label.set_color(text_color_light)
        ax_light.yaxis.label.set_color(text_color_light)
        ax_light.zaxis.label.set_color(text_color_light)
        ax_light.tick_params(axis='x', colors=text_color_light)
        ax_light.tick_params(axis='y', colors=text_color_light)
        ax_light.tick_params(axis='z', colors=text_color_light)
        
        ax_light.set_facecolor(bg_color_light)
        fig_eta_light.patch.set_facecolor(bg_color_light)
        
    except Exception as e:
        print(f"Ошибка при построении 3D графика: {e}")
        create_empty_3d_plot()

global fig_he_light
def plot_he_vs_pk_mu():
    global canvas_3d_he, fig_he_light
    
    try:
        # Очищаем display_frame
        for child in display_frame.winfo_children():
            child.destroy()
        
        # Проверка заполнения полей для πк
        min_pk_text = entry_min_he.get().strip()
        max_pk_text = entry_max_he.get().strip()
        step_pk_text = entry_step_he.get().strip()
        
        # Если не введены значения πк
        if not min_pk_text or not max_pk_text or not step_pk_text:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Не введены значения для минимального, максимального πк или шага"
            )
            return
        
        # Проверка заполнения всех необходимых полей
        required_fields = [
            inp1, inp2, inp3, inp5, inp6, inp7, inp8, inp9,
            inp10, inp11, inp12, inp13, inp15, inp16,
            inp17, inp18, inp100
        ]
        
        for field in required_fields:
            value = field.get().strip()
            if not value:
                create_empty_3d_plot()
                messagebox.showwarning(
                    "Ошибка ввода",
                    "Не все обязательные поля заполнены"
                )
                return

        # Получение диапазона μ (от 0 до 100 с шагом 2)
        mu_values = np.arange(0, 101, 2)
        
        # Получение диапазона πк (без дефолтных значений)
        try:
            min_pk = float(replace_comma_with_dot(entry_min_he.get())) 
            max_pk = float(replace_comma_with_dot(entry_max_he.get())) 
            step_pk = float(replace_comma_with_dot(entry_step_he.get())) 
        except ValueError:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк должны быть числовыми"
            )
            return

        # Проверка положительности значений πк и шага
        if min_pk <= 0 or max_pk <= 0 or step_pk <= 0:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк и шаг должны быть положительными"
            )
            return
            
        # Проверка, что минимальное πк меньше максимального
        if min_pk >= max_pk:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода", 
                "Минимальное πк должно быть меньше максимального"
            )
            return
        
        # Проверка, что шаг не слишком большой
        if step_pk > (max_pk - min_pk):
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Шаг слишком большой для заданного диапазона πк"
            )
            return
        
        pk_values = np.arange(min_pk, max_pk + step_pk, step_pk)
        
        # Создаем сетку значений
        PK, MU = np.meshgrid(pk_values, mu_values)
        HE = np.zeros_like(PK)
        
        # Сохранение исходных значений
        original_pk = inp6.get()
        original_mu = inp14.get()
        
        # Расчет He для каждой комбинации πк и μ
        for i in range(len(mu_values)):
            for j in range(len(pk_values)):
                # Устанавливаем текущее μ
                inp14.delete(0, ctk.END)
                inp14.insert(0, str(mu_values[i]))
                
                # Устанавливаем текущее πк
                inp6.delete(0, ctk.END)
                inp6.insert(0, str(pk_values[j]))
                
                # Выполняем расчет
                raschet()
                
                # Получаем He
                he_text = label200.cget("text")
                if 'He=' in he_text:
                    he_str = he_text.split('=')[1].strip().split()[0]
                    HE[i,j] = float(he_str.replace(',', '.'))
                else:
                    HE[i,j] = 0.0
        
        # Восстановление исходных значений
        inp6.delete(0, ctk.END)
        inp6.insert(0, original_pk)
        inp14.delete(0, ctk.END)
        inp14.insert(0, original_mu)
        raschet()
        
        # Создание 3D графика
        fig = plt.Figure(figsize=(14, 10), dpi=100)
        ax = fig.add_subplot(111, projection='3d')

        # Настройка стиля в зависимости от темы
        current_theme = get_current_theme()
        if current_theme == "dark":
            plt.style.use("dark_background")
            text_color = "white"
            bg_color = "#151515"
        else:
            plt.style.use("default")
            text_color = "black"
            bg_color = "white"
        
        # Построение поверхности
        surf = ax.plot_surface(PK, MU, HE, cmap='viridis', edgecolor='none')
        fig.colorbar(surf, ax=ax, shrink=0.5, aspect=10, pad=0.1)
        
        # Настройка осей
        ax.set_xlabel('Степень повышения давления (πк)', color=text_color)
        ax.set_ylabel('Степень рекуперации (μ, %)', color=text_color)
        ax.set_zlabel('Эффективная удельная работа (He, кДж/кг)', color=text_color)
        ax.set_title('Зависимость He от πк и μ', color=text_color)
        
        # Настройка цветов
        ax.xaxis.label.set_color(text_color)
        ax.yaxis.label.set_color(text_color)
        ax.zaxis.label.set_color(text_color)
        ax.tick_params(axis='x', colors=text_color)
        ax.tick_params(axis='y', colors=text_color)
        ax.tick_params(axis='z', colors=text_color)
        
        # Фон
        ax.set_facecolor(bg_color)
        fig.patch.set_facecolor(bg_color)
        
        # Встраивание графика
        canvas_3d_he = FigureCanvasTkAgg(fig, master=display_frame)
        canvas_3d_he.draw()
        canvas_3d_he.get_tk_widget().pack(fill='both', expand=True, padx=10, pady=10)

        # Создание светлой версии
        fig_he_light = plt.Figure(figsize=(14, 10), dpi=100)
        ax_light = fig_he_light.add_subplot(111, projection='3d')
        plt.style.use("default")
        text_color_light = "black"
        bg_color_light = "white"

        surf_light = ax_light.plot_surface(PK, MU, HE, cmap='viridis', edgecolor='none')
        fig_he_light.colorbar(surf_light, ax=ax_light, shrink=0.5, aspect=10, pad=0.1)

        # 2. Установка черного цвета для линий осей (X, Y, Z)
        ax_light.xaxis.line.set_color("black")  # Ось X
        ax_light.yaxis.line.set_color("black")  # Ось Y
        ax_light.zaxis.line.set_color("black")  # Ось Z

        # 3. Установка черного цвета для сетки
        ax_light.xaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
        ax_light.yaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
        ax_light.zaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
            
        ax_light.set_xlabel('Степень повышения давления (πк)', color=text_color_light)
        ax_light.set_ylabel('Степень рекуперации (μ, %)', color=text_color_light)
        ax_light.set_zlabel('Эффективная удельная работа (He, кДж/кг)', color=text_color_light)
        ax_light.set_title('Зависимость He от πк и μ', color=text_color_light)

        ax_light.xaxis.label.set_color(text_color_light)
        ax_light.yaxis.label.set_color(text_color_light)
        ax_light.zaxis.label.set_color(text_color_light)
        ax_light.tick_params(axis='x', colors=text_color_light)
        ax_light.tick_params(axis='y', colors=text_color_light)
        ax_light.tick_params(axis='z', colors=text_color_light)
        
        ax_light.set_facecolor(bg_color_light)
        fig_he_light.patch.set_facecolor(bg_color_light)
        
    except Exception as e:
        print(f"Ошибка при построении 3D графика: {e}")
        create_empty_3d_plot()

global fig_eta_mu_he_light
def plot_eta_vs_mu_he():
    global canvas_3d4, fig_eta_mu_he_light
    
    try:
        # Очищаем display_frame
        for child in display_frame.winfo_children():
            child.destroy()

        # Проверка заполнения полей для πк
        min_pk_text = entry_min_eta_he.get().strip()
        max_pk_text = entry_max_eta_he.get().strip()
        step_pk_text = entry_step_eta_he.get().strip()
        
        # Если не введены значения πк
        if not min_pk_text or not max_pk_text or not step_pk_text:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Не введены значения для минимального, максимального πк или шага"
            )
            return
        
        # Проверка заполнения всех необходимых полей
        required_fields = [
            inp1, inp2, inp3, inp5, inp6, inp7, inp8, inp9,
            inp10, inp11, inp12, inp13, inp15, inp16,
            inp17, inp18, inp100
        ]
        
        for field in required_fields:
            value = field.get().strip()
            if not value:
                create_empty_3d_plot()
                messagebox.showwarning(
                    "Ошибка ввода",
                    "Не все обязательные поля заполнены"
                )
                return

        # Получение диапазона μ (от 0 до 100 с шагом 2)
        mu_values = np.arange(0, 101, 2)
         
        # Получение диапазона πк (без дефолтных значений)
        try:
            min_pk = float(replace_comma_with_dot(entry_min_eta_he.get())) 
            max_pk = float(replace_comma_with_dot(entry_max_eta_he.get())) 
            step_pk = float(replace_comma_with_dot(entry_step_eta_he.get())) 
        except ValueError:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк должны быть числовыми"
            )
            return

        # Проверка положительности значений πк и шага
        if min_pk <= 0 or max_pk <= 0 or step_pk <= 0:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Значения πк и шаг должны быть положительными"
            )
            return
            
        # Проверка, что минимальное πк меньше максимального
        if min_pk >= max_pk:
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода", 
                "Минимальное πк должно быть меньше максимального"
            )
            return
        
        # Проверка, что шаг не слишком большой
        if step_pk > (max_pk - min_pk):
            create_empty_3d_plot()
            messagebox.showwarning(
                "Ошибка ввода",
                "Шаг слишком большой для заданного диапазона πк"
            )
            return
        
        pk_values = np.arange(min_pk, max_pk + step_pk, step_pk)
        
        # Создаем сетку значений
        PK, MU = np.meshgrid(pk_values, mu_values)
        ETA = np.zeros_like(PK)
        HE = np.zeros_like(PK)
        
        # Сохранение исходных значений
        original_pk = inp6.get()
        original_mu = inp14.get()
        
        # Расчет η и He для каждой комбинации πк и μ
        for i in range(len(mu_values)):
            for j in range(len(pk_values)):
                # Устанавливаем текущее μ
                inp14.delete(0, ctk.END)
                inp14.insert(0, str(mu_values[i]))
                
                # Устанавливаем текущее πк
                inp6.delete(0, ctk.END)
                inp6.insert(0, str(pk_values[j]))
                
                # Выполняем расчет
                raschet()
                
                # Получаем η (умножаем на 100 и ограничиваем диапазон 0-70%)
                eta_text = label34.cget("text")
                if 'ηе=' in eta_text:
                    eta_str = eta_text.split('=')[1].strip()
                    eta_value = min(70.0, max(0.0, float(eta_str) * 100))  # Ограничение 0-70%
                    ETA[i,j] = eta_value
                else:
                    ETA[i,j] = 0.0
                
                # Получаем He
                he_text = label200.cget("text")
                if 'He=' in he_text:
                    he_str = he_text.split('=')[1].strip().split()[0]
                    HE[i,j] = float(he_str.replace(',', '.'))
                else:
                    HE[i,j] = 0.0
        
        # Восстановление исходных значений
        inp6.delete(0, ctk.END)
        inp6.insert(0, original_pk)
        inp14.delete(0, ctk.END)
        inp14.insert(0, original_mu)
        raschet()
        
        # Создание 3D графика
        fig = plt.Figure(figsize=(14, 10), dpi=100)
        ax = fig.add_subplot(111, projection='3d')

        # Настройка стиля в зависимости от темы
        current_theme = get_current_theme()
        if current_theme == "dark":
            plt.style.use("dark_background")
            text_color = "white"
            bg_color = "#151515"
        else:
            plt.style.use("default")
            text_color = "black"
            bg_color = "white"
        
        # Построение поверхности (He от η и μ)
        surf = ax.plot_surface(ETA, MU, HE, cmap='viridis', edgecolor='none', vmin=0)
        
        # Установка пределов для осей
        ax.set_xlim(0, 80)  # η ограничен 0-70%
        ax.set_ylim(0, 100)  # μ в процентах от 0 до 100
        
        # Настройка цветовой шкалы
        cbar = fig.colorbar(surf, ax=ax, shrink=0.5, aspect=10, pad=0.1)
        cbar.set_label('', color=text_color)
        cbar.ax.yaxis.set_tick_params(color=text_color)
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color=text_color)
        
        # Настройка осей
        ax.set_xlabel('Эффективный КПД (ηe, %)', color=text_color)
        ax.set_ylabel('Степень рекуперации (μ, %)', color=text_color)
        ax.set_zlabel('Эффективная удельная работа (He, кДж/кг)', color=text_color)
        ax.set_title('Зависимость He от ηe и μ', color=text_color)  # Обновлен заголовок
        
        # Настройка цветов осей
        ax.xaxis.label.set_color(text_color)
        ax.yaxis.label.set_color(text_color)
        ax.zaxis.label.set_color(text_color)
        ax.tick_params(axis='x', colors=text_color)
        ax.tick_params(axis='y', colors=text_color)
        ax.tick_params(axis='z', colors=text_color)
        
        # Фон
        ax.set_facecolor(bg_color)
        fig.patch.set_facecolor(bg_color)
        
        # Встраивание графика
        canvas_3d4 = FigureCanvasTkAgg(fig, master=display_frame)
        canvas_3d4.draw()
        canvas_3d4.get_tk_widget().pack(fill='both', expand=True, padx=10, pady=10)

         # Создание светлой версии
        fig_eta_mu_he_light = plt.Figure(figsize=(14, 10), dpi=100)
        ax_light = fig_eta_mu_he_light.add_subplot(111, projection='3d')
        plt.style.use("default")
        text_color_light = "black"
        bg_color_light = "white"

        surf_light = ax_light.plot_surface(ETA, MU, HE, cmap='viridis', edgecolor='none', vmin=0)
            
        ax_light.set_xlim(0, 50)
        ax_light.set_ylim(0, 100)
            
        fig_eta_mu_he_light.colorbar(surf_light, ax=ax_light, shrink=0.5, aspect=10, pad=0.1)

        # 2. Установка черного цвета для линий осей (X, Y, Z)
        ax_light.xaxis.line.set_color("black")  # Ось X
        ax_light.yaxis.line.set_color("black")  # Ось Y
        ax_light.zaxis.line.set_color("black")  # Ось Z

        # 3. Установка черного цвета для сетки
        ax_light.xaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
        ax_light.yaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
        ax_light.zaxis._axinfo["grid"].update({"color": "gray", "linestyle": "-", "linewidth": 0.5})
            
        ax_light.set_xlabel('Эффективный КПД (ηe, %)', color=text_color_light)
        ax_light.set_ylabel('Степень рекуперации (μ, %)', color=text_color_light)
        ax_light.set_zlabel('Эффективная удельная работа (He, кДж/кг)', color=text_color_light)
        ax_light.set_title('Зависимость He от ηe и μ', color=text_color_light)

        ax_light.xaxis.label.set_color(text_color_light)
        ax_light.yaxis.label.set_color(text_color_light)
        ax_light.zaxis.label.set_color(text_color_light)
        ax_light.tick_params(axis='x', colors=text_color_light)
        ax_light.tick_params(axis='y', colors=text_color_light)
        ax_light.tick_params(axis='z', colors=text_color_light)
        
        ax_light.set_facecolor(bg_color_light)
        fig_eta_mu_he_light.patch.set_facecolor(bg_color_light)
        
    except Exception as e:
        print(f"Ошибка при построении 3D графика: {e}")
        create_empty_3d_plot()

def create_empty_plot_in_frame(parent_frame):
    """Создает пустой график в указанном фрейме"""
    # Удаляем предыдущие графики
    for child in parent_frame.winfo_children():
        child.destroy()
    # Создаем фигуру
    fig = plt.Figure(figsize=(4, 3), dpi=100)
    ax = fig.add_subplot(111)
    
    # Настройки в зависимости от темы
    current_theme = get_current_theme()
    if current_theme == "dark":
        plt.style.use("dark_background")
        text_color = "white"
        bg_color = "#151515"
    else:
        plt.style.use("default")
        text_color = "black"
        bg_color = "white"
    
    # Настройка осей и сетки
    ax.set_xlabel('', color=text_color)
    ax.set_ylabel('', color=text_color)
    ax.grid(True, color="gray" if current_theme == "dark" else "black")
    
    # Настройка цветов
    ax.set_facecolor(bg_color)
    fig.patch.set_facecolor(bg_color)
    ax.tick_params(axis='x', colors=text_color)
    ax.tick_params(axis='y', colors=text_color)
    
    for spine in ax.spines.values():
        spine.set_edgecolor(text_color)
    
    # Создаем холст и отображаем
    canvas = FigureCanvasTkAgg(fig, master=parent_frame)
    canvas.draw()
    canvas.get_tk_widget().pack(fill='both', expand=True, padx=5, pady=5)
    
    return canvas  # Возвращаем холст для возможного дальнейшего использования


# Добавление кнопки в left_method_frame1 (первая кнопка для первого графика)
toggle_graph_frame_btn = ctk.CTkButton(
    left_method_frame1,
    text="",
    command=lambda: toggle_graph_frame1(graph_frame),
    fg_color="#151515",#151515
    hover_color="#252525",
    anchor="center",
    height=185,
    image=graph_icon,  # Добавляем иконку
    compound="top",
    corner_radius=15, border_width=1, border_color="#454545"  
)
toggle_graph_frame_btn.grid(row=1, column=0, pady=(0, 5), padx=(10,10), sticky="nsew")


# Создаем фреймы с начальным цветом в зависимости от темы
initial_color = "#252525" if ctk.get_appearance_mode() == "Dark" else "#868a91"

graph_frame = ctk.CTkFrame(left_method_frame1, fg_color=initial_color)
graph_frame2 = ctk.CTkFrame(left_method_frame1, fg_color=initial_color)
graph_frame3 = ctk.CTkFrame(left_method_frame1, fg_color=initial_color)
graph_frame4 = ctk.CTkFrame(left_method_frame1, fg_color=initial_color)

# Фрейм для первого графика
graph_frame.grid(row=2, column=0, sticky="nsew", pady=(0, 5))
graph_frame.grid_remove()

# Настройка весов строк для правильного растягивания
graph_frame.grid_columnconfigure(0, weight=1)  # Кнопка 1
graph_frame.grid_columnconfigure(1, weight=1)  # Кнопка 2

# Добавление кнопки в left_method_frame1 (вторая кнопка для второго графика)
toggle_graph_frame_btn2 = ctk.CTkButton(
    left_method_frame1,
    text="",
    command=lambda: toggle_graph_frame2(graph_frame2),
    fg_color="#151515",#151515
    hover_color="#252525",
    image=graph2_icon,  # Добавляем иконку
    compound="top",
    anchor="center",
    height=185, corner_radius=15, border_width=1, border_color="#454545"  
)
toggle_graph_frame_btn2.grid(row=3, column=0, pady=(5, 5), padx=(10,10), sticky="nsew")

# Фрейм для второго графика
graph_frame2.grid(row=4, column=0, sticky="nsew", pady=(0, 10))
graph_frame2.grid_remove()

# Настройка весов строк для правильного растягивания
graph_frame2.grid_columnconfigure(0, weight=1)  # Кнопка 1
graph_frame2.grid_columnconfigure(1, weight=1)  # Кнопка 2

# Добавление кнопки в left_method_frame1 (3 кнопка для 3 графика)
toggle_graph_frame_btn3 = ctk.CTkButton(
    left_method_frame1,
    text="",
    command=lambda: toggle_graph_frame3(graph_frame3),
    fg_color="#151515",#151515
    hover_color="#252525",
    image=graph3_icon,  # Добавляем иконку
    compound="top",
    anchor="center",
    height=185, corner_radius=15, border_width=1, border_color="#454545"  
)
toggle_graph_frame_btn3.grid(row=5, column=0, pady=(5, 5), padx=(10,10), sticky="nsew")

# Фрейм для 3 графика
graph_frame3.grid(row=6, column=0, sticky="nsew", pady=(0, 10))
graph_frame3.grid_remove()

# Настройка весов строк для правильного растягивания
graph_frame3.grid_columnconfigure(0, weight=1)  # Кнопка 1
graph_frame3.grid_columnconfigure(1, weight=1)  # Кнопка 2

# Добавление кнопки в left_method_frame1 (4 кнопка для 4 графика)
toggle_graph_frame_btn4 = ctk.CTkButton(
    left_method_frame1,
    text="",
    command=lambda: toggle_graph_frame4(graph_frame4),
    fg_color="#151515",#151515
    hover_color="#252525",
    image=graph4_icon,  # Добавляем иконку
    compound="top",
    anchor="center",
    height=185, corner_radius=15, border_width=1, border_color="#454545"  
)
toggle_graph_frame_btn4.grid(row=7, column=0, pady=(5, 0), padx=(10,10), sticky="nsew")

# Фрейм для 4 графика
graph_frame4.grid(row=8, column=0, sticky="nsew", pady=(0, 10))
graph_frame4.grid_remove()

# Настройка весов строк для правильного растягивания
graph_frame4.grid_columnconfigure(0, weight=1)  # Кнопка 1
graph_frame4.grid_columnconfigure(1, weight=1)  # Кнопка 2


#==================================================================================================================================================


# Вкладка "Графики 1" (graph_frame)
lb_comp = ctk.CTkLabel(graph_frame, text="Построение графика зависимости φ от πк", font=bold_font)
lb_comp.grid(row=0, column=0, columnspan=2, padx=5, pady=5)

# Поля для графика зависимости φ от πк
label_min_phi = ctk.CTkLabel(graph_frame, text="Минимальное πк:")
label_min_phi.grid(row=1, column=0, padx=5, pady=5)
entry_min_phi = ctk.CTkEntry(graph_frame)
entry_min_phi.grid(row=1, column=1, padx=5, pady=5)

label_max_phi = ctk.CTkLabel(graph_frame, text="Максимальное πк:")
label_max_phi.grid(row=2, column=0, padx=5, pady=5)
entry_max_phi = ctk.CTkEntry(graph_frame)
entry_max_phi.grid(row=2, column=1, padx=5, pady=5)

label_step_phi = ctk.CTkLabel(graph_frame, text="Шаг πк:")
label_step_phi.grid(row=3, column=0, padx=5, pady=5)
entry_step_phi = ctk.CTkEntry(graph_frame)
entry_step_phi.grid(row=3, column=1, padx=5, pady=5)

# Вкладка "Графики" (graph_frame2)
lb_comp = ctk.CTkLabel(graph_frame2, text="Построение графика зависимости ηe от πк", font=bold_font)
lb_comp.grid(row=0, column=0, columnspan=2, padx=5, pady=5)

# Поля для графика зависимости ηe от πк
label_min_eta = ctk.CTkLabel(graph_frame2, text="Минимальное πк:")
label_min_eta.grid(row=1, column=0, padx=5, pady=5)
entry_min_eta = ctk.CTkEntry(graph_frame2)
entry_min_eta.grid(row=1, column=1, padx=5, pady=5)

label_max_eta = ctk.CTkLabel(graph_frame2, text="Максимальное πк:")
label_max_eta.grid(row=2, column=0, padx=5, pady=5)
entry_max_eta = ctk.CTkEntry(graph_frame2)
entry_max_eta.grid(row=2, column=1, padx=5, pady=5)

label_step_eta = ctk.CTkLabel(graph_frame2, text="Шаг πк:")
label_step_eta.grid(row=3, column=0, padx=5, pady=5)
entry_step_eta = ctk.CTkEntry(graph_frame2)
entry_step_eta.grid(row=3, column=1, padx=5, pady=5)

# Вкладка "Графики" (graph_frame3)
lb_comp = ctk.CTkLabel(graph_frame3, text="Построение графика зависимости He от πк", font=bold_font)
lb_comp.grid(row=0, column=0, columnspan=2, padx=5, pady=5)

# Поля для графика зависимости ηe от πк
label_min_he = ctk.CTkLabel(graph_frame3, text="Минимальное πк:")
label_min_he.grid(row=1, column=0, padx=5, pady=5)
entry_min_he = ctk.CTkEntry(graph_frame3)
entry_min_he.grid(row=1, column=1, padx=5, pady=5)

label_max_he = ctk.CTkLabel(graph_frame3, text="Максимальное πк:")
label_max_he.grid(row=2, column=0, padx=5, pady=5)
entry_max_he = ctk.CTkEntry(graph_frame3)
entry_max_he.grid(row=2, column=1, padx=5, pady=5)

label_step_he = ctk.CTkLabel(graph_frame3, text="Шаг πк:")
label_step_he.grid(row=3, column=0, padx=5, pady=5)
entry_step_he = ctk.CTkEntry(graph_frame3)
entry_step_he.grid(row=3, column=1, padx=5, pady=5)

# Вкладка "Графики" (graph_frame4)
lb_comp = ctk.CTkLabel(graph_frame4, text="Построение графика зависимости ηe от He", font=bold_font)
lb_comp.grid(row=0, column=0, columnspan=2, padx=5, pady=5)

# Поля для графика зависимости ηe от He
label_min_eta_he = ctk.CTkLabel(graph_frame4, text="Минимальное πк:")
label_min_eta_he.grid(row=1, column=0, padx=5, pady=5)
entry_min_eta_he = ctk.CTkEntry(graph_frame4)
entry_min_eta_he.grid(row=1, column=1, padx=5, pady=5)

label_max_eta_he = ctk.CTkLabel(graph_frame4, text="Максимальное πк:")
label_max_eta_he.grid(row=2, column=0, padx=5, pady=5)
entry_max_eta_he = ctk.CTkEntry(graph_frame4)
entry_max_eta_he.grid(row=2, column=1, padx=5, pady=5)

label_step_eta_he = ctk.CTkLabel(graph_frame4, text="Шаг πк:")
label_step_eta_he.grid(row=3, column=0, padx=5, pady=5)
entry_step_eta_he = ctk.CTkEntry(graph_frame4)
entry_step_eta_he.grid(row=3, column=1, padx=5, pady=5)




#==================================================================================================================================================
# Создаем пустой график при запуске
def create_empty_plot(master=None):
    """Создает пустой график в указанном фрейме или основном display_frame"""
    global canvas_e
    
    # Определяем целевой фрейм
    target_frame = master if master else display_frame
    
    # Уничтожаем предыдущие графики в целевом фрейме
    for child in target_frame.winfo_children():
        child.destroy()
    
    # Создаем фигуру
    fig = plt.Figure(figsize=(6, 4), dpi=100)
    ax = fig.add_subplot(111)
    
    # Настройки в зависимости от темы
    current_theme = get_current_theme()
    if current_theme == "dark":
        plt.style.use("dark_background")
        text_color = "white"
        bg_color = "#151515"
    else:
        plt.style.use("default")
        text_color = "black"
        bg_color = "white"
    
    # Настройка осей и сетки
    ax.set_xlabel('', color=text_color)
    ax.set_ylabel('', color=text_color)
    ax.grid(True, color="gray" if current_theme == "dark" else "black")
    
    # Настройка цветов
    ax.set_facecolor(bg_color)
    fig.patch.set_facecolor(bg_color)
    ax.tick_params(axis='x', colors=text_color)
    ax.tick_params(axis='y', colors=text_color)
    
    for spine in ax.spines.values():
        spine.set_edgecolor(text_color)
    
    # Создаем холст и отображаем
    canvas_e = FigureCanvasTkAgg(fig, master=target_frame)
    canvas_e.draw()
    canvas_e.get_tk_widget().pack(fill='both', expand=True, padx=10, pady=10)

# Вызываем функцию создания пустого графика при запуске
create_empty_plot()

def update_button_images(event=None):
    # Получаем текущую высоту фрейма
    frame_height = left_method_frame1.winfo_height()
    
    # Вычисляем новый размер изображения 
    new_size = int(frame_height * 0.8)
    if new_size < 10:  # Минимальный размер
        new_size = 10
    if new_size > 130:  # Максимальный размер
        new_size = 130
    
    # Обновляем все иконки
    graph_icon.configure(size=(new_size, new_size))
    graph2_icon.configure(size=(new_size, new_size))
    graph3_icon.configure(size=(new_size, new_size))
    graph4_icon.configure(size=(new_size, new_size))

# Привязываем функцию обновления к изменению размера фрейма
left_method_frame1.bind("<Configure>", update_button_images)

# Также вызываем при старте, чтобы установить начальные размеры
root.after(100, update_button_images)

clear1_1 = ctk.CTkButton(graph_frame, text='', image=trash_icon,  fg_color="#20252e", hover_color="#9b2d30", width=1000, command = cleargraf) 
clear1_1.grid(column=0, columnspan=2, row=4, padx=5, pady=5)

clear12_12 = ctk.CTkButton(graph_frame2, text='', image=trash_icon,  fg_color="#20252e", hover_color="#9b2d30", width=1000, command = cleargraf2) 
clear12_12.grid(column=0, columnspan=2, row=4, padx=5, pady=5)
# savegraf2_button = ctk.CTkButton(graph_frame2, text='', image=word_icon, width=140, height=25, fg_color="#009dda", hover_color="#2dcbff", corner_radius=5)#, command=on_save_button_click)
# savegraf2_button.grid(column=0, columnspan=2, row=4, padx=5, pady=5)

clear11_1 = ctk.CTkButton(graph_frame3, text='', image=trash_icon, fg_color="#20252e", hover_color="#9b2d30", width=1000, command = cleargraf3)
clear11_1.grid(column=0, columnspan=2, row=4, padx=5, pady=5)

clear11_11 = ctk.CTkButton(graph_frame4, text='', image=trash_icon, fg_color="#20252e", hover_color="#9b2d30", width=1000, command = cleargraf4)
clear11_11.grid(column=0, columnspan=2, row=4, padx=5, pady=5)

# Настройка весов для фиксации положения кнопки
settings_frame.grid_rowconfigure(98, weight=0)
settings_frame.grid_rowconfigure(99, weight=0)  # Растягиваем строку, чтобы кнопка была прижата к низу
settings_frame.grid_rowconfigure(100, weight=1)  # Растягиваем строку, чтобы кнопка была прижата к низу
settings_frame.grid_rowconfigure(101, weight=0)  # Растягиваем строку, чтобы кнопка была прижата к низу
settings_frame.grid_columnconfigure(98, weight=1)
settings_frame.grid_columnconfigure(99, weight=0)
settings_frame.grid_columnconfigure(100, weight=0)  # Растягиваем колонку, чтобы кнопка была прижата к правому краю

def save_plot_as_pdf():
    """Сохраняет 2D график в PDF через диалоговое окно."""
    global saved_fig_he
    
    # Проверка наличия фигуры
    if saved_fig_he is None:
        messagebox.showwarning("Предупреждение", "Нет доступного 2D графика зависимости He от πк для сохранения!")
        print("Ошибка: Нет доступного графика для сохранения!")
        return
    
    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()
    
    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить 2D график зависимости He от πк как PDF",
        initialfile="He от πк.pdf"
    )
    
    # Закрываем Tkinter окно
    root.destroy()
    
    # Сохраняем если путь получен
    if file_path:
        try:
            saved_fig_he.savefig(file_path, format="pdf", bbox_inches="tight")
            messagebox.showinfo("Успех", f"2D график He от πк успешно сохранен в:\n{file_path}")
            print(f"График успешно сохранен в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении 2D графика He от πк:\n{str(e)}")
            print(f"Ошибка при сохранении файла: {str(e)}")
    else:
        messagebox.showinfo("Инфо", "Сохранение 2D графика He от πк отменено пользователем.")
        print("Сохранение отменено пользователем.")

#saved_fig_eta_he
def save_plot_as_pdf4():
    """Сохраняет 2D график в PDF через диалоговое окно."""
    global saved_fig_eta_he
    
    # Проверка наличия фигуры
    if saved_fig_eta_he is None:
        messagebox.showwarning("Предупреждение", "Нет доступного 2D графика зависимости ηe от He для сохранения!")
        print("Ошибка: Нет доступного графика для сохранения!")
        return
    
    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()
    
    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить 2D график зависимости ηe от He как PDF",
        initialfile="ηe от He.pdf"
        
    )
    
    # Закрываем Tkinter окно
    root.destroy()
    
    # Сохраняем если путь получен
    if file_path:
        try:
            saved_fig_eta_he.savefig(file_path, format="pdf", bbox_inches="tight")
            messagebox.showinfo("Успех", f"2D график ηe от He успешно сохранен в:\n{file_path}")
            print(f"График успешно сохранен в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении 2D графика ηe от He:\n{str(e)}")
            print(f"Ошибка при сохранении файла: {str(e)}")
    else:
        messagebox.showinfo("Инфо", "Сохранение 2D графика ηe от He отменено пользователем.")
        print("Сохранение отменено пользователем.")

#saved_fig_eta
def save_plot_as_pdf2():
    """Сохраняет 2D график в PDF через диалоговое окно."""
    global saved_fig_eta
    
    # Проверка наличия фигуры
    if saved_fig_eta is None:
        messagebox.showwarning("Предупреждение", "Нет доступного 2D графика зависимости ηe от πк для сохранения!")
        print("Ошибка: Нет доступного графика для сохранения!")
        return
    
    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()
    
    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить 2D график зависимости ηe от πк как PDF",
        initialfile="ηe от πк.pdf"
    )
    
    # Закрываем Tkinter окно
    root.destroy()
    
    # Сохраняем если путь получен
    if file_path:
        try:
            saved_fig_eta.savefig(file_path, format="pdf", bbox_inches="tight")
            messagebox.showinfo("Успех", f"2D график ηe от πк успешно сохранен в:\n{file_path}")
            print(f"График успешно сохранен в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении 2D графика ηe от πк:\n{str(e)}")
            print(f"Ошибка при сохранении файла: {str(e)}")
    else:
        messagebox.showinfo("Инфо", "Сохранение 2D графика ηe от πк отменено пользователем.")
        print("Сохранение отменено пользователем.")

#global_save_fig
def save_plot_as_pdf1():
    """Сохраняет 2D график в PDF через диалоговое окно."""
    global global_save_fig
    
    # Проверка наличия фигуры
    if global_save_fig is None:
        messagebox.showwarning("Предупреждение", "Нет доступного 2D графика зависимости φ от πк для сохранения!")
        print("Ошибка: Нет доступного графика для сохранения!")
        return
    
    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()
    
    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить 2D график зависимости φ от πк как PDF",
        initialfile="φ от πк.pdf"
    )
    
    # Закрываем Tkinter окно
    root.destroy()
    
    # Сохраняем если путь получен
    if file_path:
        try:
            global_save_fig.savefig(file_path, format="pdf", bbox_inches="tight")
            messagebox.showinfo("Успех", f"2D график φ от πк успешно сохранен в:\n{file_path}")
            print(f"График успешно сохранен в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении 2D графика φ от πк:\n{str(e)}")
            print(f"Ошибка при сохранении файла: {str(e)}")
    else:
        messagebox.showinfo("Инфо", "Сохранение 2D графика φ от πк отменено пользователем.")
        print("Сохранение отменено пользователем.")

# fig_phi_light = None
def save_plot_as_pdf_3d1():
    """Сохраняет текущий график в PDF через диалоговое окно."""
    global fig_phi_light
    # global global_fig
    
    # Проверка наличия фигуры
    if fig_phi_light is None:
        print("Ошибка: Нет доступного графика для сохранения!")
        return
    
    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()
    
    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить график как PDF",
        initialfile="φ от πк и μ.pdf"
    )
    
    # Закрываем Tkinter окно
    root.destroy()
    
    # Сохраняем если путь получен
    if file_path:
        try:
            fig_phi_light.savefig(file_path, format="pdf", bbox_inches="tight")
            print(f"График успешно сохранен в:\n{file_path}")
        except Exception as e:
            print(f"Ошибка при сохранении файла: {str(e)}")
    else:
        print("Сохранение отменено пользователем.")
# fig_eta_light = None
def save_plot_as_pdf_3d2():
    """Сохраняет текущий 3D график в PDF через диалоговое окно."""
    global fig_eta_light
    
    # Проверка наличия фигуры
    if fig_eta_light is None:
        messagebox.showwarning("Предупреждение", "Нет доступного 3D графика ηe от πк и μ для сохранения!")
        print("Ошибка: Нет доступного графика для сохранения!")
        return
    
    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()
    
    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить 3D график зависимости ηe от πк и μ как PDF",
        initialfile="ηe от πк и μ.pdf"
    )
    
    # Закрываем Tkinter окно
    root.destroy()
    
    # Сохраняем если путь получен
    if file_path:
        try:
            fig_eta_light.savefig(file_path, format="pdf", bbox_inches="tight")
            messagebox.showinfo("Успех", f"3D график ηe от πк и μ успешно сохранен в:\n{file_path}")
            print(f"График успешно сохранен в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении 3D графика ηe от πк и μ:\n{str(e)}")
            print(f"Ошибка при сохранении файла: {str(e)}")
    else:
        messagebox.showinfo("Инфо", "Сохранение 3D графика ηe от πк и μ отменено пользователем.")
        print("Сохранение отменено пользователем.")

# fig_he_light = None
def save_plot_as_pdf_3d3():
    """Сохраняет текущий 3D график в PDF через диалоговое окно."""
    global fig_he_light
    
    # Проверка наличия фигуры
    if fig_he_light is None:
        messagebox.showwarning("Предупреждение", "Нет доступного 3D графика для сохранения!")
        print("Ошибка: Нет доступного графика для сохранения!")
        return
    
    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()
    
    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить 3D график зависимости He от πк и μ как PDF",
        initialfile="He от πк и μ.pdf"
    )
    
    # Закрываем Tkinter окно
    root.destroy()
    
    # Сохраняем если путь получен
    if file_path:
        try:
            fig_he_light.savefig(file_path, format="pdf", bbox_inches="tight")
            messagebox.showinfo("Успех", f"3D график успешно сохранен в:\n{file_path}")
            print(f"График успешно сохранен в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
            print(f"Ошибка при сохранении файла: {str(e)}")
    else:
        messagebox.showinfo("Инфо", "Сохранение 3D графика отменено пользователем.")
        print("Сохранение отменено пользователем.")

# fig_eta_mu_he_light = None
def save_plot_as_pdf_3d4():
    """Сохраняет текущий 3D график в PDF через диалоговое окно."""
    global fig_eta_mu_he_light
    
    # Проверка наличия фигуры
    if fig_eta_mu_he_light is None:
        messagebox.showwarning("Предупреждение", "Нет доступного 3D графика для сохранения!")
        return
    
    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()
    
    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить 3D график зависимости He от ηe и μ как PDF",
        initialfile="He от ηe и μ.pdf"
    )
    
    # Закрываем Tkinter окно
    root.destroy()
    
    # Сохраняем если путь получен
    if file_path:
        try:
            fig_eta_mu_he_light.savefig(file_path, format="pdf", bbox_inches="tight")
            messagebox.showinfo("Успех", f"3D график успешно сохранен в:\n{file_path}")
            print(f"График успешно сохранен в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
            print(f"Ошибка при сохранении файла: {str(e)}")
    else:
        messagebox.showinfo("Инфо", "Сохранение 3D графика отменено пользователем.")
        print("Сохранение отменено пользователем.")
#=========================================================================================================
import os
import re
import tempfile
from docx import Document
from docx.shared import Cm
import matplotlib.pyplot as plt
from tkinter import messagebox

def save_plot_as_word_2d_all():
    global global_save_fig, saved_fig_eta, saved_fig_he, saved_fig_eta_he
    
    # Создаем временную директорию для сохранения изображений
    with tempfile.TemporaryDirectory() as temp_dir:
        # Сохраняем все фигуры в файлы
        figures = {
            '$': global_save_fig,
            '&': saved_fig_eta,
            '!': saved_fig_he,
            '^': saved_fig_eta_he
        }
        
        file_paths = {}
        for symbol, fig in figures.items():
            if fig is None:
                messagebox.showwarning("Предупреждение", f"График для символа {symbol} отсутствует.")
                continue
            try:
                path = os.path.join(temp_dir, f"{symbol}.png")
                fig.savefig(path, dpi=600, bbox_inches='tight')
                file_paths[symbol] = path
                plt.close(fig)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить график {symbol}:\n{str(e)}")
                return

        # Загрузка шаблона документа
        try:
            doc = Document("шаблон_графики2д.docx")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть шаблон документа:\n{str(e)}")
            return

        # Функция для обработки параграфов
        def process_paragraph(paragraph):
            original_text = paragraph.text
            if not original_text:
                return
            
            # Создаем паттерн для поиска символов
            symbols = re.escape(''.join(file_paths.keys()))
            pattern = re.compile(f"([{symbols}])|([^{symbols}]+)")
            
            tokens = pattern.findall(original_text)
            paragraph.clear()
            
            for token in tokens:
                symbol, text = token
                if symbol:
                    if symbol in file_paths:
                        run = paragraph.add_run()
                        try:
                            run.add_picture(file_paths[symbol], width=Cm(15), height=Cm(10))
                        except Exception as e:
                            messagebox.showerror("Ошибка", f"Ошибка вставки изображения {symbol}:\n{str(e)}")
                            raise
                    else:
                        paragraph.add_run(symbol)
                elif text:
                    paragraph.add_run(text)

        try:
            # Обработка обычных параграфов
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)

            # Обработка таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
                        # Обработка вложенных таблиц (рекурсивно)
                        for nested_table in cell.tables:
                            for nested_row in nested_table.rows:
                                for nested_cell in nested_row.cells:
                                    for nested_para in nested_cell.paragraphs:
                                        process_paragraph(nested_para)

            # Сохранение документа
            root = Tk()
            root.withdraw()  # Скрываем основное окно
            root.wm_attributes('-topmost', 1)  # Окно поверх остальных
            
            output_filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Сохранить документ",
                initialfile="результат_графики2д.docx"
            )
            
            root.destroy()  # Закрываем скрытое окно

            if not output_filename:  # Пользователь нажал "Отмена"
                messagebox.showinfo("Инфо", "Сохранение отменено")
                return

            try:
                doc.save(output_filename)
                messagebox.showinfo("Успех", f"Документ успешно сохранен:\n{output_filename}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
                return

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при обработке документа:\n{str(e)}")
            return

#=========================================================================================================
def save_plot_as_pdf_2d_all():
    global global_save_fig, saved_fig_eta, saved_fig_he, saved_fig_eta_he

    # Собираем все фигуры в список и фильтруем пустые значения
    figures = [
        global_save_fig,
        saved_fig_eta,
        saved_fig_he,
        saved_fig_eta_he
    ]
    figures = [fig for fig in figures if fig is not None]

    # Проверка наличия хотя бы одной фигуры
    if not figures:
        messagebox.showwarning("Предупреждение", "Нет доступных 2D графиков для сохранения!")
        return

    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()

    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить все 2D графики как PDF",
        initialfile="результат_графики_2D.pdf"
    )

    # Закрываем Tkinter окно
    root.destroy()

    # Сохраняем если путь получен
    if file_path:
        try:
            with PdfPages(file_path) as pdf:
                for i, fig in enumerate(figures, 1):
                    # Устанавливаем размер фигуры (15см x 10см)
                    # Переводим см в дюймы (1 см = 0.393701 дюйма)
                    fig.set_size_inches(15 * 0.393701, 10 * 0.393701)
                    pdf.savefig(fig, bbox_inches='tight')
                    print(f"График {i}/{len(figures)} успешно добавлен в PDF")
            
            messagebox.showinfo("Успех", f"Все 2D графики ({len(figures)} шт.) успешно сохранены в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
    else:
        messagebox.showinfo("Инфо", "Сохранение 2D графиков отменено пользователем.")
#=========================================================================================================
def save_plot_as_word_3d_all():
    global fig_phi_light, fig_eta_light, fig_he_light, fig_eta_mu_he_light
    # global global_save_fig, saved_fig_eta, saved_fig_he, saved_fig_eta_he
    
    # Создаем временную директорию для сохранения изображений
    with tempfile.TemporaryDirectory() as temp_dir:
        # Сохраняем все фигуры в файлы
        figures = {
            '#': fig_phi_light,
            '{': fig_eta_light,
            '[': fig_he_light,
            ']': fig_eta_mu_he_light
        }
        
        file_paths = {}
        for symbol, fig in figures.items():
            if fig is None:
                messagebox.showwarning("Предупреждение", f"График для символа {symbol} отсутствует.")
                continue
            try:
                path = os.path.join(temp_dir, f"{symbol}.png")
                fig.savefig(path, dpi=600, bbox_inches='tight')
                file_paths[symbol] = path
                plt.close(fig)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить график {symbol}:\n{str(e)}")
                return

        # Загрузка шаблона документа
        try:
            doc = Document("шаблон_графики3д.docx")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть шаблон документа:\n{str(e)}")
            return

        # Функция для обработки параграфов
        def process_paragraph(paragraph):
            original_text = paragraph.text
            if not original_text:
                return
            
            # Создаем паттерн для поиска символов
            symbols = re.escape(''.join(file_paths.keys()))
            pattern = re.compile(f"([{symbols}])|([^{symbols}]+)")
            
            tokens = pattern.findall(original_text)
            paragraph.clear()
            
            for token in tokens:
                symbol, text = token
                if symbol:
                    if symbol in file_paths:
                        run = paragraph.add_run()
                        try:
                            run.add_picture(file_paths[symbol], width=Cm(17), height=Cm(14))
                        except Exception as e:
                            messagebox.showerror("Ошибка", f"Ошибка вставки изображения {symbol}:\n{str(e)}")
                            raise
                    else:
                        paragraph.add_run(symbol)
                elif text:
                    paragraph.add_run(text)

        try:
            # Обработка обычных параграфов
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)

            # Обработка таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
                        # Обработка вложенных таблиц (рекурсивно)
                        for nested_table in cell.tables:
                            for nested_row in nested_table.rows:
                                for nested_cell in nested_row.cells:
                                    for nested_para in nested_cell.paragraphs:
                                        process_paragraph(nested_para)

            # Сохранение документа
            root = Tk()
            root.withdraw()  # Скрываем основное окно
            root.wm_attributes('-topmost', 1)  # Окно поверх остальных
            
            output_filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Сохранить документ",
                initialfile="результат_графики3д.docx"
            )
            
            root.destroy()  # Закрываем скрытое окно

            if not output_filename:  # Пользователь нажал "Отмена"
                messagebox.showinfo("Инфо", "Сохранение отменено")
                return

            try:
                doc.save(output_filename)
                messagebox.showinfo("Успех", f"Документ успешно сохранен:\n{output_filename}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
                return

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при обработке документа:\n{str(e)}")
            return
#=========================================================================================================
def save_plots_as_word_combined(): 
    global fig_phi_light, fig_eta_light, fig_he_light, fig_eta_mu_he_light
    global global_save_fig, saved_fig_eta, saved_fig_he, saved_fig_eta_he
    
    # Создаем временную директорию для сохранения изображений
    with tempfile.TemporaryDirectory() as temp_dir:
        # Сохраняем все фигуры в файлы
        figures_3d = {
            '#': fig_phi_light,
            '{': fig_eta_light,
            '[': fig_he_light,
            ']': fig_eta_mu_he_light
        }
        
        figures_2d = {
            '$': global_save_fig,
            '&': saved_fig_eta,
            '!': saved_fig_he,
            '^': saved_fig_eta_he
        }
        
        file_paths = {}
        
        # Сохраняем 3D графики
        for symbol, fig in figures_3d.items():
            if fig is None:
                messagebox.showwarning("Предупреждение", f"3D график для символа {symbol} отсутствует.")
                continue
            try:
                path = os.path.join(temp_dir, f"{symbol}.png")
                fig.savefig(path, dpi=600, bbox_inches='tight')
                file_paths[symbol] = (path, '3d')  # Помечаем как 3D
                plt.close(fig)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить 3D график {symbol}:\n{str(e)}")
                return
        
        # Сохраняем 2D графики
        for symbol, fig in figures_2d.items():
            if fig is None:
                messagebox.showwarning("Предупреждение", f"2D график для символа {symbol} отсутствует.")
                continue
            try:
                path = os.path.join(temp_dir, f"{symbol}.png")
                fig.savefig(path, dpi=600, bbox_inches='tight')
                file_paths[symbol] = (path, '2d')  # Помечаем как 2D
                plt.close(fig)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить 2D график {symbol}:\n{str(e)}")
                return

        # Загрузка шаблона документа
        try:
            doc = Document("шаблон_графики2д_графики3д.docx")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть шаблон документа:\n{str(e)}")
            return

        # Функция для обработки параграфов
        def process_paragraph(paragraph):
            original_text = paragraph.text
            if not original_text:
                return
            
            # Создаем паттерн для поиска символов
            symbols = re.escape(''.join(file_paths.keys()))
            pattern = re.compile(f"([{symbols}])|([^{symbols}]+)")
            
            tokens = pattern.findall(original_text)
            paragraph.clear()
            
            for token in tokens:
                symbol, text = token
                if symbol:
                    if symbol in file_paths:
                        path, fig_type = file_paths[symbol]
                        run = paragraph.add_run()
                        try:
                            # Устанавливаем разные размеры для 2D и 3D графиков
                            if fig_type == '3d':
                                run.add_picture(path, width=Cm(17), height=Cm(14))
                            else:  # 2d
                                run.add_picture(path, width=Cm(17), height=Cm(11.3))
                        except Exception as e:
                            messagebox.showerror("Ошибка", f"Ошибка вставки изображения {symbol}:\n{str(e)}")
                            raise
                    else:
                        paragraph.add_run(symbol)
                elif text:
                    paragraph.add_run(text)

        try:
            # Обработка обычных параграфов
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)

            # Обработка таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
                        # Обработка вложенных таблиц (рекурсивно)
                        for nested_table in cell.tables:
                            for nested_row in nested_table.rows:
                                for nested_cell in nested_row.cells:
                                    for nested_para in nested_cell.paragraphs:
                                        process_paragraph(nested_para)

            # Сохранение документа
            root = Tk()
            root.withdraw()  # Скрываем основное окно
            root.wm_attributes('-topmost', 1)  # Окно поверх остальных
            
            output_filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Сохранить документ",
                initialfile="результат_графики_комбинированный.docx"
            )
            
            root.destroy()  # Закрываем скрытое окно

            if not output_filename:  # Пользователь нажал "Отмена"
                messagebox.showinfo("Инфо", "Сохранение отменено")
                return

            try:
                doc.save(output_filename)
                messagebox.showinfo("Успех", f"Документ успешно сохранен:\n{output_filename}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
                return

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при обработке документа:\n{str(e)}")
            return
#=========================================================================================================
def fill_template_():
    try:
        # Определяем режим расчета
        if lang.get() == 'без регенерации':
            
                # Открываем шаблон
                doc = Document('шаблон_без_регенерации_график.docx')

                # Словарь для замены значений после "=" в формулах
                replacements = {
                    "Температура окружающей среды": replace_dot_with_comma(inp1.get()),  # Температура окружающей среды
                    "Давление окружающей среды": replace_dot_with_comma(inp2.get()),      # Давление окружающей среды
                    "Начальная температура газа перед турбиной": replace_dot_with_comma(inp3.get()),  # Начальная температура газа перед турбиной
                    "Газовая постоянная воздуха": replace_dot_with_comma(inp4.get()),      # Газовая постоянная воздуха
                    "Показатель изоэнтропы для воздуха": replace_dot_with_comma(inp15.get()),  # Показатель изоэнтропы для воздуха
                    "Показатель изоэнтропы для газа": replace_dot_with_comma(inp16.get()),     # Показатель изоэнтропы для газа
                    "Теплоемкость воздуха": replace_dot_with_comma(inp17.get()),              # Теплоемкость воздуха
                    "Теплоемкость газа": replace_dot_with_comma(inp18.get()),                # Теплоемкость газа
                    "Коэффициент потерь на входе в компрессор": replace_dot_with_comma(inp5.get()),  # Коэффициент потерь на входе в компрессор (делим на 100)
                    "Коэффициент потерь давления воздуха в вых. устройстве": replace_dot_with_comma(inp7.get()),  # Коэффициент потерь давления воздуха в вых. устройстве (делим на 100)
                    "Коэффициент потерь давления воздуха перед КС": replace_dot_with_comma(inp8.get()),  # Коэффициент потерь давления воздуха перед КС (делим на 100)
                    "Коэффициент потерь давления воздуха в КС": replace_dot_with_comma(inp9.get()),      # Коэффициент потерь давления воздуха в КС (делим на 100)
                    "Политропный КПД турбины": divide_by_100(inp10.get()),                      # Политропный КПД турбины (делим на 100)
                    "Механический КПД турбины": divide_by_100(inp11.get()),                     # Механический КПД турбины (делим на 100)
                    "Механический КПД компрессора": divide_by_100(inp12.get()),                 # Механический КПД компрессора (делим на 100)
                    "Адиабатический КПД компрессора": divide_by_100(inp13.get()),               # Адиабатический КПД компрессора (делим на 100)
                    "Степень повышения давления в компрессоре": replace_dot_with_comma(inp6.get()),      # Степень повышения давления в компрессоре
                    "Эффективная мощность": replace_dot_with_comma(inp100.get()),                        # Эффективная мощность
                }

                # Проходим по всем параграфам в документе
                for paragraph in doc.paragraphs:
                    # Проверяем, содержит ли параграф текст, который есть в словаре replacements
                    for key in replacements:
                        if key in paragraph.text:
   

                            # Получаем XML-структуру параграфа
                            p_xml = paragraph._element.xml

                            # Парсим XML с помощью lxml
                            p_tree = etree.fromstring(p_xml)

                            # Ищем все формулы (элементы <m:t>)
                            formulas = p_tree.xpath('.//m:t', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})

                            # Проходим по каждой формуле в параграфе
                            for formula in formulas:
                                formula_text = formula.text
                                if formula_text and "=" in formula_text:
                  

                                    # Разделяем текст на части до и после "="
                                    parts = formula_text.split("=", 1)
                                    if len(parts) == 2:
                                        # Находим маркер перед "="
                                        marker = parts[0].strip()
                       

                                        # Если маркер есть в replacements, добавляем данные после "=" без пробела
                                        if key in replacements:
                                            # Сохраняем исходный текст формулы до "="
                                            original_text = parts[0].strip()
                                            # Добавляем новые данные после "=" без пробела
                                            new_text = f"{original_text}={replacements[key]}"
               

                                            # Обновляем текст формулы
                                            formula.text = new_text

                            # Обновляем XML-структуру параграфа
                            paragraph._element.getparent().replace(paragraph._element, p_tree)

                # Сохраняем измененный документ
                doc.save('отчет_без_регенерации_график.docx')
 
            
        
        elif lang.get() == 'c регенерацией':
            
                # Открываем шаблон
                doc = Document('шаблон_регенерация_график.docx')

                # Словарь для замены значений после "=" в формулах
                replacements = {
                    "Температура окружающей среды": replace_dot_with_comma(inp1.get()),  # Температура окружающей среды
                    "Давление окружающей среды": replace_dot_with_comma(inp2.get()),      # Давление окружающей среды
                    "Начальная температура газа перед турбиной": replace_dot_with_comma(inp3.get()),  # Начальная температура газа перед турбиной
                    "Газовая постоянная воздуха": replace_dot_with_comma(inp4.get()),      # Газовая постоянная воздуха
                    "Показатель изоэнтропы для воздуха": replace_dot_with_comma(inp15.get()),  # Показатель изоэнтропы для воздуха
                    "Показатель изоэнтропы для газа": replace_dot_with_comma(inp16.get()),     # Показатель изоэнтропы для газа
                    "Теплоемкость воздуха": replace_dot_with_comma(inp17.get()),              # Теплоемкость воздуха
                    "Теплоемкость газа": replace_dot_with_comma(inp18.get()),                # Теплоемкость газа
                    "Коэффициент потерь на входе в компрессор": replace_dot_with_comma(inp5.get()),  # Коэффициент потерь на входе в компрессор (делим на 100)
                    "Коэффициент потерь давления воздуха в вых. устройстве": replace_dot_with_comma(inp7.get()),  # Коэффициент потерь давления воздуха в вых. устройстве (делим на 100)
                    "Коэффициент потерь давления воздуха перед КС": replace_dot_with_comma(inp8.get()),  # Коэффициент потерь давления воздуха перед КС (делим на 100)
                    "Коэффициент потерь давления воздуха в КС": replace_dot_with_comma(inp9.get()),      # Коэффициент потерь давления воздуха в КС (делим на 100)
                    "Политропный КПД турбины": divide_by_100(inp10.get()),                      # Политропный КПД турбины (делим на 100)
                    "Механический КПД турбины": divide_by_100(inp11.get()),                     # Механический КПД турбины (делим на 100)
                    "Механический КПД компрессора": divide_by_100(inp12.get()),                 # Механический КПД компрессора (делим на 100)
                    "Адиабатический КПД компрессора": divide_by_100(inp13.get()),               # Адиабатический КПД компрессора (делим на 100)
                    "Степень повышения давления в компрессоре": replace_dot_with_comma(inp6.get()),      # Степень повышения давления в компрессоре
                    "Эффективная мощность": replace_dot_with_comma(inp100.get()),                        # Эффективная мощность
                    "Степень рекуперации": replace_dot_with_comma(inp14.get()),                         # Степень рекуперации
                }

                # Проходим по всем параграфам в документе
                for paragraph in doc.paragraphs:
                    # Проверяем, содержит ли параграф текст, который есть в словаре replacements
                    for key in replacements:
                        if key in paragraph.text:
     

                            # Получаем XML-структуру параграфа
                            p_xml = paragraph._element.xml

                            # Парсим XML с помощью lxml
                            p_tree = etree.fromstring(p_xml)

                            # Ищем все формулы (элементы <m:t>)
                            formulas = p_tree.xpath('.//m:t', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})

                            # Проходим по каждой формуле в параграфе
                            for formula in formulas:
                                formula_text = formula.text
                                if formula_text and "=" in formula_text:
                      

                                    # Разделяем текст на части до и после "="
                                    parts = formula_text.split("=", 1)
                                    if len(parts) == 2:
                                        # Находим маркер перед "="
                                        marker = parts[0].strip()
                   

                                        # Если маркер есть в replacements, добавляем данные после "=" без пробела
                                        if key in replacements:
                                            # Сохраняем исходный текст формулы до "="
                                            original_text = parts[0].strip()
                                            # Добавляем новые данные после "=" без пробела
                                            new_text = f"{original_text}={replacements[key]}"
                   

                                            # Обновляем текст формулы
                                            formula.text = new_text

                            # Обновляем XML-структуру параграфа
                            paragraph._element.getparent().replace(paragraph._element, p_tree)

                # Сохраняем измененный документ
                doc.save('отчет_с_регенерацией_график.docx')
    
            
           

        else:
            print("Режим расчета не выбран!")
    except Exception as e:
                print(f"Ошибка при заполнении шаблона: {e}")
    return

def fill_template_3d():
    try:
        # Определяем режим расчета
        if lang.get() == 'без регенерации':
            
                # Открываем шаблон
                doc = Document('шаблон_без_регенерации_график.docx')

                # Словарь для замены значений после "=" в формулах
                replacements = {
                    "Температура окружающей среды": replace_dot_with_comma(inp1.get()),  # Температура окружающей среды
                    "Давление окружающей среды": replace_dot_with_comma(inp2.get()),      # Давление окружающей среды
                    "Начальная температура газа перед турбиной": replace_dot_with_comma(inp3.get()),  # Начальная температура газа перед турбиной
                    "Газовая постоянная воздуха": replace_dot_with_comma(inp4.get()),      # Газовая постоянная воздуха
                    "Показатель изоэнтропы для воздуха": replace_dot_with_comma(inp15.get()),  # Показатель изоэнтропы для воздуха
                    "Показатель изоэнтропы для газа": replace_dot_with_comma(inp16.get()),     # Показатель изоэнтропы для газа
                    "Теплоемкость воздуха": replace_dot_with_comma(inp17.get()),              # Теплоемкость воздуха
                    "Теплоемкость газа": replace_dot_with_comma(inp18.get()),                # Теплоемкость газа
                    "Коэффициент потерь на входе в компрессор": replace_dot_with_comma(inp5.get()),  # Коэффициент потерь на входе в компрессор (делим на 100)
                    "Коэффициент потерь давления воздуха в вых. устройстве": replace_dot_with_comma(inp7.get()),  # Коэффициент потерь давления воздуха в вых. устройстве (делим на 100)
                    "Коэффициент потерь давления воздуха перед КС": replace_dot_with_comma(inp8.get()),  # Коэффициент потерь давления воздуха перед КС (делим на 100)
                    "Коэффициент потерь давления воздуха в КС": replace_dot_with_comma(inp9.get()),      # Коэффициент потерь давления воздуха в КС (делим на 100)
                    "Политропный КПД турбины": divide_by_100(inp10.get()),                      # Политропный КПД турбины (делим на 100)
                    "Механический КПД турбины": divide_by_100(inp11.get()),                     # Механический КПД турбины (делим на 100)
                    "Механический КПД компрессора": divide_by_100(inp12.get()),                 # Механический КПД компрессора (делим на 100)
                    "Адиабатический КПД компрессора": divide_by_100(inp13.get()),               # Адиабатический КПД компрессора (делим на 100)
                    "Степень повышения давления в компрессоре": replace_dot_with_comma(inp6.get()),      # Степень повышения давления в компрессоре
                    "Эффективная мощность": replace_dot_with_comma(inp100.get()),                        # Эффективная мощность
                }

                # Проходим по всем параграфам в документе
                for paragraph in doc.paragraphs:
                    # Проверяем, содержит ли параграф текст, который есть в словаре replacements
                    for key in replacements:
                        if key in paragraph.text:
       

                            # Получаем XML-структуру параграфа
                            p_xml = paragraph._element.xml

                            # Парсим XML с помощью lxml
                            p_tree = etree.fromstring(p_xml)

                            # Ищем все формулы (элементы <m:t>)
                            formulas = p_tree.xpath('.//m:t', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})

                            # Проходим по каждой формуле в параграфе
                            for formula in formulas:
                                formula_text = formula.text
                                if formula_text and "=" in formula_text:
                           

                                    # Разделяем текст на части до и после "="
                                    parts = formula_text.split("=", 1)
                                    if len(parts) == 2:
                                        # Находим маркер перед "="
                                        marker = parts[0].strip()
                                      

                                        # Если маркер есть в replacements, добавляем данные после "=" без пробела
                                        if key in replacements:
                                            # Сохраняем исходный текст формулы до "="
                                            original_text = parts[0].strip()
                                            # Добавляем новые данные после "=" без пробела
                                            new_text = f"{original_text}={replacements[key]}"


                                            # Обновляем текст формулы
                                            formula.text = new_text

                            # Обновляем XML-структуру параграфа
                            paragraph._element.getparent().replace(paragraph._element, p_tree)

                # Сохраняем измененный документ
                doc.save('отчет_без_регенерации_график.docx')

            
        
        elif lang.get() == 'c регенерацией':
            
                # Открываем шаблон
                doc = Document('шаблон_регенерация_график3д.docx')

                # Словарь для замены значений после "=" в формулах
                replacements = {
                    "Температура окружающей среды": replace_dot_with_comma(inp1.get()),  # Температура окружающей среды
                    "Давление окружающей среды": replace_dot_with_comma(inp2.get()),      # Давление окружающей среды
                    "Начальная температура газа перед турбиной": replace_dot_with_comma(inp3.get()),  # Начальная температура газа перед турбиной
                    "Газовая постоянная воздуха": replace_dot_with_comma(inp4.get()),      # Газовая постоянная воздуха
                    "Показатель изоэнтропы для воздуха": replace_dot_with_comma(inp15.get()),  # Показатель изоэнтропы для воздуха
                    "Показатель изоэнтропы для газа": replace_dot_with_comma(inp16.get()),     # Показатель изоэнтропы для газа
                    "Теплоемкость воздуха": replace_dot_with_comma(inp17.get()),              # Теплоемкость воздуха
                    "Теплоемкость газа": replace_dot_with_comma(inp18.get()),                # Теплоемкость газа
                    "Коэффициент потерь на входе в компрессор": replace_dot_with_comma(inp5.get()),  # Коэффициент потерь на входе в компрессор (делим на 100)
                    "Коэффициент потерь давления воздуха в вых. устройстве": replace_dot_with_comma(inp7.get()),  # Коэффициент потерь давления воздуха в вых. устройстве (делим на 100)
                    "Коэффициент потерь давления воздуха перед КС": replace_dot_with_comma(inp8.get()),  # Коэффициент потерь давления воздуха перед КС (делим на 100)
                    "Коэффициент потерь давления воздуха в КС": replace_dot_with_comma(inp9.get()),      # Коэффициент потерь давления воздуха в КС (делим на 100)
                    "Политропный КПД турбины": divide_by_100(inp10.get()),                      # Политропный КПД турбины (делим на 100)
                    "Механический КПД турбины": divide_by_100(inp11.get()),                     # Механический КПД турбины (делим на 100)
                    "Механический КПД компрессора": divide_by_100(inp12.get()),                 # Механический КПД компрессора (делим на 100)
                    "Адиабатический КПД компрессора": divide_by_100(inp13.get()),               # Адиабатический КПД компрессора (делим на 100)
                    "Степень повышения давления в компрессоре": replace_dot_with_comma(inp6.get()),      # Степень повышения давления в компрессоре
                    "Эффективная мощность": replace_dot_with_comma(inp100.get()),                        # Эффективная мощность
                    "Степень рекуперации": replace_dot_with_comma(inp14.get()),                         # Степень рекуперации
                }

                # Проходим по всем параграфам в документе
                for paragraph in doc.paragraphs:
                    # Проверяем, содержит ли параграф текст, который есть в словаре replacements
                    for key in replacements:
                        if key in paragraph.text:


                            # Получаем XML-структуру параграфа
                            p_xml = paragraph._element.xml

                            # Парсим XML с помощью lxml
                            p_tree = etree.fromstring(p_xml)

                            # Ищем все формулы (элементы <m:t>)
                            formulas = p_tree.xpath('.//m:t', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})

                            # Проходим по каждой формуле в параграфе
                            for formula in formulas:
                                formula_text = formula.text
                                if formula_text and "=" in formula_text:


                                    # Разделяем текст на части до и после "="
                                    parts = formula_text.split("=", 1)
                                    if len(parts) == 2:
                                        # Находим маркер перед "="
                                        marker = parts[0].strip()


                                        # Если маркер есть в replacements, добавляем данные после "=" без пробела
                                        if key in replacements:
                                            # Сохраняем исходный текст формулы до "="
                                            original_text = parts[0].strip()
                                            # Добавляем новые данные после "=" без пробела
                                            new_text = f"{original_text}={replacements[key]}"


                                            # Обновляем текст формулы
                                            formula.text = new_text

                            # Обновляем XML-структуру параграфа
                            paragraph._element.getparent().replace(paragraph._element, p_tree)

                # Сохраняем измененный документ
                doc.save('отчет_с_регенерацией_график3д.docx')

            
           

        else:
            print("Режим расчета не выбран!")
    except Exception as e:
                print(f"Ошибка при заполнении шаблона: {e}")
    return

def fill_template_full():
    try:
        # Определяем режим расчета
        if lang.get() == 'без регенерации':
            
                # Открываем шаблон
                doc = Document('шаблон_без_регенерации_график.docx')

                # Словарь для замены значений после "=" в формулах
                replacements = {
                    "Температура окружающей среды": replace_dot_with_comma(inp1.get()),  # Температура окружающей среды
                    "Давление окружающей среды": replace_dot_with_comma(inp2.get()),      # Давление окружающей среды
                    "Начальная температура газа перед турбиной": replace_dot_with_comma(inp3.get()),  # Начальная температура газа перед турбиной
                    "Газовая постоянная воздуха": replace_dot_with_comma(inp4.get()),      # Газовая постоянная воздуха
                    "Показатель изоэнтропы для воздуха": replace_dot_with_comma(inp15.get()),  # Показатель изоэнтропы для воздуха
                    "Показатель изоэнтропы для газа": replace_dot_with_comma(inp16.get()),     # Показатель изоэнтропы для газа
                    "Теплоемкость воздуха": replace_dot_with_comma(inp17.get()),              # Теплоемкость воздуха
                    "Теплоемкость газа": replace_dot_with_comma(inp18.get()),                # Теплоемкость газа
                    "Коэффициент потерь на входе в компрессор": replace_dot_with_comma(inp5.get()),  # Коэффициент потерь на входе в компрессор (делим на 100)
                    "Коэффициент потерь давления воздуха в вых. устройстве": replace_dot_with_comma(inp7.get()),  # Коэффициент потерь давления воздуха в вых. устройстве (делим на 100)
                    "Коэффициент потерь давления воздуха перед КС": replace_dot_with_comma(inp8.get()),  # Коэффициент потерь давления воздуха перед КС (делим на 100)
                    "Коэффициент потерь давления воздуха в КС": replace_dot_with_comma(inp9.get()),      # Коэффициент потерь давления воздуха в КС (делим на 100)
                    "Политропный КПД турбины": divide_by_100(inp10.get()),                      # Политропный КПД турбины (делим на 100)
                    "Механический КПД турбины": divide_by_100(inp11.get()),                     # Механический КПД турбины (делим на 100)
                    "Механический КПД компрессора": divide_by_100(inp12.get()),                 # Механический КПД компрессора (делим на 100)
                    "Адиабатический КПД компрессора": divide_by_100(inp13.get()),               # Адиабатический КПД компрессора (делим на 100)
                    "Степень повышения давления в компрессоре": replace_dot_with_comma(inp6.get()),      # Степень повышения давления в компрессоре
                    "Эффективная мощность": replace_dot_with_comma(inp100.get()),                        # Эффективная мощность
                }

                # Проходим по всем параграфам в документе
                for paragraph in doc.paragraphs:
                    # Проверяем, содержит ли параграф текст, который есть в словаре replacements
                    for key in replacements:
                        if key in paragraph.text:


                            # Получаем XML-структуру параграфа
                            p_xml = paragraph._element.xml

                            # Парсим XML с помощью lxml
                            p_tree = etree.fromstring(p_xml)

                            # Ищем все формулы (элементы <m:t>)
                            formulas = p_tree.xpath('.//m:t', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})

                            # Проходим по каждой формуле в параграфе
                            for formula in formulas:
                                formula_text = formula.text
                                if formula_text and "=" in formula_text:


                                    # Разделяем текст на части до и после "="
                                    parts = formula_text.split("=", 1)
                                    if len(parts) == 2:
                                        # Находим маркер перед "="
                                        marker = parts[0].strip()


                                        # Если маркер есть в replacements, добавляем данные после "=" без пробела
                                        if key in replacements:
                                            # Сохраняем исходный текст формулы до "="
                                            original_text = parts[0].strip()
                                            # Добавляем новые данные после "=" без пробела
                                            new_text = f"{original_text}={replacements[key]}"


                                            # Обновляем текст формулы
                                            formula.text = new_text

                            # Обновляем XML-структуру параграфа
                            paragraph._element.getparent().replace(paragraph._element, p_tree)

                # Сохраняем измененный документ
                doc.save('отчет_без_регенерации_график.docx')

            
        
        elif lang.get() == 'c регенерацией':
            
                # Открываем шаблон
                doc = Document('шаблон_полное.docx')

                # Словарь для замены значений после "=" в формулах
                replacements = {
                    "Температура окружающей среды": replace_dot_with_comma(inp1.get()),  # Температура окружающей среды
                    "Давление окружающей среды": replace_dot_with_comma(inp2.get()),      # Давление окружающей среды
                    "Начальная температура газа перед турбиной": replace_dot_with_comma(inp3.get()),  # Начальная температура газа перед турбиной
                    "Газовая постоянная воздуха": replace_dot_with_comma(inp4.get()),      # Газовая постоянная воздуха
                    "Показатель изоэнтропы для воздуха": replace_dot_with_comma(inp15.get()),  # Показатель изоэнтропы для воздуха
                    "Показатель изоэнтропы для газа": replace_dot_with_comma(inp16.get()),     # Показатель изоэнтропы для газа
                    "Теплоемкость воздуха": replace_dot_with_comma(inp17.get()),              # Теплоемкость воздуха
                    "Теплоемкость газа": replace_dot_with_comma(inp18.get()),                # Теплоемкость газа
                    "Коэффициент потерь на входе в компрессор": replace_dot_with_comma(inp5.get()),  # Коэффициент потерь на входе в компрессор (делим на 100)
                    "Коэффициент потерь давления воздуха в вых. устройстве": replace_dot_with_comma(inp7.get()),  # Коэффициент потерь давления воздуха в вых. устройстве (делим на 100)
                    "Коэффициент потерь давления воздуха перед КС": replace_dot_with_comma(inp8.get()),  # Коэффициент потерь давления воздуха перед КС (делим на 100)
                    "Коэффициент потерь давления воздуха в КС": replace_dot_with_comma(inp9.get()),      # Коэффициент потерь давления воздуха в КС (делим на 100)
                    "Политропный КПД турбины": divide_by_100(inp10.get()),                      # Политропный КПД турбины (делим на 100)
                    "Механический КПД турбины": divide_by_100(inp11.get()),                     # Механический КПД турбины (делим на 100)
                    "Механический КПД компрессора": divide_by_100(inp12.get()),                 # Механический КПД компрессора (делим на 100)
                    "Адиабатический КПД компрессора": divide_by_100(inp13.get()),               # Адиабатический КПД компрессора (делим на 100)
                    "Степень повышения давления в компрессоре": replace_dot_with_comma(inp6.get()),      # Степень повышения давления в компрессоре
                    "Эффективная мощность": replace_dot_with_comma(inp100.get()),                        # Эффективная мощность
                    "Степень рекуперации": replace_dot_with_comma(inp14.get()),                         # Степень рекуперации
                }

                # Проходим по всем параграфам в документе
                for paragraph in doc.paragraphs:
                    # Проверяем, содержит ли параграф текст, который есть в словаре replacements
                    for key in replacements:
                        if key in paragraph.text:

                            # Получаем XML-структуру параграфа
                            p_xml = paragraph._element.xml

                            # Парсим XML с помощью lxml
                            p_tree = etree.fromstring(p_xml)

                            # Ищем все формулы (элементы <m:t>)
                            formulas = p_tree.xpath('.//m:t', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})

                            # Проходим по каждой формуле в параграфе
                            for formula in formulas:
                                formula_text = formula.text
                                if formula_text and "=" in formula_text:

                                    # Разделяем текст на части до и после "="
                                    parts = formula_text.split("=", 1)
                                    if len(parts) == 2:
                                        # Находим маркер перед "="
                                        marker = parts[0].strip()

                                        # Если маркер есть в replacements, добавляем данные после "=" без пробела
                                        if key in replacements:
                                            # Сохраняем исходный текст формулы до "="
                                            original_text = parts[0].strip()
                                            # Добавляем новые данные после "=" без пробела
                                            new_text = f"{original_text}={replacements[key]}"

                                            # Обновляем текст формулы
                                            formula.text = new_text

                            # Обновляем XML-структуру параграфа
                            paragraph._element.getparent().replace(paragraph._element, p_tree)

                # Сохраняем измененный документ
                doc.save('отчет_с_регенерацией_полное.docx')
            
           

        else:
            print("Режим расчета не выбран!")
    except Exception as e:
                print(f"Ошибка при заполнении шаблона: {e}")
    return

def fill_calculated_values_():
    try:
        import re
        from lxml import etree

        # Определяем режим расчета
        if lang.get() == 'без регенерации':
            template_path = 'отчет_без_регенерации_график.docx'
        elif lang.get() == 'c регенерацией':
            template_path = 'отчет_с_регенерацией_график.docx'
        else:
            print("Режим расчета не выбран!")
            return

        # Функция для извлечения числа после знака '=' в тексте лейбла
        def extract_value(label_text):
            # Ищем значение после знака '=', включая отрицательные числа
            match = re.search(r"=\s*(-?[\d.,]+)", label_text)
            if match:
                # Возвращаем значение, заменяя точку на запятую
                return match.group(1).replace('.', ',')
            return ""  # Если значение не найдено, возвращаем пустую строку

        # Собираем значения из интерфейса
        calculated_values = {
            "Давление воздуха перед компрессором": extract_value(label6.cget("text")),
            "Температура воздуха перед компрессором": extract_value(label7.cget("text")),
            "Давление воздуха за компрессором": extract_value(label21.cget("text")),
            "Температура воздуха за компрессором": extract_value(label22.cget("text")),
            "Работа, соответствующая изоэнтропийному перепаду в компрессоре:": extract_value(label23.cget("text")),
            "Полезная работа в компрессоре": extract_value(label24.cget("text")),
            "Давление воздуха перед РВ": extract_value(label36.cget("text")),
            "Давление газа перед турбиной": extract_value(label50.cget("text")),
            "Давление газа за турбиной": extract_value(label25.cget("text")),
            "Степень расширения газа в турбине": extract_value(label26.cget("text")),
            "Работа, соответствующая изоэнтропийному перепаду в турбине:": extract_value(label27.cget("text")),
            "Полезная работа в турбине": extract_value(label28.cget("text")),
            "Температура газа за турбиной": extract_value(label29.cget("text")),
            "Расход воздуха через компрессор": extract_value(label30.cget("text")),
            "Температура воздуха перед камерой сгорания": extract_value(label31.cget("text")),
            "Теплота с учетом потерь в камере сгорания": extract_value(label32.cget("text")),
            "Расход теплоты": extract_value(label33.cget("text")),
            "Эффективная удельная работа": extract_value(label200.cget("text")),
            "Эффективный КПД установки": extract_value(label34.cget("text")),
            "Коэффициент полезной работы": extract_value(label35.cget("text")),
        }

        doc = Document(template_path)
        paragraphs = list(doc.paragraphs)  # Получаем список всех параграфов

        # Проходим по всем параграфам
        for i, paragraph in enumerate(paragraphs):
            # Проверяем каждый параметр из списка
            for param_name in calculated_values:
                if param_name in paragraph.text:
                    # Проверяем, есть ли следующий параграф
                    if i + 1 < len(paragraphs):
                        next_paragraph = paragraphs[i + 1]


                        # Получаем XML следующего параграфа
                        next_p_xml = next_paragraph._element.xml
                        next_p_tree = etree.fromstring(next_p_xml)

                        # Ищем формулы с '@'
                        formulas = next_p_tree.xpath(
                            './/m:t[contains(text(), "@")]',
                            namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
                        )

                        # Заменяем '@' на '= значение'
                        value = calculated_values[param_name]
                        for formula in formulas:
                            new_text = formula.text.replace('@', f'= {value}')
                            formula.text = new_text


                        # Обновляем XML следующего параграфа
                        next_paragraph._element.getparent().replace(next_paragraph._element, next_p_tree)

        # Сохраняем изменения
        doc.save(template_path)


    except Exception as e:
        print(f"Ошибка при сохранении расчетных значений: {e}")

def fill_calculated_values_3d():
    try:
        import re
        from lxml import etree

        # Определяем режим расчета
        if lang.get() == 'без регенерации':
            template_path = 'отчет_без_регенерации_график.docx'
        elif lang.get() == 'c регенерацией':
            template_path = 'отчет_с_регенерацией_график3д.docx'
        else:
            print("Режим расчета не выбран!")
            return

        # Функция для извлечения числа после знака '=' в тексте лейбла
        def extract_value(label_text):
            # Ищем значение после знака '=', включая отрицательные числа
            match = re.search(r"=\s*(-?[\d.,]+)", label_text)
            if match:
                # Возвращаем значение, заменяя точку на запятую
                return match.group(1).replace('.', ',')
            return ""  # Если значение не найдено, возвращаем пустую строку

        # Собираем значения из интерфейса
        calculated_values = {
            "Давление воздуха перед компрессором": extract_value(label6.cget("text")),
            "Температура воздуха перед компрессором": extract_value(label7.cget("text")),
            "Давление воздуха за компрессором": extract_value(label21.cget("text")),
            "Температура воздуха за компрессором": extract_value(label22.cget("text")),
            "Работа, соответствующая изоэнтропийному перепаду в компрессоре:": extract_value(label23.cget("text")),
            "Полезная работа в компрессоре": extract_value(label24.cget("text")),
            "Давление воздуха перед РВ": extract_value(label36.cget("text")),
            "Давление газа перед турбиной": extract_value(label50.cget("text")),
            "Давление газа за турбиной": extract_value(label25.cget("text")),
            "Степень расширения газа в турбине": extract_value(label26.cget("text")),
            "Работа, соответствующая изоэнтропийному перепаду в турбине:": extract_value(label27.cget("text")),
            "Полезная работа в турбине": extract_value(label28.cget("text")),
            "Температура газа за турбиной": extract_value(label29.cget("text")),
            "Расход воздуха через компрессор": extract_value(label30.cget("text")),
            "Температура воздуха перед камерой сгорания": extract_value(label31.cget("text")),
            "Теплота с учетом потерь в камере сгорания": extract_value(label32.cget("text")),
            "Расход теплоты": extract_value(label33.cget("text")),
            "Эффективная удельная работа": extract_value(label200.cget("text")),
            "Эффективный КПД установки": extract_value(label34.cget("text")),
            "Коэффициент полезной работы": extract_value(label35.cget("text")),
        }

        doc = Document(template_path)
        paragraphs = list(doc.paragraphs)  # Получаем список всех параграфов

        # Проходим по всем параграфам
        for i, paragraph in enumerate(paragraphs):
            # Проверяем каждый параметр из списка
            for param_name in calculated_values:
                if param_name in paragraph.text:
                    # Проверяем, есть ли следующий параграф
                    if i + 1 < len(paragraphs):
                        next_paragraph = paragraphs[i + 1]


                        # Получаем XML следующего параграфа
                        next_p_xml = next_paragraph._element.xml
                        next_p_tree = etree.fromstring(next_p_xml)

                        # Ищем формулы с '@'
                        formulas = next_p_tree.xpath(
                            './/m:t[contains(text(), "@")]',
                            namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
                        )

                        # Заменяем '@' на '= значение'
                        value = calculated_values[param_name]
                        for formula in formulas:
                            new_text = formula.text.replace('@', f'= {value}')
                            formula.text = new_text


                        # Обновляем XML следующего параграфа
                        next_paragraph._element.getparent().replace(next_paragraph._element, next_p_tree)

        # Сохраняем изменения
        doc.save(template_path)


    except Exception as e:
        print(f"Ошибка при сохранении расчетных значений: {e}")

def fill_calculated_values_full():
    try:
        import re
        from lxml import etree

        # Определяем режим расчета
        if lang.get() == 'без регенерации':
            template_path = 'отчет_без_регенерации_график.docx'
        elif lang.get() == 'c регенерацией':
            template_path = 'отчет_с_регенерацией_полное.docx'
        else:
            print("Режим расчета не выбран!")
            return

        # Функция для извлечения числа после знака '=' в тексте лейбла
        def extract_value(label_text):
            # Ищем значение после знака '=', включая отрицательные числа
            match = re.search(r"=\s*(-?[\d.,]+)", label_text)
            if match:
                # Возвращаем значение, заменяя точку на запятую
                return match.group(1).replace('.', ',')
            return ""  # Если значение не найдено, возвращаем пустую строку

        # Собираем значения из интерфейса
        calculated_values = {
            "Давление воздуха перед компрессором": extract_value(label6.cget("text")),
            "Температура воздуха перед компрессором": extract_value(label7.cget("text")),
            "Давление воздуха за компрессором": extract_value(label21.cget("text")),
            "Температура воздуха за компрессором": extract_value(label22.cget("text")),
            "Работа, соответствующая изоэнтропийному перепаду в компрессоре:": extract_value(label23.cget("text")),
            "Полезная работа в компрессоре": extract_value(label24.cget("text")),
            "Давление воздуха перед РВ": extract_value(label36.cget("text")),
            "Давление газа перед турбиной": extract_value(label50.cget("text")),
            "Давление газа за турбиной": extract_value(label25.cget("text")),
            "Степень расширения газа в турбине": extract_value(label26.cget("text")),
            "Работа, соответствующая изоэнтропийному перепаду в турбине:": extract_value(label27.cget("text")),
            "Полезная работа в турбине": extract_value(label28.cget("text")),
            "Температура газа за турбиной": extract_value(label29.cget("text")),
            "Расход воздуха через компрессор": extract_value(label30.cget("text")),
            "Температура воздуха перед камерой сгорания": extract_value(label31.cget("text")),
            "Теплота с учетом потерь в камере сгорания": extract_value(label32.cget("text")),
            "Расход теплоты": extract_value(label33.cget("text")),
            "Эффективная удельная работа": extract_value(label200.cget("text")),
            "Эффективный КПД установки": extract_value(label34.cget("text")),
            "Коэффициент полезной работы": extract_value(label35.cget("text")),
        }

        doc = Document(template_path)
        paragraphs = list(doc.paragraphs)  # Получаем список всех параграфов

        # Проходим по всем параграфам
        for i, paragraph in enumerate(paragraphs):
            # Проверяем каждый параметр из списка
            for param_name in calculated_values:
                if param_name in paragraph.text:
                    # Проверяем, есть ли следующий параграф
                    if i + 1 < len(paragraphs):
                        next_paragraph = paragraphs[i + 1]


                        # Получаем XML следующего параграфа
                        next_p_xml = next_paragraph._element.xml
                        next_p_tree = etree.fromstring(next_p_xml)

                        # Ищем формулы с '@'
                        formulas = next_p_tree.xpath(
                            './/m:t[contains(text(), "@")]',
                            namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
                        )

                        # Заменяем '@' на '= значение'
                        value = calculated_values[param_name]
                        for formula in formulas:
                            new_text = formula.text.replace('@', f'= {value}')
                            formula.text = new_text


                        # Обновляем XML следующего параграфа
                        next_paragraph._element.getparent().replace(next_paragraph._element, next_p_tree)

        # Сохраняем изменения
        doc.save(template_path)


    except Exception as e:
        print(f"Ошибка при сохранении расчетных значений: {e}")
#=========================================================================================================
def save_word_report_2d():
    fill_template_()
    fill_calculated_values_()
    global global_save_fig, saved_fig_eta, saved_fig_he, saved_fig_eta_he
    
    with tempfile.TemporaryDirectory() as temp_dir:
        figures = {
            '$': global_save_fig,
            '&': saved_fig_eta,
            '!': saved_fig_he,
            '^': saved_fig_eta_he
        }
        
        file_paths = {}
        for symbol, fig in figures.items():
            if fig is None:
                messagebox.showwarning("Предупреждение", f"График для символа {symbol} отсутствует.")
                continue
            try:
                path = os.path.join(temp_dir, f"{symbol}.png")
                fig.savefig(path, dpi=600, bbox_inches='tight')
                file_paths[symbol] = path
                plt.close(fig)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить график {symbol}:\n{str(e)}")
                return

        if lang.get() == 'без регенерации':
            doc = Document('отчет_без_регенерации_график.docx')
        elif lang.get() == 'c регенерацией':
            doc = Document('отчет_с_регенерацией_график.docx')
        else:
            print("Режим расчета не выбран!")
            return

        def contains_formula(paragraph):
            for run in paragraph.runs:
                xml = run._element.xml
                if 'm:oMath' in xml or 'm:oMathPara' in xml or 'm:' in xml:
                    return True
            return False

        def process_paragraph(paragraph):
            # Проверяем, содержит ли параграф только символы для замены
            symbols_only = all(c in file_paths.keys() for c in paragraph.text.strip())
            
            # Если параграф содержит формулы И НЕ состоит только из символов для замены
            if contains_formula(paragraph) and not symbols_only:
                return
                
            original_text = paragraph.text
            if not original_text:
                return
                
            symbols = re.escape(''.join(file_paths.keys()))
            pattern = re.compile(f"([{symbols}])|([^{symbols}]+)")
            
            tokens = pattern.findall(original_text)
            paragraph.clear()
            
            for token in tokens:
                symbol, text = token
                if symbol:
                    if symbol in file_paths:
                        run = paragraph.add_run()
                        try:
                            run.add_picture(file_paths[symbol], width=Cm(15), height=Cm(10))
                        except Exception as e:
                            messagebox.showerror("Ошибка", f"Ошибка вставки изображения {symbol}:\n{str(e)}")
                            raise
                    else:
                        paragraph.add_run(symbol)
                elif text:
                    paragraph.add_run(text)

        try:
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
                        for nested_table in cell.tables:
                            for nested_row in nested_table.rows:
                                for nested_cell in nested_row.cells:
                                    for nested_para in nested_cell.paragraphs:
                                        process_paragraph(nested_para)

            # Сначала сохраняем в директорию программы
            program_dir = os.path.dirname(os.path.abspath(__file__))
            local_copy_path = os.path.join(program_dir, "результат_отчет_графики2д.docx")
            doc.save(local_copy_path)

            root = Tk()
            root.withdraw()
            root.wm_attributes('-topmost', 1)
            
            output_filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Сохранить документ",
                initialfile="результат_отчет_графики2д.docx"
            )
            
            root.destroy()

            if not output_filename:
                messagebox.showinfo("Инфо", "Сохранение отменено")
                return

            try:
                doc.save(output_filename)
                messagebox.showinfo("Успех", f"Документ успешно сохранен:\n{output_filename}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
                return

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при обработке документа:\n{str(e)}")
            return
#=========================================================================================================
import os
import re
import tempfile
import shutil
def save_pdf_report_2d():
    # Создаем прогресс-бар
    progress_bar = CTkProgressBar(save_right_frame, mode='indeterminate')
    progress_bar.grid(row=8, column=0, columnspan=3, padx=20, pady=10, sticky="ew")
    progress_bar.start()
    
    try:
        # Подготовка данных
        fill_template_()
        fill_calculated_values_()
        global global_save_fig, saved_fig_eta, saved_fig_he, saved_fig_eta_he
        
        # Создаем временную директорию для всех временных файлов
        with tempfile.TemporaryDirectory() as temp_dir:
            # Сохраняем все графики во временную директорию
            figures = {
                '$': global_save_fig,
                '&': saved_fig_eta,
                '!': saved_fig_he,
                '^': saved_fig_eta_he
            }
            
            file_paths = {}
            for symbol, fig in figures.items():
                if fig is None:
                    messagebox.showwarning("Предупреждение", f"График для символа {symbol} отсутствует.")
                    continue
                try:
                    path = os.path.join(temp_dir, f"{symbol}.png")
                    fig.savefig(path, dpi=600, bbox_inches='tight')
                    file_paths[symbol] = path
                    plt.close(fig)
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Не удалось сохранить график {symbol}:\n{str(e)}")
                    progress_bar.stop()
                    progress_bar.destroy()
                    return

            # Выбираем шаблон документа
            template_name = 'отчет_без_регенерации_график.docx' if lang.get() == 'без регенерации' else 'отчет_с_регенерацией_график.docx'
            if not os.path.exists(template_name):
                messagebox.showerror("Ошибка", f"Шаблон документа {template_name} не найден!")
                progress_bar.stop()
                progress_bar.destroy()
                return

            # Создаем временный DOCX файл
            temp_docx_path = os.path.join(temp_dir, "temp_report.docx")
            
            # Обрабатываем документ
            doc = Document(template_name)
            
            def contains_formula(paragraph):
                for run in paragraph.runs:
                    xml = run._element.xml
                    if 'm:oMath' in xml or 'm:oMathPara' in xml or 'm:' in xml:
                        return True
                return False

            def process_paragraph(paragraph):
                symbols_only = all(c in file_paths.keys() for c in paragraph.text.strip())
                
                if contains_formula(paragraph) and not symbols_only:
                    return
                    
                original_text = paragraph.text
                if not original_text:
                    return
                    
                symbols = re.escape(''.join(file_paths.keys()))
                pattern = re.compile(f"([{symbols}])|([^{symbols}]+)")
                
                tokens = pattern.findall(original_text)
                paragraph.clear()
                
                for token in tokens:
                    symbol, text = token
                    if symbol:
                        if symbol in file_paths:
                            run = paragraph.add_run()
                            try:
                                run.add_picture(file_paths[symbol], width=Cm(15), height=Cm(10))
                            except Exception as e:
                                messagebox.showerror("Ошибка", f"Ошибка вставки изображения {symbol}:\n{str(e)}")
                                raise
                        else:
                            paragraph.add_run(symbol)
                    elif text:
                        paragraph.add_run(text)

            try:
                # Обработка всех параграфов и таблиц
                for paragraph in doc.paragraphs:
                    process_paragraph(paragraph)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                process_paragraph(paragraph)
                            for nested_table in cell.tables:
                                for nested_row in nested_table.rows:
                                    for nested_cell in nested_row.cells:
                                        for nested_para in nested_cell.paragraphs:
                                            process_paragraph(nested_para)

                # Сохраняем временный DOCX
                doc.save(temp_docx_path)

                # Создаем временный PDF
                temp_pdf_path = os.path.join(temp_dir, "temp_report.pdf")
                
                # Конвертируем в PDF с подавлением вывода в консоль
                try:
                    # Перенаправляем stdout и stderr временно
                    import sys
                    from io import StringIO
                    
                    old_stdout = sys.stdout
                    old_stderr = sys.stderr
                    sys.stdout = StringIO()
                    sys.stderr = StringIO()
                    
                    try:
                        convert(temp_docx_path, temp_pdf_path)
                    except Exception as e:
                        messagebox.showerror("Ошибка", f"Ошибка конвертации в PDF:\n{str(e)}")
                        return
                    finally:
                        # Восстанавливаем stdout и stderr
                        sys.stdout = old_stdout
                        sys.stderr = old_stderr
                        
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Ошибка конвертации в PDF:\n{str(e)}")
                    progress_bar.stop()
                    progress_bar.destroy()
                    return

                # Запрашиваем путь для сохранения у пользователя
                root = Tk()
                root.withdraw()
                root.wm_attributes('-topmost', 1)
                
                pdf_output_path = filedialog.asksaveasfilename(
                    defaultextension=".pdf",
                    filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
                    title="Сохранить PDF документ",
                    initialfile="результат_отчет_графики2д.pdf"
                )
                
                root.destroy()

                if not pdf_output_path:
                    messagebox.showinfo("Инфо", "Сохранение PDF отменено")
                    progress_bar.stop()
                    progress_bar.destroy()
                    return

                # Копируем временный PDF в указанное место
                try:
                    shutil.copy2(temp_pdf_path, pdf_output_path)
                    messagebox.showinfo("Успех", f"PDF документ успешно сохранен:\n{pdf_output_path}")
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Ошибка при сохранении PDF файла:\n{str(e)}")
                    return

            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при обработке документа:\n{str(e)}")
                return

    except Exception as e:
        messagebox.showerror("Ошибка", f"Неожиданная ошибка:\n{str(e)}")
        return
    finally:
        # Останавливаем и удаляем прогресс-бар в любом случае
        progress_bar.stop()
        progress_bar.destroy()
#=========================================================================================================
def save_word_report_3d():
    fill_template_3d()
    fill_calculated_values_3d()
    global fig_phi_light, fig_eta_light, fig_he_light, fig_eta_mu_he_light

    # Создаем временную директорию для сохранения изображений
    with tempfile.TemporaryDirectory() as temp_dir:
        # Сохраняем все фигуры в файлы
        figures = {
            '#': fig_phi_light,
            '{': fig_eta_light,
            '[': fig_he_light,
            ']': fig_eta_mu_he_light
        }
        
        file_paths = {}
        for symbol, fig in figures.items():
            if fig is None:
                messagebox.showwarning("Предупреждение", f"График для символа {symbol} отсутствует.")
                continue
            try:
                path = os.path.join(temp_dir, f"{symbol}.png")
                fig.savefig(path, dpi=600, bbox_inches='tight')
                file_paths[symbol] = path
                plt.close(fig)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить график {symbol}:\n{str(e)}")
                return

        doc = Document('отчет_с_регенерацией_график3д.docx')

        def contains_formula(paragraph):
            for run in paragraph.runs:
                xml = run._element.xml
                if 'm:oMath' in xml or 'm:oMathPara' in xml or 'm:' in xml:
                    return True
            return False

        def process_paragraph(paragraph):
            # Проверяем, содержит ли параграф только символы для замены
            symbols_only = all(c in file_paths.keys() for c in paragraph.text.strip())
            
            # Если параграф содержит формулы И НЕ состоит только из символов для замены
            if contains_formula(paragraph) and not symbols_only:
                return
                
            original_text = paragraph.text
            if not original_text:
                return
                
            symbols = re.escape(''.join(file_paths.keys()))
            pattern = re.compile(f"([{symbols}])|([^{symbols}]+)")
            
            tokens = pattern.findall(original_text)
            paragraph.clear()
            
            for token in tokens:
                symbol, text = token
                if symbol:
                    if symbol in file_paths:
                        run = paragraph.add_run()
                        try:
                            run.add_picture(file_paths[symbol], width=Cm(17), height=Cm(14))
                        except Exception as e:
                            messagebox.showerror("Ошибка", f"Ошибка вставки изображения {symbol}:\n{str(e)}")
                            raise
                    else:
                        paragraph.add_run(symbol)
                elif text:
                    paragraph.add_run(text)

        try:
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
                        for nested_table in cell.tables:
                            for nested_row in nested_table.rows:
                                for nested_cell in nested_row.cells:
                                    for nested_para in nested_cell.paragraphs:
                                        process_paragraph(nested_para)
            
            # Сначала сохраняем в директорию программы
            program_dir = os.path.dirname(os.path.abspath(__file__))
            local_copy_path = os.path.join(program_dir, "результат_отчет_графики3д.docx")
            doc.save(local_copy_path)

            root = Tk()
            root.withdraw()
            root.wm_attributes('-topmost', 1)
            
            output_filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Сохранить документ",
                initialfile="результат_отчет_графики3д.docx"
            )
            
            root.destroy()

            if not output_filename:
                messagebox.showinfo("Инфо", "Сохранение отменено")
                return

            try:
                doc.save(output_filename)
                messagebox.showinfo("Успех", f"Документ успешно сохранен:\n{output_filename}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
                return

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при обработке документа:\n{str(e)}")
            return
#=========================================================================================================
def save_pdf_report_3d():
    # Создаем прогресс-бар
    progress_bar = CTkProgressBar(save_right_frame, mode='indeterminate')
    progress_bar.grid(row=8, column=0, columnspan=3, padx=20, pady=10, sticky="ew")
    progress_bar.start()
    
    try:
        # Подготовка данных
        fill_template_3d()
        fill_calculated_values_3d()
        global fig_phi_light, fig_eta_light, fig_he_light, fig_eta_mu_he_light

        # Создаем временную директорию для всех временных файлов
        with tempfile.TemporaryDirectory() as temp_dir:
            # Сохраняем все графики во временную директорию
            figures = {
                '#': fig_phi_light,
                '{': fig_eta_light,
                '[': fig_he_light,
                ']': fig_eta_mu_he_light
            }
            
            file_paths = {}
            for symbol, fig in figures.items():
                if fig is None:
                    messagebox.showwarning("Предупреждение", f"График для символа {symbol} отсутствует.")
                    continue
                try:
                    path = os.path.join(temp_dir, f"{symbol}.png")
                    fig.savefig(path, dpi=600, bbox_inches='tight')
                    file_paths[symbol] = path
                    plt.close(fig)
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Не удалось сохранить график {symbol}:\n{str(e)}")
                    progress_bar.stop()
                    progress_bar.destroy()
                    return

            # Проверяем наличие шаблона документа
            template_name = 'отчет_с_регенерацией_график3д.docx'
            if not os.path.exists(template_name):
                messagebox.showerror("Ошибка", f"Шаблон документа {template_name} не найден!")
                progress_bar.stop()
                progress_bar.destroy()
                return

            # Создаем временный DOCX файл
            temp_docx_path = os.path.join(temp_dir, "temp_report.docx")
            
            # Обрабатываем документ
            doc = Document(template_name)
            
            def contains_formula(paragraph):
                for run in paragraph.runs:
                    xml = run._element.xml
                    if 'm:oMath' in xml or 'm:oMathPara' in xml or 'm:' in xml:
                        return True
                return False

            def process_paragraph(paragraph):
                symbols_only = all(c in file_paths.keys() for c in paragraph.text.strip())
                
                if contains_formula(paragraph) and not symbols_only:
                    return
                    
                original_text = paragraph.text
                if not original_text:
                    return
                    
                symbols = re.escape(''.join(file_paths.keys()))
                pattern = re.compile(f"([{symbols}])|([^{symbols}]+)")
                
                tokens = pattern.findall(original_text)
                paragraph.clear()
                
                for token in tokens:
                    symbol, text = token
                    if symbol:
                        if symbol in file_paths:
                            run = paragraph.add_run()
                            try:
                                run.add_picture(file_paths[symbol], width=Cm(17), height=Cm(14))
                            except Exception as e:
                                messagebox.showerror("Ошибка", f"Ошибка вставки изображения {symbol}:\n{str(e)}")
                                raise
                        else:
                            paragraph.add_run(symbol)
                    elif text:
                        paragraph.add_run(text)

            try:
                # Обработка всех параграфов и таблиц
                for paragraph in doc.paragraphs:
                    process_paragraph(paragraph)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                process_paragraph(paragraph)
                            for nested_table in cell.tables:
                                for nested_row in nested_table.rows:
                                    for nested_cell in nested_row.cells:
                                        for nested_para in nested_cell.paragraphs:
                                            process_paragraph(nested_para)

                # Сохраняем временный DOCX
                doc.save(temp_docx_path)

                # Создаем временный PDF
                temp_pdf_path = os.path.join(temp_dir, "temp_report.pdf")
                
                # Конвертируем в PDF с подавлением вывода в консоль
                try:
                    # Перенаправляем stdout и stderr временно
                    import sys
                    from io import StringIO
                    
                    old_stdout = sys.stdout
                    old_stderr = sys.stderr
                    sys.stdout = StringIO()
                    sys.stderr = StringIO()
                    
                    try:
                        convert(temp_docx_path, temp_pdf_path)
                    except Exception as e:
                        messagebox.showerror("Ошибка", f"Ошибка конвертации в PDF:\n{str(e)}")
                        return
                    finally:
                        # Восстанавливаем stdout и stderr
                        sys.stdout = old_stdout
                        sys.stderr = old_stderr
                        
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Ошибка конвертации в PDF:\n{str(e)}")
                    progress_bar.stop()
                    progress_bar.destroy()
                    return

                # Запрашиваем путь для сохранения у пользователя
                root = Tk()
                root.withdraw()
                root.wm_attributes('-topmost', 1)
                
                pdf_output_path = filedialog.asksaveasfilename(
                    defaultextension=".pdf",
                    filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
                    title="Сохранить PDF документ",
                    initialfile="результат_отчет_графики3д.pdf"
                )
                
                root.destroy()

                if not pdf_output_path:
                    messagebox.showinfo("Инфо", "Сохранение PDF отменено")
                    progress_bar.stop()
                    progress_bar.destroy()
                    return

                # Копируем временный PDF в указанное место
                try:
                    shutil.copy(temp_pdf_path, pdf_output_path)
                    messagebox.showinfo("Успех", f"PDF документ успешно сохранен:\n{pdf_output_path}")
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Ошибка при сохранении PDF файла:\n{str(e)}")
                    return

            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при обработке документа:\n{str(e)}")
                return

    except Exception as e:
        messagebox.showerror("Ошибка", f"Неожиданная ошибка:\n{str(e)}")
        return
    finally:
        # Останавливаем и удаляем прогресс-бар в любом случае
        progress_bar.stop()
        progress_bar.destroy()
#=========================================================================================================
from docx.shared import Cm
import matplotlib.pyplot as plt

def save_pdf_report_full():
    # Создаем прогресс-бар
    progress_bar = CTkProgressBar(save_right_frame, mode='indeterminate')
    progress_bar.grid(row=8, column=0, columnspan=3, padx=20, pady=10, sticky="ew")
    progress_bar.start()
    
    try:
        # Подготовка данных
        fill_template_full()
        fill_calculated_values_full()
        global fig_phi_light, fig_eta_light, fig_he_light, fig_eta_mu_he_light
        global global_save_fig, saved_fig_eta, saved_fig_he, saved_fig_eta_he

        # Создаем временную директорию для всех файлов
        with tempfile.TemporaryDirectory() as temp_dir:
            # Определяем все графики
            figures = {
                '#': (fig_phi_light, '3d'),
                '{': (fig_eta_light, '3d'),
                '[': (fig_he_light, '3d'),
                ']': (fig_eta_mu_he_light, '3d'),
                '$': (global_save_fig, '2d'),
                '&': (saved_fig_eta, '2d'),
                '!': (saved_fig_he, '2d'),
                '^': (saved_fig_eta_he, '2d')
            }
            
            file_paths = {}
            
            # Сохраняем все графики
            for symbol, (fig, fig_type) in figures.items():
                if fig is None:
                    messagebox.showwarning("Предупреждение", 
                        f"График {fig_type.upper()} для символа {symbol} отсутствует.")
                    continue
                try:
                    path = os.path.join(temp_dir, f"{symbol}.png")
                    fig.savefig(path, dpi=600, bbox_inches='tight')
                    file_paths[symbol] = (path, fig_type)
                    plt.close(fig)
                except Exception as e:
                    messagebox.showerror("Ошибка", 
                        f"Не удалось сохранить {fig_type.upper()} график {symbol}:\n{str(e)}")
                    progress_bar.stop()
                    progress_bar.destroy()
                    return

            # Проверяем наличие шаблона
            template_name = 'отчет_с_регенерацией_полное.docx'
            if not os.path.exists(template_name):
                messagebox.showerror("Ошибка", f"Шаблон документа {template_name} не найден!")
                progress_bar.stop()
                progress_bar.destroy()
                return

            # Создаем временный DOCX
            temp_docx_path = os.path.join(temp_dir, "temp_report.docx")
            
            # Обрабатываем документ
            doc = Document(template_name)
            
            def contains_formula(paragraph):
                for run in paragraph.runs:
                    xml = run._element.xml
                    if 'm:oMath' in xml or 'm:oMathPara' in xml or 'm:' in xml:
                        return True
                return False

            def process_paragraph(paragraph):
                symbols = ''.join(file_paths.keys())
                symbols_only = all(c in symbols for c in paragraph.text.strip())
                
                if contains_formula(paragraph) and not symbols_only:
                    return
                    
                original_text = paragraph.text
                if not original_text:
                    return
                    
                escaped_symbols = re.escape(symbols)
                pattern = re.compile(f"([{escaped_symbols}])|([^{escaped_symbols}]+)")
                
                tokens = pattern.findall(original_text)
                paragraph.clear()
                
                for token in tokens:
                    symbol, text = token
                    if symbol:
                        if symbol in file_paths:
                            path, fig_type = file_paths[symbol]
                            try:
                                if not os.path.exists(path):
                                    messagebox.showwarning("Предупреждение", 
                                        f"Файл изображения {symbol} не найден")
                                    continue
                                
                                run = paragraph.add_run()
                                # Разные размеры для 2D и 3D графиков
                                if fig_type == '3d':
                                    run.add_picture(path, width=Cm(17), height=Cm(14))
                                else:  # 2D
                                    run.add_picture(path, width=Cm(15), height=Cm(10))
                            except Exception as e:
                                messagebox.showerror("Ошибка", 
                                    f"Ошибка вставки изображения {symbol}:\n{str(e)}")
                                raise
                        else:
                            paragraph.add_run(symbol)
                    elif text:
                        paragraph.add_run(text)

            try:
                # Обработка всего документа
                for paragraph in doc.paragraphs:
                    process_paragraph(paragraph)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                process_paragraph(paragraph)
                            for nested_table in cell.tables:
                                for nested_row in nested_table.rows:
                                    for nested_cell in nested_row.cells:
                                        for nested_para in nested_cell.paragraphs:
                                            process_paragraph(nested_para)

                # Сохраняем временный DOCX
                doc.save(temp_docx_path)

                # Создаем временный PDF
                temp_pdf_path = os.path.join(temp_dir, "temp_report.pdf")
                
                # Конвертируем в PDF с подавлением вывода в консоль
                try:
                    # Перенаправляем stdout и stderr временно
                    import sys
                    from io import StringIO
                    
                    old_stdout = sys.stdout
                    old_stderr = sys.stderr
                    sys.stdout = StringIO()
                    sys.stderr = StringIO()
                    
                    try:
                        convert(temp_docx_path, temp_pdf_path)
                    except Exception as e:
                        messagebox.showerror("Ошибка", f"Ошибка конвертации в PDF:\n{str(e)}")
                        return
                    finally:
                        # Восстанавливаем stdout и stderr
                        sys.stdout = old_stdout
                        sys.stderr = old_stderr
                        
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Ошибка конвертации в PDF:\n{str(e)}")
                    progress_bar.stop()
                    progress_bar.destroy()
                    return

                # Запрашиваем путь сохранения у пользователя
                root = Tk()
                root.withdraw()
                root.wm_attributes('-topmost', 1)
                
                pdf_output_path = filedialog.asksaveasfilename(
                    defaultextension=".pdf",
                    filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
                    title="Сохранить полный отчет PDF",
                    initialfile="результат_отчет_полное.pdf"
                )
                
                root.destroy()

                if not pdf_output_path:
                    messagebox.showinfo("Инфо", "Сохранение PDF отменено")
                    progress_bar.stop()
                    progress_bar.destroy()
                    return

                # Копируем PDF в указанное место
                try:
                    shutil.copy(temp_pdf_path, pdf_output_path)
                    messagebox.showinfo("Успех", 
                        f"Полный отчет успешно сохранен:\n{pdf_output_path}")
                except Exception as e:
                    messagebox.showerror("Ошибка", 
                        f"Ошибка при сохранении PDF файла:\n{str(e)}")
                    return

            except Exception as e:
                messagebox.showerror("Ошибка", 
                    f"Ошибка при обработке документа:\n{str(e)}")
                return

    except Exception as e:
        messagebox.showerror("Ошибка", 
            f"Неожиданная ошибка при создании отчета:\n{str(e)}")
        return
    finally:
        # Останавливаем и удаляем прогресс-бар в любом случае
        progress_bar.stop()
        progress_bar.destroy()
#=========================================================================================================
#=========================================================================================================
def save_word_report_full():
    fill_template_full()
    fill_calculated_values_full()
 
    global fig_phi_light, fig_eta_light, fig_he_light, fig_eta_mu_he_light
    global global_save_fig, saved_fig_eta, saved_fig_he, saved_fig_eta_he
    
    # Создаем временную директорию для сохранения изображений
    with tempfile.TemporaryDirectory() as temp_dir:
        # Сохраняем все фигуры в файлы
        figures_3d = {
            '#': fig_phi_light,
            '{': fig_eta_light,
            '[': fig_he_light,
            ']': fig_eta_mu_he_light
        }
        
        figures_2d = {
            '$': global_save_fig,
            '&': saved_fig_eta,
            '!': saved_fig_he,
            '^': saved_fig_eta_he
        }
        
        file_paths = {}
        
        # Сохраняем 3D графики
        for symbol, fig in figures_3d.items():
            if fig is None:
                messagebox.showwarning("Предупреждение", f"3D график для символа {symbol} отсутствует.")
                continue
            try:
                path = os.path.join(temp_dir, f"{symbol}.png")
                fig.savefig(path, dpi=600, bbox_inches='tight')
                file_paths[symbol] = (path, '3d')  # Помечаем как 3D
                plt.close(fig)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить 3D график {symbol}:\n{str(e)}")
                return
        
        # Сохраняем 2D графики
        for symbol, fig in figures_2d.items():
            if fig is None:
                messagebox.showwarning("Предупреждение", f"2D график для символа {symbol} отсутствует.")
                continue
            try:
                path = os.path.join(temp_dir, f"{symbol}.png")
                fig.savefig(path, dpi=600, bbox_inches='tight')
                file_paths[symbol] = (path, '2d')  # Явно указываем тип '2d'
                plt.close(fig)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить 2D график {symbol}:\n{str(e)}")
                return

        # Загрузка шаблона документа
        try:
            doc = Document('отчет_с_регенерацией_полное.docx')
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть шаблон документа:\n{str(e)}")
            return

        def contains_formula(paragraph):
            for run in paragraph.runs:
                xml = run._element.xml
                if 'm:oMath' in xml or 'm:oMathPara' in xml or 'm:' in xml:
                    return True
            return False

        def process_paragraph(paragraph):
            # Проверяем, содержит ли параграф только символы для замены
            symbols = ''.join(file_paths.keys())
            symbols_only = all(c in symbols for c in paragraph.text.strip())
    
            # Если параграф содержит формулы И НЕ состоит только из символов для замены
            if contains_formula(paragraph) and not symbols_only:
                return
        
            original_text = paragraph.text
            if not original_text:
                return
        
            # Экранируем специальные символы для regex
            escaped_symbols = re.escape(symbols)
            pattern = re.compile(f"([{escaped_symbols}])|([^{escaped_symbols}]+)")
    
            tokens = pattern.findall(original_text)
            paragraph.clear()
    
            for token in tokens:
                symbol, text = token
                if symbol:
                    if symbol in file_paths:
                        path, fig_type = file_paths[symbol]  # Получаем путь и тип графика
                        try:
                            # Проверяем существование файла
                            if not os.path.exists(path):
                                messagebox.showwarning("Предупреждение", f"Файл изображения {symbol} не найден")
                                continue
                        
                            run = paragraph.add_run()
                    
                            # Устанавливаем разные размеры для 2D и 3D графиков
                            if fig_type == '3d':
                                run.add_picture(path, width=Cm(17), height=Cm(14))
                            else:  # 2D графики
                                run.add_picture(path, width=Cm(15), height=Cm(10))  # Меньшая высота для 2D
                        except Exception as e:
                            messagebox.showerror("Ошибка", f"Ошибка вставки изображения {symbol}:\n{str(e)}")
                            raise
                    else:
                        paragraph.add_run(symbol)
                elif text:
                    paragraph.add_run(text)

        try:
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
                        for nested_table in cell.tables:
                            for nested_row in nested_table.rows:
                                for nested_cell in nested_row.cells:
                                    for nested_para in nested_cell.paragraphs:
                                        process_paragraph(nested_para)
            
            # Сначала сохраняем в директорию программы
            program_dir = os.path.dirname(os.path.abspath(__file__))
            local_copy_path = os.path.join(program_dir, "результат_отчет_полное.docx")
            doc.save(local_copy_path)

            root = Tk()
            root.withdraw()
            root.wm_attributes('-topmost', 1)
            
            output_filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Сохранить документ",
                initialfile="результат_отчет_полное.docx"
            )
            
            root.destroy()

            if not output_filename:
                messagebox.showinfo("Инфо", "Сохранение отменено")
                return

            try:
                doc.save(output_filename)
                messagebox.showinfo("Успех", f"Документ успешно сохранен:\n{output_filename}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
                return

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при обработке документа:\n{str(e)}")
            return
#=========================================================================================================
def save_plot_as_pdf_3d_all():
    """Сохраняет все доступные 3D графики в многостраничный PDF через диалоговое окно."""
    global fig_phi_light, fig_eta_light, fig_he_light, fig_eta_mu_he_light

    # Собираем все фигуры в список и фильтруем пустые значения
    figures = [
        fig_phi_light,
        fig_eta_light,
        fig_he_light,
        fig_eta_mu_he_light
    ]
    figures = [fig for fig in figures if fig is not None]

    # Проверка наличия хотя бы одной фигуры
    if not figures:
        messagebox.showwarning("Предупреждение", "Нет доступных графиков для сохранения!")
        return

    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()

    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить все 3D графики как PDF",
        initialfile="результат_графики_3D.pdf"
    )

    # Закрываем Tkinter окно
    root.destroy()

    # Сохраняем если путь получен
    if file_path:
        try:
            with PdfPages(file_path) as pdf:
                for i, fig in enumerate(figures, 1):
                    pdf.savefig(fig, bbox_inches='tight')
                    print(f"График {i}/{len(figures)} успешно добавлен в PDF")
            
            messagebox.showinfo("Успех", f"Все графики ({len(figures)} шт.) успешно сохранены в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
    else:
        messagebox.showinfo("Инфо", "Сохранение отменено пользователем.")
#=============================================================================================================================================
def save_plot_all_plots_pdf():
    global global_save_fig, saved_fig_eta, saved_fig_he, saved_fig_eta_he, fig_phi_light, fig_eta_light, fig_he_light, fig_eta_mu_he_light

    # Собираем все фигуры в список и фильтруем пустые значения
    figures = [
        global_save_fig,
        saved_fig_eta,
        saved_fig_he,
        saved_fig_eta_he,
        fig_phi_light,
        fig_eta_light,
        fig_he_light,
        fig_eta_mu_he_light
    ]
    figures = [fig for fig in figures if fig is not None]

    # Проверка наличия хотя бы одной фигуры
    if not figures:
        messagebox.showwarning("Предупреждение", "Нет доступных графиков для сохранения!")
        return

    # Создаем скрытое окно Tkinter
    root = Tk()
    root.withdraw()

    # Запрашиваем путь для сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        title="Сохранить все графики как PDF",
        initialfile="результат_графики_2D_и_3D.pdf"
    )

    # Закрываем Tkinter окно
    root.destroy()

    # Сохраняем если путь получен
    if file_path:
        try:
            with PdfPages(file_path) as pdf:
                for i, fig in enumerate(figures, 1):
                    pdf.savefig(fig, bbox_inches='tight')
                    print(f"График {i}/{len(figures)} успешно добавлен в PDF")
            
            messagebox.showinfo("Успех", f"Все графики ({len(figures)} шт.) успешно сохранены в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
    else:
        messagebox.showinfo("Инфо", "Сохранение отменено пользователем.")
#=============================================================================================================================================

update_label_positions()  

class TabManager:
    def __init__(self):
        self.current_tab = None
        self.tabs = {
            'calc': {
                'frame': frame_calc,
                'btn': calc_btn,
                'icon': calc_icon,
                'active_icon': calc_white_icon,
                'name': 'Вариантный расчет',
                'entries': [
                    (inp1, get1),
                    (inp2, get2),
                    (inp3, get3),
                    (inp4, get4),
                    (inp15, get15),
                    (inp16, get16),
                    (inp17, get17),
                    (inp18, get18),
                    (inp5, get5),
                    (inp7, get7),
                    (inp8, get8),
                    (inp9, get9),
                    (inp10, get10),
                    (inp11, get11),
                    (inp12, get12),
                    (inp13, get13),
                    (inp6, get6),
                    (inp100, get100),
                    (inp14, get14)
                ]
            },
            'grafik': {
                'frame': frame_grafik,
                'btn': grafik_btn,
                'icon': grafik_icon,
                'active_icon': grafik_white_icon,
                'name': 'Построение графиков',
                'entries': [
                    (entry_min_phi, plot_phi_vs_pk),
                    (entry_max_phi, plot_phi_vs_pk),
                    (entry_step_phi, plot_phi_vs_pk),
                    (entry_min_eta, plot_eta_vs_pk),
                    (entry_max_eta, plot_eta_vs_pk),
                    (entry_step_eta, plot_eta_vs_pk),
                    (entry_min_he, plot_he_vs_pk),
                    (entry_max_he, plot_he_vs_pk),
                    (entry_step_he, plot_he_vs_pk),
                    (entry_min_eta_he, plot_eta_he_vs_pk),
                    (entry_max_eta_he, plot_eta_he_vs_pk),
                    (entry_step_eta_he, plot_eta_he_vs_pk)
                ]
            },
            'settings': {
                'frame': frame_settings,
                'btn': settings_btn,
                'icon': settings_icon,
                'active_icon': settings_white_icon,
                'name': 'Настройки',
                'entries': []
            },
            'save': {
                'frame': frame_save,
                'btn': save_btn,
                'icon': save_menu_icon,
                'active_icon': save_menu_white_icon,
                'name': 'Сохранение',
                'entries': []
            }
        }
        
        self.setup_tabs()
        self.setup_bindings()
    
    def setup_tabs(self):
        # Привязка команд к кнопкам
        calc_btn.configure(command=lambda: self.switch_tab('calc'))
        grafik_btn.configure(command=lambda: self.switch_tab('grafik'))
        settings_btn.configure(command=lambda: self.switch_tab('settings'))
        
        # Показываем вкладку по умолчанию
        self.switch_tab('calc')
    
    def switch_tab(self, tab_name):
        # Скрываем текущую вкладку
        if self.current_tab:
            self.tabs[self.current_tab]['frame'].grid_forget()
            self.update_button_style(self.current_tab, active=False)
        
        # Показываем новую вкладку
        tab = self.tabs[tab_name]
        tab['frame'].grid(row=0, column=0, sticky="nsew")
        self.update_button_style(tab_name, active=True)
        self.current_tab = tab_name
        
        # Обновляем заголовок окна
        root.title(f'Расчет газовой турбины - {tab["name"]}')
        
        # Устанавливаем фокус на первое поле вкладки
        if tab['entries']:
            tab['entries'][0][0].focus()
    
    def update_button_style(self, tab_name, active):
        tab = self.tabs[tab_name]
        btn = tab['btn']
        current_width = left_frame.cget("width")
        
        if active:
            btn.configure(text_color="white", image=tab['active_icon'])
            btn.active = True
        else:
            current_theme = ctk.get_appearance_mode()
            btn.configure(text_color="#6a6a6a" if current_theme == "Dark" else "#a0a4ac", 
                         image=tab['icon'])
            btn.active = False
        
        # Обновление текста/иконки в зависимости от ширины меню
        if current_width == 70:
            btn.configure(text="", anchor="c")
        else:
            btn.configure(text=tab['name'] if active else tab['name'], anchor="w")
    
    def setup_bindings(self):
        # Глобальные бинды для навигации
        root.bind("<Down>", lambda e: self.navigate(1))
        root.bind("<Up>", lambda e: self.navigate(-1))
        
        # Бинды для полей ввода
        for tab in self.tabs.values():
            for i, (entry, func) in enumerate(tab['entries']):
                if hasattr(entry, '_entry'):  # Для CTkEntry
                    entry._entry.bind("<FocusIn>", lambda e, idx=i, t=tab['name']: self.set_current_index(idx, t))
                    entry._entry.bind("<Return>", lambda e, f=func: f())
                else:  # Для обычных виджетов
                    entry.bind("<FocusIn>", lambda e, idx=i, t=tab['name']: self.set_current_index(idx, t))
                    entry.bind("<Return>", lambda e, f=func: f())
    
    def set_current_index(self, index, tab_name):
        """Устанавливает текущий индекс для активной вкладки"""
        for key, tab in self.tabs.items():
            if tab['name'] == tab_name:
                self.current_tab = key
                self.current_index = index
                break
    
    def navigate(self, direction):
        """Навигация с учетом вложенности right_content_frame в CTkScrollableFrame"""
        if self.current_tab not in self.tabs:
            return

        tab = self.tabs[self.current_tab]
        entries = tab['entries']
        if not entries:
            return

        new_index = self.current_index + direction
        if not 0 <= new_index < len(entries):
            return

        # Устанавливаем фокус на элемент
        entry = entries[new_index][0]
        entry.focus()
        self.current_index = new_index

        # Получаем родительский CTkScrollableFrame
        scrollable_frame = self._find_parent_scrollable(entry)
        if not scrollable_frame or not hasattr(scrollable_frame, '_parent_canvas'):
            print("Не найден родительский CTkScrollableFrame")
            return

        canvas = scrollable_frame._parent_canvas
        scroll_frame = scrollable_frame

        # Принудительное обновление геометрии
        scroll_frame.update_idletasks()
        canvas.update_idletasks()

        # Координаты entry относительно CTkScrollableFrame
        entry_y = entry.winfo_y() + right_content_frame.winfo_y()  # Учитываем позицию right_content_frame
        
        # Вычисляем видимую область
        visible_top = canvas.canvasy(0)
        visible_bottom = canvas.canvasy(canvas.winfo_height())
        entry_height = entry.winfo_height()

        # Увеличенный отступ для прокрутки 
        scroll_margin = 270

        # Проверяем видимость элемента с учетом увеличенного отступа
        if entry_y + entry_height > visible_bottom - scroll_margin:
            # Прокручиваем вниз с учетом отступа
            scroll_to = (entry_y + entry_height - canvas.winfo_height() + scroll_margin) / scroll_frame.winfo_height()
            canvas.yview_moveto(max(0, min(1, scroll_to)))
        elif entry_y < visible_top + scroll_margin:
            # Прокручиваем вверх с учетом отступа
            scroll_to = (entry_y - scroll_margin) / scroll_frame.winfo_height()
            canvas.yview_moveto(max(0, min(1, scroll_to)))

    def _find_parent_scrollable(self, widget):
        """Рекурсивно ищет родительский CTkScrollableFrame"""
        parent = widget.master
        while parent:
            if hasattr(parent, '_parent_canvas'):  # Это CTkScrollableFrame
                return parent
            parent = parent.master
        return None

# Инициализация менеджера вкладок в конце программы
tab_manager = TabManager()

# Удаляем старые обработчики, если они были
calc_btn.configure(command=lambda: tab_manager.switch_tab('calc'))
grafik_btn.configure(command=lambda: tab_manager.switch_tab('grafik'))
settings_btn.configure(command=lambda: tab_manager.switch_tab('settings'))
save_btn.configure(command=lambda: tab_manager.switch_tab('save'))

root.mainloop()